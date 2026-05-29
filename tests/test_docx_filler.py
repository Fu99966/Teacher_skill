from __future__ import annotations

from docx import Document

from teacher_agent.docx_filler import fill_docx_template


def _paragraph_texts(path):
    document = Document(path)
    return [paragraph.text for paragraph in document.paragraphs]


def test_replaces_placeholder_template(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "out.docx"
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    document.add_paragraph("{{teaching_goals}}")
    document.add_paragraph("{{teaching_process}}")
    document.save(template)

    report = fill_docx_template(
        template,
        {
            "lesson_title": "桂林山水",
            "teaching_goals": "理解课文内容",
            "teaching_process": "导入\n探究\n总结",
        },
        output,
    )

    assert _paragraph_texts(output) == ["桂林山水", "理解课文内容", "导入\n探究\n总结"]
    assert report.remaining_placeholders == []


def test_empty_string_does_not_clear_placeholder(tmp_path):
    template = tmp_path / "empty-placeholder.docx"
    output = tmp_path / "empty-placeholder-out.docx"
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    document.save(template)

    report = fill_docx_template(template, {"lesson_title": ""}, output)

    assert Document(output).paragraphs[0].text == "{{lesson_title}}"
    assert "lesson_title" not in report.filled_fields
    assert "lesson_title" in report.empty_fields
    assert "lesson_title" in report.skipped_empty_fields


def test_chinese_placeholder_can_be_filled(tmp_path):
    template = tmp_path / "chinese-placeholder.docx"
    output = tmp_path / "chinese-placeholder-out.docx"
    document = Document()
    document.add_paragraph("{{教学目标}}")
    document.save(template)

    report = fill_docx_template(template, {"教学目标": "理解课文内容"}, output)

    assert Document(output).paragraphs[0].text == "理解课文内容"
    assert report.filled_fields == ["教学目标"]


def test_table_label_fills_blank_right_cell(tmp_path):
    template = tmp_path / "table.docx"
    output = tmp_path / "out.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "教学目标"
    table.cell(0, 1).text = ""
    document.save(template)

    report = fill_docx_template(template, {"teaching_goals": "目标一\n目标二"}, output)
    result = Document(output)

    assert result.tables[0].cell(0, 0).text == "教学目标"
    assert result.tables[0].cell(0, 1).text == "目标一\n目标二"
    assert report.table_fields_filled == ["teaching_goals"]


def test_empty_string_does_not_clear_table_target(tmp_path):
    template = tmp_path / "empty-table.docx"
    output = tmp_path / "empty-table-out.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "教学目标"
    table.cell(0, 1).text = "请填写教学目标"
    document.save(template)

    report = fill_docx_template(template, {"teaching_goals": ""}, output)
    result = Document(output)

    assert result.tables[0].cell(0, 1).text == "请填写教学目标"
    assert "teaching_goals" not in report.filled_fields
    assert "teaching_goals" in report.empty_fields


def test_unknown_chinese_table_label_can_be_filled(tmp_path):
    template = tmp_path / "unknown-table.docx"
    output = tmp_path / "unknown-table-out.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "教学评价"
    table.cell(0, 1).text = ""
    document.save(template)

    report = fill_docx_template(template, {"教学评价": "采用课堂观察和任务单评价。"}, output)
    result = Document(output)

    assert result.tables[0].cell(0, 1).text == "采用课堂观察和任务单评价。"
    assert "教学评价" in report.table_fields_filled


def test_mixed_template_fills_placeholder_and_table(tmp_path):
    template = tmp_path / "mixed.docx"
    output = tmp_path / "out.docx"
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "作业设计"
    table.cell(0, 1).text = ""
    document.save(template)

    report = fill_docx_template(template, {"lesson_title": "观潮", "homework": "分层作业"}, output)
    result = Document(output)

    assert result.paragraphs[0].text == "观潮"
    assert result.tables[0].cell(0, 1).text == "分层作业"
    assert set(report.filled_fields) == {"lesson_title", "homework"}


def test_cross_run_placeholder_is_replaced(tmp_path):
    template = tmp_path / "runs.docx"
    output = tmp_path / "out.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("{{lesson")
    paragraph.add_run("_title}}")
    document.save(template)

    report = fill_docx_template(template, {"lesson_title": "跨 run 课题"}, output)

    assert Document(output).paragraphs[0].text == "跨 run 课题"
    assert report.remaining_placeholders == []


def test_missing_field_keeps_placeholder_and_reports(tmp_path):
    template = tmp_path / "missing.docx"
    output = tmp_path / "out.docx"
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    document.save(template)

    report = fill_docx_template(template, {}, output)

    assert Document(output).paragraphs[0].text == "{{lesson_title}}"
    assert report.missing_fields == ["lesson_title"]
    assert report.remaining_placeholders == ["lesson_title"]


def test_blank_template_risk_is_reported(tmp_path):
    template = tmp_path / "blank-risk.docx"
    output = tmp_path / "blank-risk-out.docx"
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    document.add_paragraph("{{teaching_goals}}")
    document.add_paragraph("{{teaching_process}}")
    document.save(template)

    report = fill_docx_template(
        template,
        {"lesson_title": "", "teaching_goals": "", "teaching_process": ""},
        output,
    )

    assert report.filled_non_empty_count == 0
    assert report.errors
    assert "空白模板" in report.errors[0]


def test_header_table_label_can_be_filled(tmp_path):
    template = tmp_path / "header-table.docx"
    output = tmp_path / "header-table-out.docx"
    document = Document()
    table = document.sections[0].header.add_table(rows=1, cols=2, width=1000000)
    table.cell(0, 0).text = "课题"
    table.cell(0, 1).text = ""
    document.save(template)

    report = fill_docx_template(template, {"lesson_title": "页眉课题"}, output)
    filled = Document(output)

    assert filled.sections[0].header.tables[0].cell(0, 1).text == "页眉课题"
    assert "lesson_title" in report.table_fields_filled

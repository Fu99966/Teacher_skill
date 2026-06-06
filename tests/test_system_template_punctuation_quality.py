from __future__ import annotations

import re

from docx import Document

from test_stm32_smart_car_prompt_generation import _run_export


def _docx_text(path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_stm32_system_template_has_clean_final_word_punctuation(monkeypatch, tmp_path):
    output_path, export = _run_export(monkeypatch, tmp_path, template_mode="system")
    text = _docx_text(output_path)
    compact_lines = re.sub(r"[ \t]+", " ", text)

    assert "本项目共 32 课时，围绕“完成一辆基于 STM32 的智能小车设计与调试”展开。" in compact_lines
    assert "围绕 完成" not in text
    assert "“ STM32" not in text
    assert "调试 展开。 ”" not in text
    assert "未命名课题" not in text
    assert "帮我生成一份" not in text
    assert "学生在真实智能小车项目实践中完成设计、检查、修改和展示" in text
    assert export["output_name"]

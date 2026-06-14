from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

from docx import Document

from teacher_agent.sample_template import create_sample_template
from teacher_agent.web_app import TeacherAgentHandler
from teacher_agent.workflow import TeacherWorkflow
import teacher_agent.web_app as web_app


FIELDS = {
    "lesson_title": "传感器基础",
    "subject": "物联网",
    "grade": "24级物联网班",
    "class_type": "新授课",
    "class_hour": "2课时",
    "teaching_goals": "理解传感器基础概念。",
    "teaching_key_difficult": "重点：传感器分类。难点：信号转换。",
    "teaching_aids": "传感器实物、任务单。",
    "teaching_process": "导入、探究、练习、总结。",
    "teaching_method": "任务驱动法、演示教学法。",
    "homework": "完成传感器分类表。",
    "reflection": "关注学生对信号转换的理解。",
}


def test_repeated_exports_never_overwrite_same_second(tmp_path):
    template = tmp_path / "sample_lesson_template.docx"
    create_sample_template(template)
    output_dir = tmp_path / "outputs"
    preview_dir = tmp_path / "previews"
    workflow = TeacherWorkflow()

    first = workflow.export_document(FIELDS, template, output_dir, preview_dir)
    second = workflow.export_document(FIELDS, template, output_dir, preview_dir)

    assert first["output_name"] != second["output_name"]
    assert (output_dir / first["output_name"]).exists()
    assert (output_dir / second["output_name"]).exists()
    assert "传感器基础" in first["output_name"]
    assert "传感器基础" in second["output_name"]


def test_same_named_template_uploads_get_distinct_paths(monkeypatch, tmp_path):
    upload_dir = tmp_path / "uploads"
    monkeypatch.setattr(web_app, "UPLOAD_DIR", upload_dir)
    handler = object.__new__(TeacherAgentHandler)
    first_docx = BytesIO()
    first_document = Document()
    first_document.add_paragraph("第一份模板")
    first_document.save(first_docx)
    second_docx = BytesIO()
    second_document = Document()
    second_document.add_paragraph("第二份模板")
    second_document.save(second_docx)

    first = handler._save_template(
        {"template": SimpleNamespace(filename="学校教案模板.docx", file=BytesIO(first_docx.getvalue()))}
    )
    second = handler._save_template(
        {"template": SimpleNamespace(filename="学校教案模板.docx", file=BytesIO(second_docx.getvalue()))}
    )

    assert first != second
    assert Document(str(first)).paragraphs[0].text == "第一份模板"
    assert Document(str(second)).paragraphs[0].text == "第二份模板"
    assert first.name.endswith(".docx")
    assert second.name.endswith(".docx")

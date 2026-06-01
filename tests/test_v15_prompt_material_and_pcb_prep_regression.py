from __future__ import annotations

import json
import re
import threading
import urllib.error
import urllib.request
from http.server import ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document

import teacher_agent.web_app as web_app


AGENT_REQUEST = "帮我生成一份 24级物联网 PCB板设计的 32课时的教案"


def _post_json(base_url: str, path: str, payload: dict[str, Any]) -> tuple[int, dict[str, Any]]:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        base_url + path,
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    return _open_json(request)


def _post_multipart(base_url: str, path: str, fields: dict[str, str]) -> tuple[int, dict[str, Any]]:
    boundary = "----teacher-skill-v15-material-regression"
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("ascii"))
        chunks.append(value.encode("utf-8"))
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode("ascii"))
    request = urllib.request.Request(
        base_url + path,
        data=b"".join(chunks),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    return _open_json(request)


def _open_json(request: urllib.request.Request) -> tuple[int, dict[str, Any]]:
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            return response.status, json.loads(response.read().decode("utf-8") or "{}")
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            data = json.loads(raw or "{}")
        except json.JSONDecodeError:
            data = {"error": raw}
        return exc.code, data


def _patch_web_paths(monkeypatch, tmp_path: Path) -> Path:
    output_dir = tmp_path / "outputs"
    template_dir = tmp_path / "templates"
    monkeypatch.setattr(web_app, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(web_app, "UPLOAD_DIR", output_dir / "uploads")
    monkeypatch.setattr(web_app, "PREVIEW_DIR", output_dir / "previews")
    monkeypatch.setattr(web_app, "HISTORY_DB", output_dir / "history.sqlite3")
    monkeypatch.setattr(web_app, "AGENT_MEMORY_DB", output_dir / "memory.sqlite3")
    monkeypatch.setattr(web_app, "TEMPLATE_DIR", template_dir)
    monkeypatch.setattr(web_app, "SAMPLE_TEMPLATE", template_dir / "sample_lesson_template.docx")
    return output_dir


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    parts.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_prompt_like_material_is_removed_and_pcb_project_prep_is_specific(monkeypatch, tmp_path):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    output_dir = _patch_web_paths(monkeypatch, tmp_path)

    server = ThreadingHTTPServer(("127.0.0.1", 0), web_app.TeacherAgentHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base_url = f"http://127.0.0.1:{server.server_address[1]}"

    try:
        preview_status, preview = _post_json(base_url, "/api/agent-preview", {"agent_request": AGENT_REQUEST})
        assert preview_status == 200, preview
        assert preview["agent_task"]["title"] in {"PCB板设计", "PCB设计"}
        assert not preview["needs_input"], preview

        run_status, data = _post_multipart(
            base_url,
            "/api/agent-run",
            {
                "agent_request": AGENT_REQUEST,
                "template_mode": "system",
                "strict_ai": "0",
            },
        )
        assert run_status == 200, data
        assert data["generation_backend"] == "local_fallback"
        fields = data["fields"]
        assert fields["lesson_title"] in {"PCB板设计", "PCB设计"}

        export_status, export = _post_json(
            base_url,
            "/api/export",
            {
                "template_id": data["template_id"],
                "fields": fields,
                "request_context": data.get("agent_task") or {},
                "generation_backend": data.get("generation_backend"),
            },
        )
        assert export_status == 200, export
        output_path = output_dir / export["output_name"]
        assert output_path.exists(), export

        text = _docx_text(output_path)
        compact = re.sub(r"\s+", " ", text)

        assert "课题 PCB板设计" in compact or "课题 PCB 板设计" in compact or "课题 PCB设计" in compact
        assert "帮我生成一份" not in compact
        assert "24级物联网 PCB板设计的 32课时的教案" not in compact
        assert "教材依据：帮我生成" not in compact
        assert "物联网应用技术专业课程要求" in compact
        assert "PCB设计项目任务书" in compact
        assert "EDA软件操作规范" in compact

        prep_keywords = ["计算机机房", "EDA", "PCB示例", "原理图素材", "封装库", "DRC", "项目任务单", "评价表"]
        hits = [keyword for keyword in prep_keywords if keyword in compact]
        assert len(hits) >= 5, hits

        assert "教学方法的运用" in compact
        assert "项目教学法" in compact or "任务驱动" in compact
        assert "32课时" in compact or "32 课时" in compact
        assert "Gerber" in compact
        assert "DRC" in compact
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)

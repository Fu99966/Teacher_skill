from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from teacher_agent.agent_core.memory import AgentMemoryStore
from teacher_agent.agent_core.repair import repair_state
from teacher_agent.agent_core.state import AgentRunState
from teacher_agent.agent_observer import build_teacher_diagnostic_report
from teacher_agent.lesson_generator import JSON_FIELD_NAMES, draft_lesson_document_fields_with_source
from teacher_agent.material_ingestion import extract_material_bytes
from teacher_agent.rag_context import build_knowledge_context
from teacher_agent.template_parser import analyze_template
from teacher_agent.template_profile import TemplateProfileStore


def _docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _placeholder_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    document.add_paragraph("{{teaching_process}}")
    document.add_paragraph("{{teaching_method}}")
    document.save(str(path))


def test_teacher_diagnostic_report_summarizes_writes_and_failures():
    report = build_teacher_diagnostic_report(
        template_analysis={"mapped_fields": ["lesson_title", "teaching_process", "teaching_method"]},
        fill_report={
            "field_write_counts": {"lesson_title": 1, "teaching_method": 1},
            "empty_fields": ["teaching_process"],
            "warnings": ["主要教学内容为空"],
        },
        evaluation_report={"passed": True},
        fields={"lesson_title": "PCB板设计", "teaching_method": "项目教学法", "teaching_process": ""},
    )

    data = report.to_dict()
    assert data["status"] == "needs_review"
    assert len(data["written_fields"]) == 2
    assert data["unwritten_fields"][0]["field"] == "teaching_process"
    assert "识别" in data["summary"]


def test_template_profile_reuses_successful_mapping(tmp_path):
    template = tmp_path / "school.docx"
    _placeholder_template(template)
    analysis = analyze_template(template)

    store = TemplateProfileStore(tmp_path / "profiles")
    profile_id = store.template_fingerprint(template)
    first = store.get_or_create(profile_id, analysis)
    assert first["profile_hit"] is False

    store.save_successful_mapping(
        profile_id,
        analysis.get("table_mappings", {}),
        {"filled_non_empty_count": 3, "field_write_counts": {"lesson_title": 1}},
        mapped_fields=analysis["mapped_fields"],
        repeat_fill_mode="first_only",
    )
    second = store.get_or_create(profile_id, analysis)
    assert second["profile_hit"] is True
    applied = store.apply_profile({"mapped_fields": [], "table_mappings": {}}, second)
    assert "lesson_title" in applied["mapped_fields"]
    assert applied["template_profile"]["profile_hit"] is True


def test_material_docx_upload_feeds_lightweight_rag():
    extraction = extract_material_bytes(
        "教材.docx",
        _docx_bytes("课程标准：学生需要掌握传感器分类、实验步骤和评价要求。"),
    )
    assert "传感器分类" in extraction.text

    context = build_knowledge_context(
        extraction.text,
        subject="物联网",
        title="传感器基础",
        class_type="实验课",
        teaching_style="探究式",
    )
    assert context.chunks
    assert context.lesson_pattern["key"] == "experiment_lesson"
    assert "实验" in context.enhanced_material(extraction.text)


def test_course_type_fallbacks_are_not_one_generic_template(monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    fields, backend = draft_lesson_document_fields_with_source(
        "物联网",
        "24级物联网班",
        "传感器数据采集",
        "",
        "2课时",
        "实验课",
        "探究式",
        "常规混合水平",
        "标准",
        JSON_FIELD_NAMES,
        False,
    )
    assert backend == "local_fallback"
    assert "实验" in fields["teaching_process"]
    assert "数据" in fields["teaching_process"] or "现象" in fields["teaching_process"]
    assert "探究式教学" in fields["teaching_method"]

    review_fields, _ = draft_lesson_document_fields_with_source(
        "物联网",
        "24级物联网班",
        "传感器单元复习",
        "",
        "2课时",
        "复习课",
        "常规启发式",
        "常规混合水平",
        "标准",
        JSON_FIELD_NAMES,
        False,
    )
    assert "错因分析" in review_fields["teaching_process"]
    assert "分层训练" in review_fields["teaching_method"]


def test_material_title_does_not_override_requested_title(monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    fields, _ = draft_lesson_document_fields_with_source(
        "物联网",
        "24级物联网班",
        "PCB板设计",
        "教材示例：《观潮》是一篇课文。",
        "32课时",
        "项目实训课",
        "项目式教学",
        "常规混合水平",
        "标准",
        JSON_FIELD_NAMES,
        False,
    )
    assert fields["lesson_title"] == "PCB板设计"
    assert "观潮" not in fields["lesson_title"]


def test_repair_loop_backfills_method_and_removes_prompt_leak(monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    state = AgentRunState(
        session_id="s1",
        status="failed",
        task={
            "subject": "物联网",
            "grade": "24级物联网班",
            "title": "PCB板设计",
            "class_hour": "32课时",
            "class_type": "项目实训课",
            "raw_text": "帮我生成一份 PCB板设计教案",
            "material": "",
        },
        current_node="evaluate_delivery",
        next_action="",
        template_analysis={"mapped_fields": ["lesson_title", "teaching_process", "teaching_method"]},
        fields={
            "lesson_title": "",
            "teaching_process": "帮我生成一份 PCB板设计教案\n一、项目总任务：完成PCB设计。",
            "teaching_method": "",
        },
    )
    repaired = repair_state(state)
    assert repaired.status == "fields_generated"
    assert repaired.fields
    assert repaired.fields["lesson_title"] == "PCB板设计"
    assert "帮我生成" not in repaired.fields["teaching_process"]
    assert repaired.fields["teaching_method"].strip()


def test_agent_memory_store_remembers_teacher_edits(tmp_path):
    store = AgentMemoryStore(tmp_path / "memory.sqlite3")
    store.remember_teacher_edit(
        template_id="school-template",
        task={"subject": "物联网", "grade": "24级物联网班", "title": "PCB板设计", "class_type": "实训课"},
        fields={"teaching_process": "老师修改后的项目过程", "teaching_method": "任务驱动法"},
    )
    examples = store.find_teacher_edit_examples(subject="物联网", title="PCB板设计", template_id="school-template")
    assert examples
    assert examples[0]["fields"]["teaching_method"] == "任务驱动法"

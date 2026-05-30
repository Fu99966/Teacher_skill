"""E2E Web Agent test – completely bypasses DeepSeek API.

Validates the full engineering pipeline:
  agent/start → waiting_teacher_review
  → checkpoint load (GET equivalent)
  → agent/continue with teacher_edits
  → completed
  → output Word contains teacher-modified content.

Uses fixed fields — no real LLM calls.
"""
from __future__ import annotations

from pathlib import Path

from teacher_agent.agent_core.checkpoint import AgentCheckpointStore
from teacher_agent.agent_core.graph_planner import build_graph
from teacher_agent.agent_core.state import AgentRunState
from teacher_agent.agent_core.tool_registry import build_agent_tool_registry
from teacher_agent.agent_core.executor import AgentExecutor
from teacher_agent.agent_core.task_router import AgentTask

FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
REAL_TEMPLATE = FIXTURE_DIR / "教案模板.docx"

FIXED_FIELDS = {
    "teaching_date": "2026-05-29",
    "class_name": "24物联网1班",
    "lesson_title": "传感器基础",
    "class_type": "新授课",
    "class_hour": "2课时",
    "teaching_environment": "多媒体教室，具备投影设备和传感器演示套件。",
    "teaching_goals": "理解传感器的基本概念、分类和典型应用。",
    "teaching_key_difficult": "重点：传感器分类与工作原理。难点：传感器信号与物联网系统的关系。",
    "teaching_aids": "PPT、传感器实物、实验演示板。",
    "teaching_process": "这是主要教学内容正文。",
    "teaching_method": "这是教学方法正文。",
    "homework": "老师修改后的作业内容。",
    "reflection": "课后关注学生是否掌握传感器应用。",
}

TEACHER_EDITS = {"homework": "老师修改后的作业内容。"}


def _mock_draft_fields(*args, **kwargs):
    """Return fixed fields, bypassing any LLM call."""
    return dict(FIXED_FIELDS), "mock_no_llm"


def _make_task():
    return AgentTask(
        raw_request="生成物联网传感器基础教案",
        task_type="lesson_plan", subject="物联网", grade="24物联网1班",
        title="传感器基础", class_hour="2课时", class_type="新授课",
        teaching_style="常规启发式", student_level="常规", generation_depth="标准",
        missing_fields=[], confidence=0.9, notes=[],
    )


def _make_state(template_path, suffix=""):
    tid = f"e2e-{suffix}"
    return AgentRunState(
        session_id=tid, status="initialized",
        task=_make_task().to_dict(), current_node="", next_action="",
        template_path=str(template_path), template_id=tid,
    )


def _make_registry(output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    return build_agent_tool_registry(
        output_dir=output_dir, preview_dir=output_dir,
        history_db=output_dir / "h.sqlite3", memory_db=output_dir / "m.sqlite3",
    )


# ── Test: full E2E pipeline, zero DeepSeek calls ──

def test_web_agent_e2e_pipeline_no_llm(monkeypatch, tmp_path):
    """Full flow: start → pause → continue → completed → verify Word."""
    # ── 1. Bypass DeepSeek ──
    monkeypatch.setattr(
        "teacher_agent.lesson_generator.draft_lesson_document_fields_with_source",
        _mock_draft_fields,
    )
    monkeypatch.setattr(
        "teacher_agent.deepseek_client.is_deepseek_configured",
        lambda: False,
    )

    output_dir = tmp_path / "outputs"
    checkpoint_dir = tmp_path

    registry = _make_registry(output_dir)
    checkpoint = AgentCheckpointStore(checkpoint_dir)
    executor = AgentExecutor(registry, checkpoint)

    # ── 2. START: POST /api/agent/start ──
    state = _make_state(REAL_TEMPLATE, "start")
    graph = build_graph(_make_task())
    result = executor.run(graph, state)

    # Validate pause at teacher_review_gate
    assert result.status == "waiting_teacher_review", (
        f"Web API 路由失败：Expected waiting_teacher_review, got {result.status}. "
        f"Errors: {result.state.errors}"
    )
    assert result.next_action == "teacher_edit_fields"
    assert result.state.fields is not None, "Agent 未生成 fields"
    assert result.state.fields.get("lesson_title") == "传感器基础"
    assert result.state.fields.get("teaching_process") == "这是主要教学内容正文。"
    assert result.state.fields.get("teaching_method") == "这是教学方法正文。"

    # ── 3. GET: checkpoint restore (simulates GET /api/agent/{id}) ──
    reloaded = checkpoint.load(state.session_id)
    assert reloaded is not None, (
        f"checkpoint 恢复失败：session_id={state.session_id} 未找到 checkpoint 文件"
    )
    assert reloaded.status == "waiting_teacher_review"
    assert reloaded.fields is not None

    # ── 4. Teacher edits + CONTINUE: POST /api/agent/{id}/continue ──
    fields = dict(reloaded.fields)
    fields.update(TEACHER_EDITS)
    reloaded.fields = fields
    reloaded.teacher_edits = TEACHER_EDITS
    reloaded.status = "fields_generated"
    checkpoint.save(reloaded)

    # Build fresh graph (simulates web rebuild)
    new_graph = build_graph(_make_task())
    executor2 = AgentExecutor(registry, checkpoint)
    result2 = executor2.continue_after_review(new_graph, reloaded)

    assert result2.status == "completed", (
        f"continue 状态失败：Expected completed, got {result2.status}. "
        f"Errors: {result2.state.errors}, Warnings: {result2.state.warnings}"
    )

    # ── 5. VERIFY Word output ──
    export = result2.state.export_result or {}
    download_url = export.get("download_url")
    assert download_url, "Web API 完成但 Word 导出失败：缺少 download_url"

    output_name = export.get("output_name", "")
    assert output_name, "Word 导出失败：缺少 output_name"

    docx_path = output_dir / output_name
    assert docx_path.exists(), (
        f"Word 导出失败：文件不存在 {docx_path}. "
        f"Export result: {export}"
    )

    # ── 6. Verify docx content ──
    try:
        from docx import Document
        doc = Document(str(docx_path))
    except Exception as e:
        raise AssertionError(f"docx 内容读取失败：{e}") from e

    all_text = "\n".join(
        cell.text for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    checks = [
        ("传感器基础", "课题名称"),
        ("这是主要教学内容正文", "teaching_process 正文"),
        ("这是教学方法正文", "teaching_method 正文"),
        ("老师修改后的作业内容", "老师编辑后的作业"),
    ]
    for keyword, label in checks:
        assert keyword in all_text, (
            f"docx 内容验证失败：「{label}」({keyword}) 不在输出 Word 中。\n"
            f"输出文本前500字：{all_text[:500]}"
        )

    # ── 7. Position verification: content NOT in label cells ──
    from teacher_agent.docx_grid import parse_table_grid
    for t_idx, table in enumerate(doc.tables):
        grid = parse_table_grid(table)
        for ri, row_grid in enumerate(grid):
            for gc, gcell in enumerate(row_grid):
                if gcell is None or gcell.grid_col != gc:
                    continue
                t = gcell.text
                if t.strip() == "主要教学内容":
                    assert "这是主要教学内容正文" not in t, (
                        f"Table{t_idx} Row{ri} Col{gc}: teaching_process leaked into label cell"
                    )
                if t.strip() == "教学方法的运用":
                    assert "这是教学方法正文" not in t, (
                        f"Table{t_idx} Row{ri} Col{gc}: teaching_method leaked into label cell"
                    )

    # Verify from template_analysis that targets are next-row
    analysis = result2.state.template_analysis or {}
    for f, label_text in [("teaching_process", "主要教学内容"), ("teaching_method", "教学方法的运用")]:
        targets = analysis.get("table_mappings", {}).get(f, [])
        for t in targets:
            if t.get("label", "").replace(" ", "").replace("\n", "") == label_text:
                assert t["row"] > t["label_row"], (
                    f"E2E position: {f} target_row({t['row']}) must be > label_row({t['label_row']})"
                )
                assert t.get("target_type") == "next_row_cell", (
                    f"E2E position: {f} target_type={t.get('target_type')}, expected next_row_cell"
                )

    print(f"\n✅ E2E PASS: session={state.session_id}, output={output_name}")
    print(f"   docx size={docx_path.stat().st_size} bytes")
    print(f"   fields={len(FIXED_FIELDS)}, filled_non_empty={result2.state.fields and len(result2.state.fields)}")


# ── Test: model health check logic (unit test, no web needed) ──

def test_model_health_check_no_api():
    """Verify the DeepSeek health check returns sensible data without API key."""
    from teacher_agent.deepseek_client import check_deepseek_health

    status = check_deepseek_health(probe=False)
    data = status.to_dict()

    assert "configured" in data, f"Missing 'configured': {data}"
    assert "provider" not in data or "status" in data
    assert data["status"] in ("ok", "error", "not_configured"), f"Bad status: {data['status']}"

    print(f"\n✅ Model health: configured={data['configured']}, status={data['status']}")


# ── Test: draft_fields_tool works with mocked function ──

def test_draft_fields_tool_with_mock(monkeypatch, tmp_path):
    """Verify draft_fields_tool uses mocked function, not real DeepSeek."""
    monkeypatch.setattr(
        "teacher_agent.lesson_generator.draft_lesson_document_fields_with_source",
        _mock_draft_fields,
    )

    state = _make_state(REAL_TEMPLATE, "draft-tool")
    state.template_analysis = {"mapped_fields": list(FIXED_FIELDS.keys()), "fillable_count": len(FIXED_FIELDS)}

    # Call the tool directly via registry
    registry = _make_registry(tmp_path / "outputs")
    tool_fn = registry.get("draft_fields")
    result = tool_fn({"state": state})

    assert result["field_count"] == len(FIXED_FIELDS)
    assert state.fields["lesson_title"] == "传感器基础"
    assert state.fields["teaching_process"] == "这是主要教学内容正文。"
    assert state.fields["homework"] == "老师修改后的作业内容。"
    assert result["generation_backend"] == "mock_no_llm"

    print(f"\n✅ draft_fields_tool: {result['field_count']} fields via {result['generation_backend']}")

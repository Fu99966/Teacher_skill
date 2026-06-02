from __future__ import annotations

import json
import re
import threading
import urllib.error
import urllib.request
from http.server import ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document

import teacher_agent.web_app as web_app
from teacher_agent.docx_grid import find_cell_by_grid, parse_table_grid


AGENT_REQUEST = "帮我生成一份 24级物联网班 PCB板设计课的 32课时的教案"
FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
REAL_TEMPLATE = FIXTURE_DIR / "教案模板.docx"


def _post_json(base_url: str, path: str, payload: dict[str, Any]) -> tuple[int, dict[str, Any]]:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        base_url + path,
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    return _open_json(request)


def _post_multipart(
    base_url: str,
    path: str,
    fields: dict[str, str],
    files: dict[str, tuple[str, bytes, str]],
) -> tuple[int, dict[str, Any]]:
    boundary = "----teacher-skill-uploaded-method-cell"
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("ascii"))
        chunks.append(value.encode("utf-8"))
        chunks.append(b"\r\n")
    for name, (filename, content, content_type) in files.items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(
            f'Content-Disposition: form-data; name="{name}"; filename="{filename}"\r\n'.encode("utf-8")
        )
        chunks.append(f"Content-Type: {content_type}\r\n\r\n".encode("ascii"))
        chunks.append(content)
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode("ascii"))
    request = urllib.request.Request(
        base_url + path,
        data=b"".join(chunks),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    return _open_json(request)


def _open_json(request: urllib.request.Request) -> tuple[int, dict[str, Any]]:
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            return response.status, json.loads(response.read().decode("utf-8") or "{}")
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            data = json.loads(raw or "{}")
        except json.JSONDecodeError:
            data = {"error": raw}
        return exc.code, data


def _patch_web_paths(monkeypatch, tmp_path: Path) -> Path:
    output_dir = tmp_path / "outputs"
    template_dir = tmp_path / "templates"
    monkeypatch.setattr(web_app, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(web_app, "UPLOAD_DIR", output_dir / "uploads")
    monkeypatch.setattr(web_app, "PREVIEW_DIR", output_dir / "previews")
    monkeypatch.setattr(web_app, "HISTORY_DB", output_dir / "history.sqlite3")
    monkeypatch.setattr(web_app, "AGENT_MEMORY_DB", output_dir / "memory.sqlite3")
    monkeypatch.setattr(web_app, "TEMPLATE_DIR", template_dir)
    monkeypatch.setattr(web_app, "SAMPLE_TEMPLATE", template_dir / "sample_lesson_template.docx")
    return output_dir


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs]
    parts.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _parallel_content_cells(document: Document):
    for table in document.tables:
        grid = parse_table_grid(table)
        for row_index, row_grid in enumerate(grid[:-1]):
            process_label = None
            method_label = None
            for grid_col, gcell in enumerate(row_grid):
                if gcell is None or gcell.grid_col != grid_col:
                    continue
                if "主要教学内容" in gcell.text:
                    process_label = gcell
                if "教学方法的运用" in gcell.text:
                    method_label = gcell
            if process_label is None or method_label is None:
                continue
            process_cell = find_cell_by_grid(table, row_index + 1, process_label.grid_col, process_label.physical_col)
            method_cell = find_cell_by_grid(table, row_index + 1, method_label.grid_col, method_label.physical_col)
            if process_cell is not None and method_cell is not None:
                yield process_cell, method_cell


def test_uploaded_template_teaching_method_writes_right_next_row_cell(monkeypatch, tmp_path):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    output_dir = _patch_web_paths(monkeypatch, tmp_path)

    server = ThreadingHTTPServer(("127.0.0.1", 0), web_app.TeacherAgentHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base_url = f"http://127.0.0.1:{server.server_address[1]}"

    try:
        preview_status, preview = _post_json(base_url, "/api/agent-preview", {"agent_request": AGENT_REQUEST})
        assert preview_status == 200, preview

        run_status, data = _post_multipart(
            base_url,
            "/api/agent-run",
            {
                "agent_request": AGENT_REQUEST,
                "template_mode": "upload",
                "strict_ai": "0",
            },
            {
                "template": (
                    REAL_TEMPLATE.name,
                    REAL_TEMPLATE.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert run_status == 200, data

        export_status, export = _post_json(
            base_url,
            "/api/export",
            {
                "template_id": data["template_id"],
                "fields": data["fields"],
                "request_context": data.get("agent_task") or {},
                "generation_backend": data.get("generation_backend"),
            },
        )
        assert export_status == 200, export
        output_path = output_dir / export["output_name"]
        assert output_path.exists(), export

        text = _docx_text(output_path)
        compact = re.sub(r"\s+", " ", text)
        assert "课题" in compact
        assert "PCB板设计" in compact or "PCB 板设计" in compact
        assert "主要教学内容" in compact
        assert "教学方法的运用" in compact
        assert "项目教学法" in compact or "任务驱动法" in compact
        assert "巡回指导" in compact
        assert "作品展示评价" in compact
        assert "32课时" in compact or "32 课时" in compact
        assert "Gerber" in compact
        assert "DRC" in compact
        assert "教材依据：帮我生成" not in compact

        doc = Document(str(output_path))
        passed_locations = 0
        for process_cell, method_cell in _parallel_content_cells(doc):
            process_text = process_cell.text
            method_text = method_cell.text
            method_hits = sum(
                keyword in method_text
                for keyword in ["项目教学法", "任务驱动法", "演示教学法", "分组协作", "巡回指导", "作品展示评价"]
            )
            if "本项目共" in process_text and ("32课时" in process_text or "32 课时" in process_text):
                assert "Gerber" in process_text
                assert "DRC" in process_text
                assert method_hits >= 4, method_text
                passed_locations += 1
        assert passed_locations >= 2

        fill_report = export["fill_report"]
        assert fill_report["field_write_counts"]["teaching_method"] >= 2
        method_report = fill_report["field_reports"]["teaching_method"]
        assert method_report["label"] == "教学方法的运用"
        assert method_report["required"] is True
        assert method_report["target_type"] == "next_row_cell"
        assert method_report["status"] == "passed"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)

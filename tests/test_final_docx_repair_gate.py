from __future__ import annotations

from docx import Document

from teacher_agent.agent_core.executor import AgentExecutor
from teacher_agent.agent_core.evaluator import evaluate_delivery
from teacher_agent.agent_core.graph_planner import GraphNode
from teacher_agent.agent_core.repair import diagnose_failure
from teacher_agent.agent_core.state import AgentRunState
from teacher_agent.agent_core.tool_registry import build_agent_tool_registry
from teacher_agent.agent_core.tool_spec import ToolRegistry, ToolSpec
from teacher_agent.agent_observer import build_teacher_diagnostic_report
from teacher_agent.output_quality import inspect_docx_delivery_quality


def test_delivery_fails_when_final_docx_quality_report_fails():
    report = evaluate_delivery(
        fields={"lesson_title": "PCB板设计", "teaching_process": "项目过程", "teaching_method": "项目教学法"},
        output_path=None,
        download_url="/download/output.docx",
        template_analysis={"mapped_fields": ["lesson_title", "teaching_process", "teaching_method"]},
        fill_report={
            "errors": [],
            "missing_fields": [],
            "remaining_placeholders": [],
            "filled_non_empty_count": 3,
            "field_write_counts": {"teaching_process": 1, "teaching_method": 1},
        },
        output_quality_report={
            "passed": False,
            "score": 76,
            "errors": ["检测到生成指令泄漏到 Word 正文。"],
            "checks": {"no_prompt_leak": False},
        },
    )

    assert report["passed"] is False
    final_check = next(check for check in report["delivery_checks"] if check["name"] == "final_docx_quality")
    assert final_check["passed"] is False
    assert "生成指令泄漏" in final_check["detail"]


def test_repair_diagnosis_reads_final_docx_quality_failures():
    state = AgentRunState(
        session_id="final-docx-diagnosis",
        status="failed",
        task={"title": "PCB板设计"},
        current_node="evaluate_delivery",
        next_action="",
        fields={"lesson_title": "PCB板设计"},
        export_result={
            "fill_report": {"filled_non_empty_count": 1},
            "output_quality_report": {
                "passed": False,
                "errors": ["教学方法的运用未写入非空内容。", "检测到生成指令泄漏到 Word 正文。"],
            },
        },
    )

    diagnosis = diagnose_failure(state)

    assert any("教学方法的运用未写入" in issue for issue in diagnosis["issues"])
    assert any("生成指令泄漏" in issue for issue in diagnosis["issues"])


def test_delivery_evaluator_goes_to_repair_instead_of_blind_retry(tmp_path):
    registry = build_agent_tool_registry(
        output_dir=tmp_path / "outputs",
        preview_dir=tmp_path / "previews",
        history_db=tmp_path / "history.sqlite3",
        memory_db=tmp_path / "memory.sqlite3",
    )

    spec = registry.get_spec("evaluate_delivery")

    assert spec is not None
    assert spec.retryable is False
    assert spec.critical is True


def test_final_docx_failure_triggers_field_repair_and_reexport(tmp_path):
    raw_request = "帮我生成一份 PCB板设计 32课时教案"
    export_calls: list[str] = []
    registry = ToolRegistry()

    def export_docx(context):
        state = context["state"]
        process = str((state.fields or {}).get("teaching_process") or "")
        export_calls.append(process)
        path = tmp_path / f"repair-{len(export_calls)}.docx"
        document = Document()
        document.add_paragraph(process)
        document.save(path)
        leaked = raw_request in process
        state.export_result = {
            "output_name": path.name,
            "download_url": f"/download/{path.name}",
            "output_path": str(path),
            "fill_report": {
                "errors": [],
                "missing_fields": [],
                "remaining_placeholders": [],
                "filled_non_empty_count": 3,
                "field_write_counts": {"teaching_process": 1, "teaching_method": 1},
            },
            "output_quality_report": {
                "passed": not leaked,
                "errors": ["检测到生成指令泄漏到 Word 正文。"] if leaked else [],
            },
        }
        return {"output_name": path.name}

    def delivery_gate(context):
        state = context["state"]
        export = state.export_result or {}
        report = evaluate_delivery(
            fields=state.fields or {},
            output_path=tmp_path / str(export.get("output_name")),
            download_url=export.get("download_url"),
            template_analysis=state.template_analysis,
            fill_report=export.get("fill_report"),
            output_quality_report=export.get("output_quality_report"),
        )
        state.evaluation_report = report
        if not report["passed"]:
            raise ValueError(report["summary"])
        return {"passed": True}

    registry.register("export_docx", export_docx, ToolSpec("export_docx", "export", retryable=False))
    registry.register("evaluate_delivery", delivery_gate, ToolSpec("evaluate_delivery", "evaluate", retryable=False))
    state = AgentRunState(
        session_id="final-docx-repair",
        status="fields_generated",
        task={
            "subject": "物联网",
            "grade": "24级物联网班",
            "title": "PCB板设计",
            "class_hour": "32课时",
            "class_type": "项目实训课",
            "raw_text": raw_request,
            "material": "",
        },
        current_node="",
        next_action="",
        template_analysis={"mapped_fields": ["lesson_title", "teaching_process", "teaching_method"]},
        fields={
            "lesson_title": "PCB板设计",
            "teaching_process": raw_request,
            "teaching_method": "项目教学法、任务驱动法。",
        },
        max_retries=1,
    )
    graph = [
        GraphNode("export_docx", "export_docx", "导出 Word"),
        GraphNode("evaluate_delivery", "evaluate_delivery", "交付检查"),
    ]

    result = AgentExecutor(registry).run(graph, state)

    assert result.status == "completed"
    assert len(export_calls) == 2
    assert raw_request in export_calls[0]
    assert raw_request not in export_calls[1]
    assert result.state.evaluation_report["passed"] is True


def test_final_docx_quality_requires_generated_field_content(tmp_path):
    path = tmp_path / "labels-only.docx"
    document = Document()
    for heading in ("课题", "教学目的", "重点难点", "主要教学内容", "教学方法的运用", "作业", "课后小记"):
        document.add_paragraph(heading)
    document.save(path)

    report = inspect_docx_delivery_quality(
        path,
        expected_fields={
            "lesson_title": "PCB板设计",
            "teaching_goals": "理解PCB设计流程与工程规范。",
            "teaching_key_difficult": "重点是布局布线，难点是DRC问题修正。",
            "teaching_process": "完成原理图、布局布线和Gerber输出。",
            "teaching_method": "项目教学法与任务驱动法。",
            "homework": "提交PCB设计文件。",
            "reflection": "关注学生工程规范意识。",
        },
    )

    assert report["passed"] is False
    assert report["checks"]["field_content_lesson_title"] is False
    assert any("lesson_title" in error for error in report["errors"])


def test_teacher_diagnostic_surfaces_final_docx_quality_errors():
    report = build_teacher_diagnostic_report(
        template_analysis={"mapped_fields": ["lesson_title", "teaching_process", "teaching_method"]},
        fill_report={
            "field_write_counts": {"lesson_title": 1, "teaching_process": 1, "teaching_method": 1},
            "errors": [],
            "warnings": [],
        },
        evaluation_report={"passed": True},
        output_quality_report={
            "passed": False,
            "errors": ["最终 Word 中未找到字段内容：teaching_method"],
        },
        fields={
            "lesson_title": "PCB板设计",
            "teaching_process": "项目过程",
            "teaching_method": "项目教学法",
        },
    )
    data = report.to_dict()

    assert data["status"] == "failed"
    assert any("teaching_method" in reason for reason in data["reasons"])
    assert any("最终 Word" in action for action in data["next_actions"])

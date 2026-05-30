from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]


def test_cli_fill_template_outputs_report(tmp_path):
    template = tmp_path / "template.docx"
    data = tmp_path / "data.json"
    output = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("{{lesson_title}}")
    doc.save(str(template))

    data.write_text('{"lesson_title": "桂林山水"}', encoding="utf-8")

    result = subprocess.run(
        [sys.executable, "-m", "teacher_agent.cli", "fill-template",
         "--template", str(template), "--data", str(data), "--output", str(output)],
        capture_output=True, text=True, encoding="utf-8", cwd=str(PROJECT_ROOT),
    )
    stdout = result.stdout

    assert result.returncode == 0, f"stderr: {result.stderr}\nstdout: {stdout}"
    assert output.exists(), f"Output file not created: {output}"
    assert "fill_report" in stdout
    assert '"success": true' in stdout
    assert Document(str(output)).paragraphs[0].text == "桂林山水"


def test_cli_generate_uses_fallback_and_fills_word(tmp_path, monkeypatch):
    monkeypatch.setattr("teacher_agent.lesson_generator.is_deepseek_configured", lambda: False)
    template = tmp_path / "template.docx"
    material = tmp_path / "material.md"
    output = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("{{lesson_title}}")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "教学目标"
    table.cell(0, 1).text = ""
    doc.save(str(template))
    material.write_text("教材内容", encoding="utf-8")

    result = subprocess.run(
        [sys.executable, "-m", "teacher_agent.cli", "generate",
         "--template", str(template),
         "--subject", "语文",
         "--grade", "四年级",
         "--title", "桂林山水",
         "--material-file", str(material),
         "--output", str(output),
         "--no-strict-ai"],
        capture_output=True, text=True, encoding="utf-8", cwd=str(PROJECT_ROOT),
    )
    stdout = result.stdout

    assert result.returncode == 0, f"stderr: {result.stderr}\nstdout: {stdout}"
    assert '"generation_backend": "local_fallback"' in stdout
    assert '"success": true' in stdout
    # CLI may output to PROJECT_ROOT/outputs/ when cwd=PROJECT_ROOT
    actual_output = output
    if not output.exists():
        alt = PROJECT_ROOT / "outputs" / output.name
        if alt.exists():
            actual_output = alt
    assert actual_output.exists(), f"Output not found at {output} or alt"
    assert Document(str(actual_output)).paragraphs[0].text == "桂林山水"
    assert Document(str(actual_output)).tables[0].cell(0, 1).text

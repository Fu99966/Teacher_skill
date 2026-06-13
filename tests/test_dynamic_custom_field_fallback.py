from __future__ import annotations

from docx import Document

from teacher_agent.docx_filler import fill_docx_template
from teacher_agent.lesson_generator import coerce_dynamic_fields, draft_lesson_document_fields_with_source
from teacher_agent.template_parser import analyze_template


def test_chinese_and_custom_template_fields_get_semantic_local_fallback(monkeypatch, tmp_path):
    monkeypatch.setattr("teacher_agent.lesson_generator.is_deepseek_configured", lambda: False)
    template = tmp_path / "custom-school-template.docx"
    output = tmp_path / "output.docx"
    document = Document()
    for field in (
        "课题",
        "教学目标",
        "教学目的与要求",
        "教材分析",
        "学习任务",
        "主要教学内容安排",
        "教学方法运用",
        "教学评价",
        "安全教育",
        "二次备课",
        "审批意见",
    ):
        document.add_paragraph(f"{{{{{field}}}}}")
    document.save(str(template))

    analysis = analyze_template(template)
    fields, backend = draft_lesson_document_fields_with_source(
        subject="物联网",
        grade="24级物联网班",
        title="传感器数据采集",
        material="教材重点：温湿度传感器、采样周期、数据误差分析和物联网平台上传。",
        class_hour="2课时",
        class_type="实验课",
        teaching_style="探究式",
        student_level="常规混合水平",
        generation_depth="标准",
        template_fields=analysis["mapped_fields"],
        strict_ai=False,
        template_context=analysis.get("field_context"),
    )

    assert backend == "local_fallback"
    assert set(fields) == set(analysis["mapped_fields"])
    assert fields["课题"] == "传感器数据采集"
    assert "知识目标" in fields["教学目标"]
    assert "知识目标" in fields["教学目的与要求"]
    assert "温湿度传感器" in fields["教材分析"]
    assert "任务" in fields["学习任务"]
    assert "实验" in fields["主要教学内容安排"]
    assert fields["教学方法运用"].strip()
    assert "评价" in fields["教学评价"]
    assert "安全" in fields["安全教育"] or "规范" in fields["安全教育"]
    assert "调整" in fields["二次备课"] or "关注" in fields["二次备课"]
    assert fields["审批意见"] == "待审批人填写。"
    assert all("生成“" not in value for value in fields.values())

    report = fill_docx_template(template, fields, output)
    assert not report.errors
    text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)
    assert "{{" not in text
    assert "传感器数据采集" in text
    assert "待审批人填写" in text


def test_known_english_custom_fields_keep_useful_local_fallback(monkeypatch):
    monkeypatch.setattr("teacher_agent.lesson_generator.is_deepseek_configured", lambda: False)
    dynamic_fields = ["lesson_title", "warm_up", "assessment", "safety_rules", "learning_evidence"]

    fields, backend = draft_lesson_document_fields_with_source(
        "科学",
        "五年级",
        "水的沸腾",
        "观察水温变化，记录实验数据并解释沸腾现象。",
        "1课时",
        "实验课",
        "探究式",
        "常规混合水平",
        "标准",
        dynamic_fields,
        False,
    )

    assert backend == "local_fallback"
    assert fields["lesson_title"] == "水的沸腾"
    assert fields["warm_up"].strip()
    assert "评价" in fields["assessment"]
    assert fields["safety_rules"].strip()
    assert "学习证据" in fields["learning_evidence"] or "成果" in fields["learning_evidence"]


def test_ai_cannot_fabricate_human_approval_fields():
    fields = coerce_dynamic_fields(
        {
            "审批意见": "AI生成：同意通过。",
            "教研组意见": "AI生成：建议实施。",
            "教学目标": "理解核心知识。",
        },
        ["审批意见", "教研组意见", "教学目标"],
    )

    assert fields["审批意见"] == "待审批人填写。"
    assert fields["教研组意见"] == "待教研组填写。"
    assert fields["教学目标"] == "理解核心知识。"

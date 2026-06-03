from __future__ import annotations

import json
import re
import threading
import urllib.error
import urllib.request
from http.server import ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document

import teacher_agent.web_app as web_app
from teacher_agent.agent_core.task_router import route_task
from teacher_agent.docx_grid import find_cell_by_grid, parse_table_grid
from teacher_agent.lesson_generator import infer_lesson_scope, is_stm32_smart_car_topic


AGENT_REQUEST = "帮我生成一份24级物联网班 STM32智能小车课程 32课时的教案。"
REAL_TEMPLATE = Path(__file__).resolve().parent / "fixtures" / "教案模板.docx"


def _post_json(base_url: str, path: str, payload: dict[str, Any]) -> tuple[int, dict[str, Any]]:
    request = urllib.request.Request(
        base_url + path,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    return _open_json(request)


def _post_multipart(
    base_url: str,
    path: str,
    fields: dict[str, str],
    files: dict[str, tuple[str, bytes, str]] | None = None,
) -> tuple[int, dict[str, Any]]:
    boundary = "----teacher-skill-stm32-smart-car"
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("ascii"))
        chunks.append(value.encode("utf-8"))
        chunks.append(b"\r\n")
    for name, (filename, content, content_type) in (files or {}).items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(
            f'Content-Disposition: form-data; name="{name}"; filename="{filename}"\r\n'.encode("utf-8")
        )
        chunks.append(f"Content-Type: {content_type}\r\n\r\n".encode("ascii"))
        chunks.append(content)
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode("ascii"))
    request = urllib.request.Request(
        base_url + path,
        data=b"".join(chunks),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    return _open_json(request)


def _open_json(request: urllib.request.Request) -> tuple[int, dict[str, Any]]:
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            return response.status, json.loads(response.read().decode("utf-8") or "{}")
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            data = json.loads(raw or "{}")
        except json.JSONDecodeError:
            data = {"error": raw}
        return exc.code, data


def _patch_web_paths(monkeypatch, tmp_path: Path) -> Path:
    output_dir = tmp_path / "outputs"
    template_dir = tmp_path / "templates"
    monkeypatch.setattr(web_app, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(web_app, "UPLOAD_DIR", output_dir / "uploads")
    monkeypatch.setattr(web_app, "PREVIEW_DIR", output_dir / "previews")
    monkeypatch.setattr(web_app, "HISTORY_DB", output_dir / "history.sqlite3")
    monkeypatch.setattr(web_app, "AGENT_MEMORY_DB", output_dir / "memory.sqlite3")
    monkeypatch.setattr(web_app, "TEMPLATE_DIR", template_dir)
    monkeypatch.setattr(web_app, "SAMPLE_TEMPLATE", template_dir / "sample_lesson_template.docx")
    return output_dir


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _run_export(
    monkeypatch,
    tmp_path: Path,
    *,
    template_mode: str,
    repeat_fill_mode: str | None = None,
) -> tuple[Path, dict[str, Any]]:
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    output_dir = _patch_web_paths(monkeypatch, tmp_path)

    server = ThreadingHTTPServer(("127.0.0.1", 0), web_app.TeacherAgentHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base_url = f"http://127.0.0.1:{server.server_address[1]}"

    try:
        preview_status, preview = _post_json(base_url, "/api/agent-preview", {"agent_request": AGENT_REQUEST})
        assert preview_status == 200, preview

        files = None
        if template_mode == "upload":
            files = {
                "template": (
                    REAL_TEMPLATE.name,
                    REAL_TEMPLATE.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
        run_fields = {
            "agent_request": AGENT_REQUEST,
            "template_mode": template_mode,
            "strict_ai": "0",
        }
        if repeat_fill_mode:
            run_fields["repeat_fill_mode"] = repeat_fill_mode

        run_status, data = _post_multipart(
            base_url,
            "/api/agent-run",
            run_fields,
            files,
        )
        assert run_status == 200, data

        export_payload = {
            "template_id": data["template_id"],
            "fields": data["fields"],
            "request_context": data.get("agent_task") or {},
            "generation_backend": data.get("generation_backend"),
        }
        if repeat_fill_mode:
            export_payload["repeat_fill_mode"] = repeat_fill_mode

        export_status, export = _post_json(
            base_url,
            "/api/export",
            export_payload,
        )
        assert export_status == 200, export
        output_path = output_dir / export["output_name"]
        assert output_path.exists(), export
        return output_path, export
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)


def _find_hour_table(document: Document):
    for table in document.tables:
        if table.rows and [cell.text.strip() for cell in table.rows[0].cells] == ["阶段", "主要内容", "课时", "阶段产出"]:
            return table
    return None


def test_stm32_prompt_parses_title_class_and_scope():
    task = route_task(AGENT_REQUEST)

    assert task.title in {"STM32智能小车课程", "STM32智能小车"}
    assert task.grade == "24级物联网班"
    assert task.subject == "物联网"
    assert task.class_hour == "32课时"
    assert infer_lesson_scope(task.class_hour, task.class_type) == "project_lesson"
    assert is_stm32_smart_car_topic(task.title)


def test_stm32_system_template_e2e_without_deepseek(monkeypatch, tmp_path):
    output_path, _export = _run_export(monkeypatch, tmp_path, template_mode="system")
    text = _docx_text(output_path)
    compact = re.sub(r"\s+", " ", text)

    assert "未命名课题" not in compact
    for keyword in ["STM32", "智能小车", "GPIO", "PWM", "电机驱动", "循迹", "避障", "传感器"]:
        assert keyword in compact
    assert "32课时" in compact or "32 课时" in compact
    assert "Keil" in compact or "STM32CubeMX" in compact
    assert "项目教学法" in compact or "任务驱动" in compact
    assert "帮我生成" not in compact

    table = _find_hour_table(Document(str(output_path)))
    assert table is not None
    table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    for keyword in [
        "阶段",
        "主要内容",
        "课时",
        "阶段产出",
        "STM32开发环境搭建",
        "电机驱动与PWM调速控制",
        "循迹与避障传感器调试",
        "智能小车综合联调",
    ]:
        assert keyword in table_text


def test_stm32_uploaded_template_writes_process_and_method_cells(monkeypatch, tmp_path):
    output_path, export = _run_export(monkeypatch, tmp_path, template_mode="upload", repeat_fill_mode="all")
    text = _docx_text(output_path)
    compact = re.sub(r"\s+", " ", text)

    assert "未命名课题" not in compact
    assert "STM32智能小车" in compact or "STM32 智能小车" in compact

    document = Document(str(output_path))
    passed_locations = 0
    for table in document.tables:
        grid = parse_table_grid(table)
        for row_index, row_grid in enumerate(grid[:-1]):
            process_label = None
            method_label = None
            for gcell in row_grid:
                if gcell is None:
                    continue
                if "主要教学内容" in gcell.text:
                    process_label = gcell
                if "教学方法的运用" in gcell.text:
                    method_label = gcell
            if process_label is None or method_label is None:
                continue
            process_cell = find_cell_by_grid(table, row_index + 1, process_label.grid_col, process_label.physical_col)
            method_cell = find_cell_by_grid(table, row_index + 1, method_label.grid_col, method_label.physical_col)
            if process_cell is None or method_cell is None:
                continue
            process_text = process_cell.text
            method_text = method_cell.text
            process_ok = all(
                keyword in process_text
                for keyword in ["本项目共32课时", "STM32", "智能小车", "PWM", "电机驱动", "循迹", "避障"]
            )
            method_hits = sum(
                keyword in method_text
                for keyword in ["项目教学法", "任务驱动法", "演示教学法", "分组协作", "巡回指导", "作品展示评价"]
            )
            if process_ok and method_hits >= 4:
                passed_locations += 1

    assert passed_locations >= 2
    assert export["fill_report"]["field_write_counts"]["teaching_method"] >= 2

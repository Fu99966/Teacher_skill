from __future__ import annotations

import threading
from http.server import ThreadingHTTPServer

from docx import Document

import teacher_agent.web_app as web_app
from test_uploaded_template_stm32_teaching_method_cell import (
    AGENT_REQUEST,
    _patch_web_paths,
    _post_json,
    _post_multipart,
)


BAD_PROCESS = (
    "本项目共 32 课时，围绕 完成一辆基于 “ STM32 的智能小车设计与调试 展开。 ”\n"
    "一、项目总任务：完成智能小车设计。\n"
    "二、课时分配表：见下表。\n"
    "三、阶段任务：完成GPIO、PWM、电机驱动、循迹和避障调试。\n"
    "四、项目产出：工程文件和展示汇报。\n"
    "五、评价方式：过程评价与作品评价。\n"
    "六、总结提升：复盘DRC、Gerber等工程规范。"
)


def _docx_text(path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_system_template_cleans_malformed_project_punctuation_in_final_docx(monkeypatch, tmp_path):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    output_dir = _patch_web_paths(monkeypatch, tmp_path)

    server = ThreadingHTTPServer(("127.0.0.1", 0), web_app.TeacherAgentHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base_url = f"http://127.0.0.1:{server.server_address[1]}"

    try:
        preview_status, preview = _post_json(
            base_url,
            "/api/agent-preview",
            {"agent_request": AGENT_REQUEST},
        )
        assert preview_status == 200, preview

        run_status, data = _post_multipart(
            base_url,
            "/api/agent-run",
            {
                "agent_request": AGENT_REQUEST,
                "template_mode": "system",
                "strict_ai": "0",
            },
            {},
        )
        assert run_status == 200, data

        fields = dict(data["fields"])
        fields["teaching_process"] = BAD_PROCESS
        export_status, export = _post_json(
            base_url,
            "/api/export",
            {
                "template_id": data["template_id"],
                "fields": fields,
                "request_context": data.get("agent_task") or {},
                "generation_backend": data.get("generation_backend"),
            },
        )
        assert export_status == 200, export

        text = _docx_text(output_dir / export["output_name"])
        assert "围绕 完成" not in text
        assert "“ STM32" not in text
        assert "调试 展开。 ”" not in text
        assert "围绕“完成一辆基于 STM32 的智能小车设计与调试”展开。" in text
        assert export["output_quality_report"]["checks"]["punctuation_clean"] is True
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)

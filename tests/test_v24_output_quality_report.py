from __future__ import annotations

from pathlib import Path

from docx import Document

from teacher_agent.output_quality import inspect_docx_delivery_quality


LONG_METHOD_SENTENCE = "学生在真实智能小车项目实践中完成设计、检查、修改和展示"


def _make_parallel_template(path, method_text: str) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.cell(0, 0).merge(table.cell(0, 2)).text = "主要教学内容"
    table.cell(0, 3).text = "教学方法的运用"
    table.cell(1, 0).merge(table.cell(1, 2)).text = "本项目共32课时，包含阶段任务和项目产出。"
    table.cell(1, 3).text = method_text
    document.add_paragraph("课题")
    document.add_paragraph("STM32智能小车课程")
    document.add_paragraph("教学目的")
    document.add_paragraph("掌握GPIO、PWM、电机驱动、循迹和避障控制。")
    document.add_paragraph("重点难点")
    document.add_paragraph("重点是PWM，难点是多模块联调。")
    document.add_paragraph("作业")
    document.add_paragraph("完成项目记录。")
    document.add_paragraph("课后小记")
    document.add_paragraph("复盘项目实施效果。")
    document.save(path)


def test_output_quality_rejects_verbose_method_even_when_under_120_chars(tmp_path):
    path = tmp_path / "verbose-method.docx"
    method = f"项目教学法、任务驱动法。{LONG_METHOD_SENTENCE}。"
    assert len(method) < 120
    _make_parallel_template(path, method)

    report = inspect_docx_delivery_quality(path, repeat_fill_mode="first_only")

    assert report["checks"]["teaching_method_fit_for_narrow_cell"] is False
    assert report["passed"] is False
    assert "教学方法栏内容过长，可能在窄栏中严重换行，请使用短版教学方法。" in (
        report["errors"] + report["warnings"]
    )


def test_output_quality_accepts_compact_method_and_blank_second_section(tmp_path):
    path = tmp_path / "compact-method.docx"
    document = Document()
    for index in range(2):
        table = document.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        table.cell(0, 0).merge(table.cell(0, 2)).text = "主要教学内容"
        table.cell(0, 3).text = "教学方法的运用"
        if index == 0:
            table.cell(1, 0).merge(table.cell(1, 2)).text = "本项目共32课时，包含阶段任务和项目产出。"
            table.cell(1, 3).text = "项目教学法、任务驱动法、演示教学法、小组协作、巡回指导、作品展示评价。"
        else:
            table.cell(1, 0).merge(table.cell(1, 2)).text = ""
            table.cell(1, 3).text = ""
    for heading, body in (
        ("课题", "STM32智能小车课程"),
        ("教学目的", "掌握GPIO、PWM、电机驱动、循迹和避障控制。"),
        ("重点难点", "重点是PWM，难点是多模块联调。"),
        ("作业", "完成项目记录。"),
        ("课后小记", "复盘项目实施效果。"),
    ):
        document.add_paragraph(heading)
        document.add_paragraph(body)
    document.save(path)

    report = inspect_docx_delivery_quality(path, repeat_fill_mode="first_only")

    assert report["checks"]["teaching_method_fit_for_narrow_cell"] is True
    assert report["checks"]["duplicate_first_only_preserved"] is True
    assert report["passed"] is True, report


def test_web_exposes_teacher_friendly_narrow_method_warning():
    root = Path(__file__).resolve().parents[1]
    script = (root / "web" / "static" / "app.js").read_text(encoding="utf-8")

    assert "教学方法栏内容过长，可能在窄栏中严重换行，请使用短版教学方法。" in script

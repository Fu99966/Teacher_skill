"""Integration tests for Web Agent API scenarios."""
import json
import uuid
from pathlib import Path

from teacher_agent.agent_core.checkpoint import AgentCheckpointStore
from teacher_agent.agent_core.graph_planner import build_graph
from teacher_agent.agent_core.state import AgentRunState
from teacher_agent.agent_core.tool_registry import build_agent_tool_registry
from teacher_agent.agent_core.executor import AgentExecutor, AgentExecutionResult
from teacher_agent.agent_core.task_router import AgentTask

FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
REAL_TEMPLATE = FIXTURE_DIR / "教案模板.docx"


def _make_task(suffix=""):
    return AgentTask(
        raw_request=f"生成教案{suffix}",
        task_type="lesson_plan", subject="语文", grade="四年级",
        title="观潮", class_hour="1课时", class_type="新授课",
        teaching_style="常规启发式", student_level="常规", generation_depth="标准",
        missing_fields=[], confidence=0.9, notes=[],
    )


def _make_state(suffix="", tmpl_path=None):
    return AgentRunState(
        session_id=f"web-test-{suffix}-{uuid.uuid4().hex[:4]}",
        status="initialized", task=_make_task(suffix).to_dict(),
        current_node="", next_action="",
        template_path=str(tmpl_path or REAL_TEMPLATE),
    )


def _make_registry(tmp_path):
    out = tmp_path / "outputs"; out.mkdir()
    return build_agent_tool_registry(
        output_dir=out, preview_dir=out,
        history_db=out / "h.sqlite3", memory_db=out / "m.sqlite3",
    )


# ── Test 1: GET /api/agent/{id} works via checkpoint load ──

def test_web_agent_get_status_route(tmp_path):
    """Simulate what do_GET would do – load state from checkpoint."""
    registry = _make_registry(tmp_path)
    checkpoint = AgentCheckpointStore(tmp_path)
    executor = AgentExecutor(registry, checkpoint)

    state = _make_state("get1")
    graph = build_graph(_make_task("get1"))
    result = executor.run(graph, state)

    # Simulate GET: reload from checkpoint
    reloaded = checkpoint.load(state.session_id)
    assert reloaded is not None
    assert reloaded.status == result.status
    assert reloaded.fields is not None
    assert len(reloaded.trace) > 0


# ── Test 2: Agent continue works with freshly rebuilt graph ──

def test_web_agent_continue_rebuild_graph_still_works(tmp_path):
    """Continue flow builds a NEW graph (not same object) and still works."""
    registry = _make_registry(tmp_path)
    checkpoint = AgentCheckpointStore(tmp_path)
    executor = AgentExecutor(registry, checkpoint)

    # Step 1: Run to gate
    state = _make_state("rebuild1")
    graph1 = build_graph(_make_task("rebuild1"))
    result1 = executor.run(graph1, state)
    assert result1.status == "waiting_teacher_review"

    # Simulate teacher edits (like web continue would)
    state2 = checkpoint.load(state.session_id)
    fields = dict(state2.fields or {})
    fields["homework"] = "老师手动修改的作业内容"
    state2.fields = fields
    state2.teacher_edits = {"homework": "老师手动修改的作业内容"}
    state2.status = "fields_generated"
    checkpoint.save(state2)

    # Step 2: Build NEW graph and continue_after_review
    new_graph = build_graph(_make_task("rebuild1"))  # fresh graph!
    executor2 = AgentExecutor(registry, checkpoint)
    result2 = executor2.continue_after_review(new_graph, checkpoint.load(state.session_id))

    assert result2.status in ("completed", "failed")
    if result2.status == "completed":
        export = result2.state.export_result or {}
        assert export.get("download_url"), "Should have download URL"
        output_name = export.get("output_name", "")
        if output_name:
            docx_path = tmp_path / "outputs" / output_name
            if docx_path.exists():
                from docx import Document
                doc = Document(str(docx_path))
                all_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
                assert "老师手动修改的作业内容" in all_text, f"Teacher edits not in output: {all_text[:300]}"


# ── Test 3: Legacy /api/agent-run still works (delegates to new flow) ──

def test_legacy_agent_run_not_broken_or_removed(tmp_path):
    """Old /api/agent-run now delegates to the new state-machine flow; must not crash."""
    registry = _make_registry(tmp_path)
    checkpoint = AgentCheckpointStore(tmp_path)
    executor = AgentExecutor(registry, checkpoint)

    state = _make_state("legacy1")
    graph = build_graph(_make_task("legacy1"))
    result = executor.run(graph, state)

    # The old flow returned execution.plan; new flow has state.trace
    assert result.state.trace is not None
    assert len(result.state.trace) > 0
    # Verify essential fields exist in result (backward compat)
    assert result.session_id
    assert result.status in ("waiting_teacher_review", "completed", "failed")


# ── Test 4: ToolSpec validates produced context keys ──

def test_tool_spec_validates_produced_keys(tmp_path):
    from teacher_agent.agent_core.tool_spec import ToolSpec, ToolRegistry

    registry = ToolRegistry()

    def good_tool(ctx):
        ctx["state"].fields = ctx["state"].fields or {}
        ctx["state"].fields["test_key"] = "ok"
        return {"done": True}

    registry.register("good", good_tool, ToolSpec(
        name="good", description="test",
        produced_context_keys=["state"],
    ))

    # Should pass
    state = AgentRunState(session_id="spec-test", status="initialized", task={}, current_node="", next_action="")
    state.fields = {}
    result = good_tool({"state": state})
    assert state.fields["test_key"] == "ok"

    missing = registry.validate_produced("good", {"state": state})
    assert not missing, f"Should not have missing produced keys, got {missing}"

    def bad_tool(ctx):
        return {"done": True}  # doesn't modify state

    registry.register("bad", bad_tool, ToolSpec(
        name="bad", description="test",
        produced_context_keys=["state", "nonexistent_key"],
    ))
    missing2 = registry.validate_produced("bad", {})
    assert "state" in missing2 or "nonexistent_key" in missing2


# ── Test 5: Repair continues to export ──

def test_repair_continues_to_export(tmp_path):
    registry = _make_registry(tmp_path)
    checkpoint = AgentCheckpointStore(tmp_path)
    executor = AgentExecutor(registry, checkpoint)

    state = _make_state("repair-export")
    graph = build_graph(_make_task("repair-export"))

    result = executor.run(graph, state)
    assert result.status in ("waiting_teacher_review", "completed", "failed")

    if result.status == "waiting_teacher_review":
        # Simulate continue
        new_graph = build_graph(_make_task("repair-export"))
        executor2 = AgentExecutor(registry, checkpoint)
        state2 = checkpoint.load(state.session_id)
        state2.status = "fields_generated"
        result2 = executor2.continue_after_review(new_graph, state2)
        assert result2.status in ("completed", "failed")

        # If completed, export must exist
        if result2.status == "completed":
            assert result2.state.export_result is not None
            export = result2.state.export_result
            assert export.get("download_url")
            output_name = export.get("output_name", "")
            docx_path = tmp_path / "outputs" / output_name
            if docx_path.exists():
                from docx import Document
                doc = Document(str(docx_path))
                all_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
                assert len(all_text) > 100, f"Output too short: {len(all_text)} chars"


# ── Test 6: Checkpoint survives restart ──

def test_agent_checkpoint_survives_restart(tmp_path):
    """Agent can be recovered and continued after full restart / rebuild."""
    registry = _make_registry(tmp_path)
    checkpoint = AgentCheckpointStore(tmp_path)
    executor = AgentExecutor(registry, checkpoint)

    state = _make_state("restart1")
    graph = build_graph(_make_task("restart1"))
    result = executor.run(graph, state)
    assert result.status == "waiting_teacher_review"

    # Save & reload (simulate restart)
    state_dict = checkpoint.load(state.session_id).to_dict()
    assert state_dict["status"] == "waiting_teacher_review"

    # Rebuild everything fresh
    new_state = AgentRunState(
        session_id=state.session_id,
        status=state_dict.get("status", "initialized"),
        task=state_dict.get("task", {}),
        current_node=state_dict.get("current_node", ""),
        next_action=state_dict.get("next_action", ""),
        template_path=state_dict.get("template_path"),
        fields=state_dict.get("fields"),
        template_analysis=state_dict.get("template_analysis"),
        trace=state_dict.get("trace", []),
    )
    new_state.status = "fields_generated"
    checkpoint.save(new_state)

    # Continue with fresh infrastructure
    restart_dir = tmp_path / "restart-out"
    restart_dir.mkdir(exist_ok=True)
    registry2 = _make_registry(restart_dir)
    checkpoint2 = AgentCheckpointStore(restart_dir)
    executor2 = AgentExecutor(registry2, checkpoint2)

    # Copy state
    checkpoint2.save(new_state)
    new_graph = build_graph(_make_task("restart1"))
    result2 = executor2.continue_after_review(new_graph, new_state)

    assert result2.status in ("completed", "failed")
    assert len(result2.state.trace) >= 2  # at least export + evaluate

"""Tests using the real 教案模板.docx fixture."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from teacher_agent.docx_filler import fill_docx_template
from teacher_agent.template_parser import analyze_template, parse_table_grid

FIXTURE_DIR = Path(__file__).resolve().parent / "fixtures"
REAL_TEMPLATE = FIXTURE_DIR / "教案模板.docx"


def _extract_all_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


# ── Test 1: Real template field detection (MUST include reflection) ────

def test_real_template_fields_are_detected():
    analysis = analyze_template(REAL_TEMPLATE)
    fields = analysis["mapped_fields"]

    print(f"\nDetected fields: {fields}")

    assert "lesson_title" in fields, f"Missing lesson_title: {fields}"
    assert "teaching_goals" in fields, f"Missing teaching_goals: {fields}"
    assert "teaching_key_difficult" in fields, f"Missing teaching_key_difficult: {fields}"
    assert "teaching_process" in fields, f"Missing teaching_process: {fields}"
    assert "teaching_method" in fields, f"Missing teaching_method: {fields}"
    assert "homework" in fields, f"Missing homework: {fields}"
    assert "reflection" in fields, f"Missing reflection (课后小记): {fields}"

    assert analysis["fillable_count"] > 0
    assert not analysis["needs_template_markers"]


# ── Test 2: 课后小记 MUST be detected as reflection ──────────────────

def test_real_template_detects_reflection():
    analysis = analyze_template(REAL_TEMPLATE)
    fields = analysis["mapped_fields"]
    assert "reflection" in fields, f"Missing reflection in fields: {fields}"

    # Verify the label mapping
    reflection_targets = analysis["table_mappings"].get("reflection", [])
    labels = [m.get("label", "").replace(" ", "").replace("\n", "") for m in reflection_targets]
    print(f"\nReflection labels: {labels}")
    assert any("课后小记" in lbl for lbl in labels), f"课后小记 not in reflection labels: {labels}"


# ── Test 3: 主要教学内容 is NOT lesson_title ───────────────────────────

def test_main_teaching_content_is_not_lesson_title():
    analysis = analyze_template(REAL_TEMPLATE)

    lesson_mappings = analysis["table_mappings"].get("lesson_title", [])
    labels = [m.get("label", "").replace(" ", "").replace("\n", "") for m in lesson_mappings]
    assert "主要教学内容" not in labels, f"主要教学内容 wrongly in lesson_title: {labels}"

    process_mappings = analysis["table_mappings"].get("teaching_process", [])
    process_labels = [m.get("label", "").replace(" ", "").replace("\n", "") for m in process_mappings]
    assert any(lbl == "主要教学内容" for lbl in process_labels), f"主要教学内容 not in teaching_process: {process_labels}"


# ── Test 4: 主要教学内容 targets next row ────────────────────────────

def test_main_teaching_content_targets_next_row():
    analysis = analyze_template(REAL_TEMPLATE)
    mappings = analysis["table_mappings"].get("teaching_process", [])

    target = None
    for m in mappings:
        label = m.get("label", "").replace(" ", "").replace("\n", "")
        if label == "主要教学内容":
            target = m
            break

    assert target is not None, "No teaching_process mapping for 主要教学内容"
    assert target["target_type"] in {"next_row_cell", "table_cell", "right_cell"}, f"Unexpected target_type: {target}"
    assert target["row"] > target["label_row"], (
        f"teaching_process target_row({target['row']}) must be > label_row({target['label_row']}): {target}"
    )
    print(f"\nteaching_process target: label_row={target['label_row']}, row={target['row']}, "
          f"target_type={target['target_type']}, grid_col={target.get('grid_col')}")


# ── Test 5: 教学方法的运用 targets next row ─────────────────────────

def test_teaching_method_targets_next_row():
    analysis = analyze_template(REAL_TEMPLATE)
    mappings = analysis["table_mappings"].get("teaching_method", [])

    target = None
    for m in mappings:
        label = m.get("label", "").replace(" ", "").replace("\n", "")
        if label == "教学方法的运用":
            target = m
            break

    assert target is not None, "No teaching_method mapping for 教学方法的运用"
    assert target["row"] > target["label_row"], (
        f"teaching_method target_row({target['row']}) must be > label_row({target['label_row']}): {target}"
    )
    print(f"\nteaching_method target: label_row={target['label_row']}, row={target['row']}, "
          f"target_type={target['target_type']}, grid_col={target.get('grid_col')}")


# ── Test 6: Fixed fields write to real template ──────────────────────

def test_fixed_fields_write_to_real_template(tmp_path):
    output_path = tmp_path / "output.docx"
    fields = {
        "teaching_date": "2026年5月29日",
        "class_name": "24物联网1班",
        "lesson_title": "传感器基础",
        "class_type": "新授课",
        "class_hour": "2课时",
        "teaching_environment": "多媒体教室，具备投影设备和传感器演示套件。",
        "teaching_goals": "理解传感器的基本概念、分类和典型应用。",
        "teaching_key_difficult": "重点：传感器分类与工作原理。难点：传感器信号与物联网系统的关系。",
        "teaching_aids": "PPT、传感器实物、实验演示板。",
        "teaching_process": "这是主要教学内容正文：一、导入展示案例。二、新授讲解概念。三、实践观察分析。四、总结梳理作用。",
        "teaching_method": "案例教学、任务驱动、小组讨论、实物演示。",
        "homework": "完成传感器分类表，并举出三个生活中的传感器应用案例。",
        "reflection": "课后关注学生是否能把传感器与物联网应用场景建立联系。",
    }

    report = fill_docx_template(REAL_TEMPLATE, fields, output_path)
    output_text = _extract_all_text(output_path)

    print(f"\nFill report: filled={report.filled_fields}, errors={report.errors}, table_write={report.table_write_count}")
    print(f"Output text preview: {output_text[:500]}")

    # Check key content
    assert "传感器基础" in output_text, "Missing '传感器基础'"
    assert "理解传感器的基本概念" in output_text, "Missing teaching_goals content"
    assert "案例教学" in output_text, "Missing teaching_method content"
    assert "完成传感器分类表" in output_text, "Missing homework content"
    assert "课后关注学生" in output_text, "Missing reflection content"
    assert "这是主要教学内容正文" in output_text, "Missing teaching_process content"

    # Verify teaching_process is NOT in the label row
    doc = Document(str(output_path))
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                t = cell.text.strip()
                if t == "主要教学内容":
                    # The cell containing the label must NOT also contain the fill content
                    assert "这是主要教学内容正文" not in t, (
                        f"teaching_process content leaked into label cell at row {row_idx}"
                    )

    assert "lesson_title" in report.filled_fields
    assert report.filled_non_empty_count > 0
    assert not report.errors, f"Unexpected errors: {report.errors}"


# ── Test 7: Content is NOT in the same cell as labels ──────────────────

def test_content_not_in_label_cells(tmp_path):
    """Ensure 主要教学内容 and 教学方法的运用 content is in next row, not label row."""
    output_path = tmp_path / "output_pos.docx"
    fields = {
        "lesson_title": "传感器基础",
        "teaching_process": "【教学正文A】导入新课→探究新知→应用巩固",
        "teaching_method": "【教学方法B】案例+任务驱动+小组讨论",
        "teaching_goals": "理解传感器",
        "teaching_key_difficult": "重点难点内容",
        "homework": "课后练习",
        "reflection": "课后反思内容",
    }

    report = fill_docx_template(REAL_TEMPLATE, fields, output_path)
    assert not report.errors, f"Errors: {report.errors}"

    doc = Document(str(output_path))
    for table in doc.tables:
        grid = parse_table_grid(table)
        for row_idx, row_grid in enumerate(grid):
            for gc, gcell in enumerate(row_grid):
                if gcell is None or gcell.grid_col != gc:
                    continue
                text = gcell.text
                if "主要教学内容" in text:
                    assert "【教学正文A】" not in text, (
                        f"teaching_process content leaked into label at row={row_idx}, grid_col={gc}"
                    )
                if "教学方法的运用" in text:
                    assert "【教学方法B】" not in text, (
                        f"teaching_method content leaked into label at row={row_idx}, grid_col={gc}"
                    )

    output_text = _extract_all_text(output_path)
    assert "【教学正文A】" in output_text
    assert "【教学方法B】" in output_text


# ── Test 8: Duplicate table writing ────────────────────────────────────

def test_duplicate_tables_get_filled(tmp_path):
    output_path = tmp_path / "output_dup.docx"
    fields = {
        "lesson_title": "传感器基础",
        "teaching_goals": "理解传感器的基本概念",
        "teaching_key_difficult": "重点：分类。难点：信号关系。",
        "teaching_process": "【正文】导入→新授→实践→总结",
        "teaching_method": "案例教学法",
        "homework": "完成练习题",
        "reflection": "课后反思内容",
        "teaching_environment": "多媒体教室",
        "teaching_aids": "PPT课件",
        "class_hour": "2课时",
    }

    report = fill_docx_template(REAL_TEMPLATE, fields, output_path, repeat_fill_mode="all")
    output_text = _extract_all_text(output_path)

    count = output_text.count("传感器基础")
    process_count = output_text.count("【正文】")
    method_count = output_text.count("案例教学法")
    print(f"\n传感器基础: {count}, 【正文】: {process_count}, 案例教学法: {method_count}")

    assert count >= 2, f"Expected >=2, got {count}"
    assert process_count >= 2, f"Expected >=2, got {process_count}"
    assert method_count >= 2, f"Expected >=2, got {method_count}"

    assert "lesson_title" in report.table_fields_filled
    assert report.table_write_count >= 6, f"Expected table_write_count >= 6, got {report.table_write_count}"


# ── Test 9: Blank output is rejected ────────────────────────────────────

def test_blank_output_is_rejected():
    fields = {
        "lesson_title": "", "teaching_goals": "", "teaching_key_difficult": "",
        "teaching_process": "", "teaching_method": "", "homework": "",
        "reflection": "",
    }
    output_path = Path("__blank_test_output.docx")
    try:
        report = fill_docx_template(REAL_TEMPLATE, fields, output_path)
        assert report.errors, f"Expected errors: {report.to_dict()}"
        assert "空白模板" in report.errors[0] or report.filled_non_empty_count == 0
    finally:
        if output_path.exists():
            output_path.unlink()


# ── Test 10: field_write_counts assertions ─────────────────────────────

def test_field_write_counts(tmp_path):
    output_path = tmp_path / "output_fwc.docx"
    fields = {
        "lesson_title": "传感器基础",
        "teaching_process": "【正文】",
        "teaching_method": "【方法】",
        "teaching_goals": "理解传感器",
        "teaching_key_difficult": "重点难点",
        "homework": "课后练习",
        "reflection": "课后反思",
        "teaching_environment": "多媒体教室",
        "teaching_aids": "PPT",
        "class_hour": "2课时",
        "teaching_date": "2026-05-29",
        "class_name": "24物联网1班",
        "class_type": "新授课",
    }

    report = fill_docx_template(REAL_TEMPLATE, fields, output_path, repeat_fill_mode="all")
    assert not report.errors, f"Errors: {report.errors}"

    fwc = report.field_write_counts or {}
    print(f"\nfield_write_counts: {fwc}")

    assert report.filled_non_empty_count >= 10, f"Too few filled: {report.filled_non_empty_count}"
    assert fwc["lesson_title"] >= 2, f"lesson_title count={fwc.get('lesson_title')}: {fwc}"
    assert fwc["teaching_process"] >= 2, f"teaching_process count={fwc.get('teaching_process')}: {fwc}"
    assert fwc["teaching_method"] >= 2, f"teaching_method count={fwc.get('teaching_method')}: {fwc}"
    assert fwc["reflection"] >= 2, f"reflection count={fwc.get('reflection')}: {fwc}"
    assert fwc["homework"] >= 2, f"homework count={fwc.get('homework')}: {fwc}"
    assert fwc["teaching_goals"] >= 2, f"teaching_goals count={fwc.get('teaching_goals')}: {fwc}"


# ── Test 11: teacher_report.md generated by diagnose-template ──────────

def test_teacher_report_is_generated(tmp_path):
    import subprocess, sys
    output_dir = tmp_path / "diag"
    result = subprocess.run(
        [sys.executable, "-m", "teacher_agent.cli", "diagnose-template",
         "--template", str(REAL_TEMPLATE),
         "--subject", "物联网", "--grade", "24物联网1班", "--title", "传感器基础",
         "--material-file", str(FIXTURE_DIR.parent.parent / "examples" / "sample_material.md"),
         "--output-dir", str(output_dir),
         "--no-strict-ai"],
        capture_output=True, text=True, encoding="utf-8",
        cwd=str(Path(__file__).resolve().parents[1]),
    )
    print(f"diagnose stdout: {result.stdout[:200]}")
    print(f"diagnose stderr: {result.stderr[:200]}")
    assert result.returncode == 0, f"diagnose failed:\n{result.stderr}"

    report_path = output_dir / "teacher_report.md"
    assert report_path.exists(), "teacher_report.md not generated"

    content = report_path.read_text(encoding="utf-8")
    print(f"teacher_report.md preview: {content[:500]}")
    assert "主要教学内容" in content
    assert "教学方法的运用" in content
    assert "课后小记" in content
    assert "填写位置为下一行" in content
    assert "成功写入" in content


# ── Test 12: Content at correct cell positions ─────────────────────────

def test_content_at_correct_grid_positions(tmp_path):
    output_path = tmp_path / "output_pos2.docx"
    fields = {
        "lesson_title": "传感器基础",
        "teaching_process": "【主要教学正文位置测试】",
        "teaching_method": "【教学方法正文位置测试】",
        "teaching_goals": "理解传感器",
        "teaching_key_difficult": "重点难点",
        "homework": "课后练习",
        "reflection": "课后反思",
    }

    report = fill_docx_template(REAL_TEMPLATE, fields, output_path)
    assert not report.errors, f"Errors: {report.errors}"

    doc = Document(str(output_path))
    for t_idx, table in enumerate(doc.tables):
        grid = parse_table_grid(table)
        for ri, row_grid in enumerate(grid):
            for gc, gcell in enumerate(row_grid):
                if gcell is None or gcell.grid_col != gc:
                    continue
                t = gcell.text
                # label cells must NOT contain fill content
                if t.strip() == "主要教学内容":
                    assert "【主要教学正文位置测试】" not in t, (
                        f"Table{t_idx} Row{ri} Col{gc}: teaching_process leaked into label cell"
                    )
                if t.strip() == "教学方法的运用":
                    assert "【教学方法正文位置测试】" not in t, (
                        f"Table{t_idx} Row{ri} Col{gc}: teaching_method leaked into label cell"
                    )

    # Also verify from template_analysis that targets are next-row
    analysis = analyze_template(REAL_TEMPLATE)
    for f, label_text in [("teaching_process", "主要教学内容"), ("teaching_method", "教学方法的运用")]:
        targets = analysis["table_mappings"].get(f, [])
        for t in targets:
            if t.get("label", "").replace(" ", "").replace("\n", "") == label_text:
                assert t["row"] > t["label_row"], (
                    f"{f} target_row({t['row']}) must be > label_row({t['label_row']})"
                )
                assert t.get("target_type") == "next_row_cell", (
                    f"{f} target_type={t.get('target_type')}, expected next_row_cell"
                )

    output_text = _extract_all_text(output_path)
    assert "【主要教学正文位置测试】" in output_text
    assert "【教学方法正文位置测试】" in output_text

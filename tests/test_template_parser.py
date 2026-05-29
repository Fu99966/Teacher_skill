from __future__ import annotations

from docx import Document

from teacher_agent.template_parser import analyze_template


def test_placeholder_template_fields_are_ordered(tmp_path):
    path = tmp_path / "placeholder.docx"
    document = Document()
    document.add_paragraph("课题：{{lesson_title}}")
    document.add_paragraph("目标：{{teaching_goals}}")
    document.add_paragraph("过程：{{teaching_process}}")
    document.save(path)

    analysis = analyze_template(path)

    assert analysis["mapped_fields"] == ["lesson_title", "teaching_goals", "teaching_process"]
    assert analysis["needs_template_markers"] is False


def test_table_label_maps_to_blank_right_cell(tmp_path):
    path = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "教学目标"
    table.cell(0, 1).text = ""
    document.save(path)

    analysis = analyze_template(path)

    assert analysis["mapped_fields"] == ["teaching_goals"]
    assert analysis["table_mappings"]["teaching_goals"]["col"] == 1


def test_mixed_template_keeps_placeholders_and_table_fields(tmp_path):
    path = tmp_path / "mixed.docx"
    document = Document()
    document.add_paragraph("{{lesson_title}}")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "作业设计"
    table.cell(0, 1).text = ""
    document.save(path)

    analysis = analyze_template(path)

    assert analysis["mapped_fields"] == ["lesson_title", "homework"]
    assert analysis["mode"] == "mixed"


def test_custom_fields_enter_dynamic_fields(tmp_path):
    path = tmp_path / "custom.docx"
    document = Document()
    document.add_paragraph("{{warm_up}}")
    document.add_paragraph("{{assessment}}")
    document.save(path)

    analysis = analyze_template(path)

    assert analysis["mapped_fields"] == ["warm_up", "assessment"]


def test_no_fields_returns_clear_error(tmp_path):
    path = tmp_path / "empty.docx"
    document = Document()
    document.add_paragraph("普通学校模板")
    document.save(path)

    analysis = analyze_template(path)

    assert analysis["needs_template_markers"] is True
    assert analysis["errors"]


def test_header_table_label_is_detected(tmp_path):
    path = tmp_path / "header-table.docx"
    document = Document()
    table = document.sections[0].header.add_table(rows=1, cols=2, width=1000000)
    table.cell(0, 0).text = "课题"
    table.cell(0, 1).text = ""
    document.save(path)

    analysis = analyze_template(path)

    assert analysis["mapped_fields"] == ["lesson_title"]
    assert analysis["table_mappings"]["lesson_title"]["location"] == "header_table"

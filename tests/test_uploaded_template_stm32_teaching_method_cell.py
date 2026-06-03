from __future__ import annotations

import json
import re
import threading
import urllib.error
import urllib.request
from http.server import ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document

import teacher_agent.web_app as web_app
from teacher_agent.docx_grid import find_cell_by_grid, parse_table_grid


AGENT_REQUEST = "帮我生成一份24级物联网班 STM32智能小车课程 32课时的教案。"


def _post_json(base_url: str, path: str, payload: dict[str, Any]) -> tuple[int, dict[str, Any]]:
    request = urllib.request.Request(
        base_url + path,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    return _open_json(request)


def _post_multipart(
    base_url: str,
    path: str,
    fields: dict[str, str],
    files: dict[str, tuple[str, bytes, str]],
) -> tuple[int, dict[str, Any]]:
    boundary = "----teacher-skill-uploaded-stm32-method-cell"
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("ascii"))
        chunks.append(value.encode("utf-8"))
        chunks.append(b"\r\n")
    for name, (filename, content, content_type) in files.items():
        chunks.append(f"--{boundary}\r\n".encode("ascii"))
        chunks.append(
            f'Content-Disposition: form-data; name="{name}"; filename="{filename}"\r\n'.encode("utf-8")
        )
        chunks.append(f"Content-Type: {content_type}\r\n\r\n".encode("ascii"))
        chunks.append(content)
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode("ascii"))
    request = urllib.request.Request(
        base_url + path,
        data=b"".join(chunks),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    return _open_json(request)


def _open_json(request: urllib.request.Request) -> tuple[int, dict[str, Any]]:
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            return response.status, json.loads(response.read().decode("utf-8") or "{}")
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            data = json.loads(raw or "{}")
        except json.JSONDecodeError:
            data = {"error": raw}
        return exc.code, data


def _patch_web_paths(monkeypatch, tmp_path: Path) -> Path:
    output_dir = tmp_path / "outputs"
    template_dir = tmp_path / "templates"
    monkeypatch.setattr(web_app, "OUTPUT_DIR", output_dir)
    monkeypatch.setattr(web_app, "UPLOAD_DIR", output_dir / "uploads")
    monkeypatch.setattr(web_app, "PREVIEW_DIR", output_dir / "previews")
    monkeypatch.setattr(web_app, "HISTORY_DB", output_dir / "history.sqlite3")
    monkeypatch.setattr(web_app, "AGENT_MEMORY_DB", output_dir / "memory.sqlite3")
    monkeypatch.setattr(web_app, "TEMPLATE_DIR", template_dir)
    monkeypatch.setattr(web_app, "SAMPLE_TEMPLATE", template_dir / "sample_lesson_template.docx")
    return output_dir


def _add_realistic_school_table(document: Document) -> None:
    table = document.add_table(rows=11, cols=4)
    table.style = "Table Grid"
    rows = [
        ("授课日期", "", "授课班级", ""),
        ("课题（含章节号）", "", "", ""),
        ("授课类型", "", "课时数", ""),
        ("对教学环境的要求", "", "", ""),
        ("教学目的", "", "", ""),
        ("重点难点", "", "", ""),
        ("教具挂图", "", "", ""),
        ("主 要\n教 学 内 容", "", "", "教学方法\n的 运 用"),
        ("", "", "", ""),
        ("作业", "", "", ""),
        ("课后小记", "", "", ""),
    ]
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.rows[row_index].cells[col_index].text = value

    # Common school template shape: label at left, wide merged fill area to the right.
    for row_index in (1, 3, 4, 5, 6, 9, 10):
        table.rows[row_index].cells[1].merge(table.rows[row_index].cells[3])

    # Parallel heading row and target row:
    # left wide area => teaching_process, right small cell => teaching_method.
    table.rows[7].cells[0].merge(table.rows[7].cells[2])
    table.rows[8].cells[0].merge(table.rows[8].cells[2])
    document.add_paragraph("")


def _make_complex_template(tmp_path: Path) -> Path:
    path = tmp_path / "stm32_complex_school_template.docx"
    document = Document()
    _add_realistic_school_table(document)
    _add_realistic_school_table(document)
    document.save(path)
    return path


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _parallel_heading_cells(document: Document):
    for table in document.tables:
        grid = parse_table_grid(table)
        for row_index, row_grid in enumerate(grid[:-1]):
            process_label = None
            method_label = None
            for gcell in row_grid:
                if gcell is None:
                    continue
                if "主要教学内容" in gcell.normalized_text:
                    process_label = gcell
                if "教学方法的运用" in gcell.normalized_text:
                    method_label = gcell
            if process_label is None or method_label is None:
                continue
            process_cell = find_cell_by_grid(table, row_index + 1, process_label.grid_col, process_label.physical_col)
            method_cell = find_cell_by_grid(table, row_index + 1, method_label.grid_col, method_label.physical_col)
            if process_cell is not None and method_cell is not None:
                yield process_cell, method_cell


def test_uploaded_template_stm32_teaching_method_right_cell_and_class_name(monkeypatch, tmp_path):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    output_dir = _patch_web_paths(monkeypatch, tmp_path)
    template = _make_complex_template(tmp_path)

    server = ThreadingHTTPServer(("127.0.0.1", 0), web_app.TeacherAgentHandler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base_url = f"http://127.0.0.1:{server.server_address[1]}"

    try:
        preview_status, preview = _post_json(base_url, "/api/agent-preview", {"agent_request": AGENT_REQUEST})
        assert preview_status == 200, preview
        assert preview["agent_task"]["grade"] == "24级物联网班"

        run_status, data = _post_multipart(
            base_url,
            "/api/agent-run",
            {
                "agent_request": AGENT_REQUEST,
                "template_mode": "upload",
                "strict_ai": "0",
            },
            {
                "template": (
                    template.name,
                    template.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert run_status == 200, data

        export_status, export = _post_json(
            base_url,
            "/api/export",
            {
                "template_id": data["template_id"],
                "fields": data["fields"],
                "request_context": data.get("agent_task") or {},
                "generation_backend": data.get("generation_backend"),
            },
        )
        assert export_status == 200, export
        output_path = output_dir / export["output_name"]
        assert output_path.exists(), export

        text = _docx_text(output_path)
        compact = re.sub(r"\s+", " ", text)
        assert "未命名课题" not in compact
        assert "STM32智能小车课程" in compact or "STM32 智能小车课程" in compact
        assert "24级物联网班" in compact
        assert "24物联网1班" not in compact
        assert "32课时" in compact or "32 课时" in compact
        for keyword in ["GPIO", "PWM", "电机驱动", "循迹", "避障"]:
            assert keyword in compact
        assert "教材依据：帮我生成" not in compact

        passed_locations = 0
        for process_cell, method_cell in _parallel_heading_cells(Document(str(output_path))):
            process_text = process_cell.text
            method_text = method_cell.text
            process_ok = all(
                keyword in process_text
                for keyword in ["本项目共32课时", "STM32", "智能小车", "PWM", "电机驱动", "循迹", "避障"]
            )
            method_hits = sum(
                keyword in method_text
                for keyword in ["项目教学法", "任务驱动法", "演示教学法", "分组协作", "巡回指导", "作品展示评价"]
            )
            if process_ok and method_hits >= 4:
                passed_locations += 1
        assert passed_locations >= 2

        field_reports = export["fill_report"]["field_reports"]
        assert field_reports["teaching_method"]["written_count"] >= 2
        assert field_reports["teaching_method"]["required"] is True
        assert field_reports["teaching_method"]["target_type"] == "next_row_cell"
        assert field_reports["teaching_method"]["status"] == "passed"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=5)

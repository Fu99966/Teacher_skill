from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_NAME = "Microsoft YaHei"

INFO_FIELDS = [
    ("课题", "lesson_title"),
    ("学科", "subject"),
    ("年级", "grade"),
    ("课时", "class_hour"),
]

SECTION_FIELDS = [
    ("教学目标", "teaching_goals"),
    ("教学重点", "key_points"),
    ("教学难点", "difficult_points"),
    ("教学准备", "teaching_preparation"),
    ("教学过程", "teaching_process"),
    ("板书设计", "blackboard_design"),
    ("作业设计", "homework"),
    ("教学反思", "reflection"),
]


def _placeholder(name: str) -> str:
    return "{{" + name + "}}"


def _set_run_font(run) -> None:
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _format_run(run, *, bold: bool = False, size: float = 10.5, color: RGBColor | None = None) -> None:
    _set_run_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _set_style_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10.5)
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _write_cell(cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _format_run(run, bold=bold, size=10.5)


def _add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("标准教案模板")
    _format_run(run, bold=True, size=18, color=RGBColor(17, 24, 39))


def _add_info_table(document: Document) -> None:
    table = document.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    pairs = [
        INFO_FIELDS[0],
        INFO_FIELDS[1],
        INFO_FIELDS[2],
        INFO_FIELDS[3],
    ]
    rows = [
        [pairs[0][0], _placeholder(pairs[0][1]), pairs[1][0], _placeholder(pairs[1][1])],
        [pairs[2][0], _placeholder(pairs[2][1]), pairs[3][0], _placeholder(pairs[3][1])],
    ]

    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        for cell_index, value in enumerate(row_values):
            is_label = cell_index % 2 == 0
            _write_cell(cell=row.cells[cell_index], text=value, bold=is_label, fill="EAF2FF" if is_label else None)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _add_section(document: Document, label: str, field: str) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(3)
    run = heading.add_run(label)
    _format_run(run, bold=True, size=12, color=RGBColor(31, 94, 220))

    body = document.add_paragraph()
    body.paragraph_format.line_spacing = 1.35
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.left_indent = Inches(0.08)
    run = body.add_run(_placeholder(field))
    _format_run(run, size=10.5, color=RGBColor(17, 24, 39))


def create_sample_template(output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_style_font(document)
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    document.core_properties.title = "标准教案模板"
    _add_title(document)
    _add_info_table(document)
    for label, field in SECTION_FIELDS:
        _add_section(document, label, field)

    document.save(str(output_path))
    return output_path

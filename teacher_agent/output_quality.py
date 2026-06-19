from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from .docx_grid import find_cell_by_grid, parse_table_grid


PROMPT_MARKERS = ("帮我生成一份", "请生成一份", "写一份教案", "生成一份教案")
BAD_PUNCTUATION_PATTERNS = (
    "围绕 完成",
    "“ STM32",
    "调试 展开。 ”",
    "“ PCB",
    "设计 展开。 ”",
)
NARROW_METHOD_FORBIDDEN_PHRASES = (
    "学生在真实智能小车项目实践中完成设计、检查、修改和展示",
)
NARROW_METHOD_WARNING = "教学方法栏内容过长，可能在窄栏中严重换行，请使用短版教学方法。"


def _normalize_label(text: str) -> str:
    return re.sub(r"[\s:：；;、，,。.·\-—–（）()\[\]【】<>《》]+", "", str(text or ""))


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _expected_content_present(expected: str, normalized_document: str) -> bool:
    normalized_expected = _normalize_label(expected)
    if not normalized_expected:
        return True
    if len(normalized_expected) <= 10:
        return normalized_expected in normalized_document
    if normalized_expected[:24] in normalized_document:
        return True
    grams = {
        normalized_expected[index : index + 6]
        for index in range(0, len(normalized_expected) - 5, 6)
    }
    matches = sum(1 for gram in grams if gram in normalized_document)
    return matches >= min(2, len(grams))


def _paragraph_body_after_heading(document: Document, headings: tuple[str, ...]) -> str:
    normalized_headings = {_normalize_label(heading) for heading in headings}
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if _normalize_label(paragraph.text) not in normalized_headings:
            continue
        for candidate in paragraphs[index + 1 :]:
            if candidate.text.strip():
                return candidate.text.strip()
    return ""


def _parallel_process_method_cells(document: Document) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        grid = parse_table_grid(table)
        for row_index, row_grid in enumerate(grid[:-1]):
            unique_cells: list[Any] = []
            seen: set[int] = set()
            for grid_cell in row_grid:
                if grid_cell is None:
                    continue
                cell_id = id(grid_cell.cell._tc)
                if cell_id in seen:
                    continue
                seen.add(cell_id)
                unique_cells.append(grid_cell)

            process_label = next(
                (cell for cell in unique_cells if "主要教学内容" in cell.normalized_text),
                None,
            )
            method_label = next(
                (cell for cell in unique_cells if "教学方法的运用" in cell.normalized_text),
                None,
            )
            if process_label is None or method_label is None:
                continue

            process_cell = find_cell_by_grid(
                table,
                row_index + 1,
                process_label.grid_col,
                process_label.physical_col,
            )
            method_cell = find_cell_by_grid(
                table,
                row_index + 1,
                method_label.grid_col,
                method_label.physical_col,
            )
            if process_cell is None or method_cell is None:
                continue
            sections.append(
                {
                    "table": table_index,
                    "label_row": row_index,
                    "process_text": process_cell.text.strip(),
                    "method_text": method_cell.text.strip(),
                    "process_grid_span": process_label.grid_span,
                    "method_grid_span": method_label.grid_span,
                }
            )
    return sections


def _method_fits_narrow_cell(text: str) -> bool:
    method_text = str(text or "").strip()
    return bool(
        method_text
        and len(method_text) <= 120
        and not any(phrase in method_text for phrase in NARROW_METHOD_FORBIDDEN_PHRASES)
    )


def inspect_docx_delivery_quality(
    docx_path: str | Path,
    *,
    repeat_fill_mode: str | None = None,
    expected_fields: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Inspect the final DOCX artifact instead of trusting intermediate field data."""
    path = Path(docx_path)
    checks: dict[str, bool] = {}
    errors: list[str] = []
    warnings: list[str] = []

    try:
        document = Document(str(path))
    except Exception as exc:
        return {
            "passed": False,
            "score": 0,
            "errors": [f"无法读取导出的 Word：{exc}"],
            "warnings": [],
            "checks": {"docx_readable": False},
        }

    text = _document_text(document)
    normalized = _normalize_label(text)
    checks["docx_readable"] = True
    checks["no_prompt_leak"] = not any(marker in text for marker in PROMPT_MARKERS)
    checks["no_unnamed_title"] = "未命名课题" not in text
    checks["punctuation_clean"] = not any(pattern in text for pattern in BAD_PUNCTUATION_PATTERNS)
    checks["material_basis_clean"] = "教材依据：帮我生成" not in text

    required_labels = {
        "has_lesson_title": ("课题",),
        "has_teaching_goals": ("教学目的", "教学目标"),
        "has_key_difficult": ("重点难点", "教学重点", "教学难点"),
        "has_teaching_process": ("主要教学内容", "教学过程"),
        "has_teaching_method": ("教学方法的运用",),
        "has_homework": ("作业", "作业设计"),
        "has_reflection": ("课后小记", "教学反思"),
    }
    for check_name, labels in required_labels.items():
        checks[check_name] = any(_normalize_label(label) in normalized for label in labels)

    expected_content_checks: list[str] = []
    for field_name in (
        "lesson_title",
        "teaching_goals",
        "teaching_key_difficult",
        "teaching_process",
        "teaching_method",
        "homework",
        "reflection",
    ):
        expected = _normalize_label(str((expected_fields or {}).get(field_name) or ""))
        if not expected:
            continue
        check_name = f"field_content_{field_name}"
        expected_content_checks.append(check_name)
        checks[check_name] = _expected_content_present(expected, normalized)

    parallel_sections = _parallel_process_method_cells(document)
    if parallel_sections:
        populated_methods = [section["method_text"] for section in parallel_sections if section["method_text"]]
        populated_processes = [section["process_text"] for section in parallel_sections if section["process_text"]]
        checks["teaching_process_written"] = bool(populated_processes)
        checks["teaching_method_written"] = bool(populated_methods)
        method_sections = (
            parallel_sections[:1]
            if repeat_fill_mode == "first_only"
            else parallel_sections
        )
        checks["teaching_method_fit_for_narrow_cell"] = bool(method_sections) and all(
            _method_fits_narrow_cell(section["method_text"])
            for section in method_sections
        )
        if repeat_fill_mode == "first_only" and len(parallel_sections) > 1:
            checks["duplicate_first_only_preserved"] = bool(
                parallel_sections[0]["process_text"]
                and parallel_sections[0]["method_text"]
                and all(
                    not section["process_text"] and not section["method_text"]
                    for section in parallel_sections[1:]
                )
            )
        else:
            checks["duplicate_first_only_preserved"] = True
    else:
        checks["teaching_process_written"] = bool(
            _paragraph_body_after_heading(document, ("教学过程", "主要教学内容"))
        )
        checks["teaching_method_written"] = bool(
            _paragraph_body_after_heading(document, ("教学方法的运用",))
        )
        checks["teaching_method_fit_for_narrow_cell"] = True
        checks["duplicate_first_only_preserved"] = True

    critical_checks = (
        "no_prompt_leak",
        "no_unnamed_title",
        "punctuation_clean",
        "material_basis_clean",
        "has_lesson_title",
        "has_teaching_goals",
        "has_key_difficult",
        "has_teaching_process",
        "has_teaching_method",
        "has_homework",
        "has_reflection",
        "teaching_process_written",
        "teaching_method_written",
        "teaching_method_fit_for_narrow_cell",
        "duplicate_first_only_preserved",
    ) + tuple(expected_content_checks)
    messages = {
        "no_prompt_leak": "检测到生成指令泄漏到 Word 正文。",
        "no_unnamed_title": "课题仍为“未命名课题”。",
        "punctuation_clean": "检测到中文引号或空格排版异常。",
        "material_basis_clean": "教材依据仍包含生成指令。",
        "has_lesson_title": "未识别到课题栏目。",
        "has_teaching_goals": "未识别到教学目的或教学目标栏目。",
        "has_key_difficult": "未识别到重点难点栏目。",
        "has_teaching_process": "未识别到主要教学内容或教学过程栏目。",
        "has_teaching_method": "未识别到教学方法的运用栏目。",
        "has_homework": "未识别到作业栏目。",
        "has_reflection": "未识别到课后小记或教学反思栏目。",
        "teaching_process_written": "主要教学内容未写入非空内容。",
        "teaching_method_written": "教学方法的运用未写入非空内容。",
        "teaching_method_fit_for_narrow_cell": NARROW_METHOD_WARNING,
        "duplicate_first_only_preserved": "first_only 模式下第二套重复教案区域仍被填充。",
    }
    for check_name in critical_checks:
        if not checks.get(check_name, False):
            errors.append(messages.get(check_name, f"最终 Word 中未找到字段内容：{check_name.removeprefix('field_content_')}"))

    score = max(0, 100 - len(errors) * 12 - len(warnings) * 5)
    return {
        "passed": not errors,
        "score": score,
        "errors": errors,
        "warnings": warnings,
        "checks": checks,
    }

from __future__ import annotations

import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from .docx_grid import find_cell_by_grid
from .template_parser import PLACEHOLDER_PATTERN, analyze_template, iter_paragraphs


@dataclass
class FillReport:
    output_path: str
    filled_fields: list[str] = field(default_factory=list)
    missing_fields: list[str] = field(default_factory=list)
    empty_fields: list[str] = field(default_factory=list)
    skipped_empty_fields: list[str] = field(default_factory=list)
    unfilled_template_fields: list[str] = field(default_factory=list)
    filled_non_empty_count: int = 0
    remaining_placeholders: list[str] = field(default_factory=list)
    placeholder_fields_filled: list[str] = field(default_factory=list)
    table_fields_filled: list[str] = field(default_factory=list)
    table_write_count: int = 0
    field_write_counts: dict[str, int] = field(default_factory=dict)
    section_write_counts: dict[str, int] = field(default_factory=dict)
    repeated_sections_detected: int = 0
    repeat_fill_mode: str = "first_only"
    filled_sections: int = 0
    field_reports: dict[str, dict[str, Any]] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    @property
    def path(self) -> Path:
        return Path(self.output_path)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    def __fspath__(self) -> str:
        return self.output_path

    def __str__(self) -> str:
        return self.output_path


def _docx_text(value: Any) -> str:
    if value is None: return ""
    if isinstance(value, str): return value.replace("\r\n", "\n").replace("\r", "\n")
    if isinstance(value, (list, tuple)): return "\n".join(_docx_text(item) for item in value)
    if isinstance(value, dict): return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def _is_effectively_empty(value: Any) -> bool:
    return not _docx_text(value).strip()


def _add_ordered(target: list[str], field_name: str) -> None:
    if field_name and field_name not in target:
        target.append(field_name)


def _increment_field_write(report: FillReport, field_name: str) -> None:
    report.field_write_counts[field_name] = report.field_write_counts.get(field_name, 0) + 1


LESSON_SECTION_MARKERS = {
    "lesson_title",
    "teaching_goals",
    "teaching_process",
    "teaching_method",
    "homework",
    "reflection",
}

COMPACT_TEACHING_METHOD = "项目教学法、任务驱动法、演示教学法、小组协作、巡回指导、作品展示评价。"
NARROW_METHOD_VERBOSE_PHRASES = (
    "学生在真实智能小车项目实践中完成设计、检查、修改和展示",
)


def _section_key_from_target(target: dict[str, Any]) -> str | None:
    if target.get("location") != "body_table":
        return None
    table_index = target.get("table")
    if table_index is None:
        return None
    return f"body_table:{table_index}"


def _detect_repeated_lesson_sections(analysis: dict[str, Any]) -> list[str]:
    section_fields: dict[str, set[str]] = {}
    for field_name, targets in (analysis.get("table_mappings") or {}).items():
        for target in targets:
            key = _section_key_from_target(target)
            if not key:
                continue
            section_fields.setdefault(key, set()).add(field_name)

    sections: list[str] = []
    for key, fields in section_fields.items():
        has_core = {"lesson_title", "teaching_process", "teaching_method"}.issubset(fields)
        has_support = bool(fields & {"teaching_goals", "teaching_key_difficult", "homework", "reflection"})
        if has_core and has_support:
            sections.append(key)

    return sorted(sections, key=lambda value: int(value.rsplit(":", 1)[1]))


def _normalize_repeat_fill_mode(mode: str | None) -> str:
    return "all" if str(mode or "").strip() == "all" else "first_only"


def _filter_mappings_for_repeat_mode(
    mappings: dict[str, list[dict[str, Any]]],
    repeated_sections: list[str],
    repeat_fill_mode: str,
) -> dict[str, list[dict[str, Any]]]:
    if repeat_fill_mode == "all" or len(repeated_sections) <= 1:
        return mappings

    allowed_section = repeated_sections[0]
    repeated = set(repeated_sections)
    filtered: dict[str, list[dict[str, Any]]] = {}
    for field_name, targets in mappings.items():
        kept: list[dict[str, Any]] = []
        for target in targets:
            key = _section_key_from_target(target)
            if key in repeated and key != allowed_section:
                continue
            kept.append(target)
        filtered[field_name] = kept
    return filtered


def _replace_placeholders_in_run(run, data: dict[str, Any], report: FillReport) -> bool:
    original = run.text
    changed = False
    def replace(match):
        nonlocal changed
        key = match.group(1).strip()
        if key not in data: _add_ordered(report.missing_fields, key); return match.group(0)
        if _is_effectively_empty(data[key]): _add_ordered(report.empty_fields, key); _add_ordered(report.skipped_empty_fields, key); return match.group(0)
        changed = True; _add_ordered(report.filled_fields, key); _add_ordered(report.placeholder_fields_filled, key); _increment_field_write(report, key); return _docx_text(data[key])
    new_text = PLACEHOLDER_PATTERN.sub(replace, original)
    if changed: run.text = new_text
    return changed


def _run_ranges(paragraph):
    ranges, cursor = [], 0
    for i, run in enumerate(paragraph.runs): ranges.append((i, cursor, cursor + len(run.text))); cursor += len(run.text)
    return ranges


def _replace_cross_run_placeholder(paragraph, match, data, report) -> bool:
    key = match.group(1).strip()
    if key not in data: _add_ordered(report.missing_fields, key); return False
    if _is_effectively_empty(data[key]): _add_ordered(report.empty_fields, key); _add_ordered(report.skipped_empty_fields, key); return False
    ranges = _run_ranges(paragraph)
    first = last = None
    for i, s, e in ranges:
        if e <= match.start() or s >= match.end(): continue
        if first is None: first = i
        last = i
    if first is None or last is None: return False
    fr = paragraph.runs[first]; fs = ranges[first][1]
    lr = paragraph.runs[last]; ls = ranges[last][1]
    fr.text = fr.text[:max(0, match.start() - fs)] + _docx_text(data[key]) + lr.text[max(0, match.end() - ls):]
    for i in range(first + 1, last + 1): paragraph.runs[i].text = ""
    _add_ordered(report.filled_fields, key); _add_ordered(report.placeholder_fields_filled, key); _increment_field_write(report, key)
    return True


def _replace_paragraph(paragraph, data, report):
    if "{{" not in paragraph.text: return
    for run in paragraph.runs: _replace_placeholders_in_run(run, data, report)
    full_text = "".join(r.text for r in paragraph.runs)
    for match in reversed(list(PLACEHOLDER_PATTERN.finditer(full_text))):
        _replace_cross_run_placeholder(paragraph, match, data, report)


def _copy_paragraph_style(source, target):
    target.style = source.style; target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing


def _write_paragraph_preserving_style(paragraph, text: str):
    if not paragraph.runs: paragraph.add_run(text); return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]: run.text = ""


def _write_cell_preserving_layout(cell, value: Any):
    text = _docx_text(value)
    if not cell.paragraphs: cell.add_paragraph(text); return
    lines = text.split("\n")
    tmpl = cell.paragraphs[0]; _write_paragraph_preserving_style(tmpl, lines[0] if lines else "")
    for p in cell.paragraphs[1:]: _write_paragraph_preserving_style(p, "")
    for line in lines[1:]:
        p = cell.add_paragraph(); _copy_paragraph_style(tmpl, p); p.add_run(line)


def _append_cell_preserving_label(cell, value: Any):
    text = _docx_text(value)
    if not text: return
    if not cell.paragraphs: cell.add_paragraph(text); return
    lines = text.split("\n")
    target = None
    for p in cell.paragraphs[1:]:
        if not p.text.strip(): target = p; break
    if target is None: target = cell.add_paragraph(); _copy_paragraph_style(cell.paragraphs[0], target)
    _write_paragraph_preserving_style(target, lines[0] if lines else "")
    for line in lines[1:]:
        p = cell.add_paragraph(); _copy_paragraph_style(target, p); p.add_run(line)


def _resolve_table(document, target):
    loc = target.get("location", "body_table"); ti = int(target.get("table", -1))
    if ti < 0: return None
    if loc == "header_table":
        si = int(target.get("section") or 0)
        if si >= len(document.sections): return None
        tables = document.sections[si].header.tables
    elif loc == "footer_table":
        si = int(target.get("section") or 0)
        if si >= len(document.sections): return None
        tables = document.sections[si].footer.tables
    else:
        tables = document.tables
    return tables[ti] if ti < len(tables) else None


def _fill_one_table_target(document, field_name, value, target, report) -> bool:
    mapping_type = target.get("type")
    if mapping_type not in ("table_cell", "table_cell_append"): return False

    row_index = int(target.get("row", -1))
    if row_index < 0: return False

    table = _resolve_table(document, target)
    if table is None:
        report.warnings.append(f"字段 {field_name} 的目标表格未找到（table={target.get('table')}）。")
        return False

    grid_col = target.get("grid_col")
    physical_col = target.get("physical_col")

    cell = None
    if grid_col is not None:
        cell = find_cell_by_grid(table, row_index, grid_col, physical_col)
    if cell is None and physical_col is not None:
        if row_index < len(table.rows) and physical_col < len(table.rows[row_index].cells):
            cell = table.rows[row_index].cells[physical_col]

    if cell is None:
        report.warnings.append(f"字段 {field_name} 的目标单元格未找到，可能是模板结构复杂或 grid 定位失败（row={row_index}, grid_col={grid_col}）。")
        return False

    if mapping_type == "table_cell_append":
        _append_cell_preserving_label(cell, value)
    else:
        _write_cell_preserving_layout(cell, value)

    report.table_write_count += 1
    _increment_field_write(report, field_name)
    section_key = _section_key_from_target(target)
    if section_key:
        report.section_write_counts[section_key] = report.section_write_counts.get(section_key, 0) + 1
    return True


def _is_narrow_teaching_method_target(
    target: dict[str, Any],
    mappings: dict[str, list[dict[str, Any]]],
) -> bool:
    if target.get("target_type") != "next_row_cell":
        return False
    method_span = int(target.get("grid_span") or 1)
    for process_target in mappings.get("teaching_process", []):
        same_region = (
            process_target.get("location") == target.get("location")
            and process_target.get("section") == target.get("section")
            and process_target.get("table") == target.get("table")
            and process_target.get("label_row") == target.get("label_row")
            and process_target.get("row") == target.get("row")
        )
        if not same_region:
            continue
        process_span = int(process_target.get("grid_span") or 1)
        if method_span * 2 <= process_span:
            return True
        # Some school templates store equal grid spans/tcW values even though the
        # right-hand teaching-method column is visually constrained. A parallel
        # next-row target to the right of teaching_process should stay concise.
        return int(target.get("grid_col") or 0) > int(process_target.get("grid_col") or 0)
    return False


def _should_compact_teaching_method(value: Any) -> bool:
    text = _docx_text(value).strip()
    return len(text) > 80 or any(phrase in text for phrase in NARROW_METHOD_VERBOSE_PHRASES)


def _build_field_reports(analysis: dict[str, Any], report: FillReport) -> None:
    mappings = analysis.get("table_mappings", {})
    required_fields = set(analysis.get("required_fields", []))
    for field_name in analysis.get("mapped_fields", []):
        targets = mappings.get(field_name, [])
        first_target = targets[0] if targets else {}
        required = field_name in required_fields or any(target.get("required") for target in targets)
        written_count = report.field_write_counts.get(field_name, 0)
        target_type = first_target.get("target_type") or ("placeholder" if field_name in report.placeholder_fields_filled else "")
        status = "passed" if written_count > 0 else ("failed" if required else "warning")
        report.field_reports[field_name] = {
            "label": first_target.get("label") or field_name,
            "required": required,
            "written_count": written_count,
            "target_type": target_type,
            "status": status,
        }
        if required and written_count == 0:
            report.warnings.append(f"必填字段 {field_name} 未写入任何目标单元格。")


def _fill_table_mappings(document, data, mappings, report):
    for field_name, targets in mappings.items():
        if field_name not in data: _add_ordered(report.missing_fields, field_name); continue
        if _is_effectively_empty(data[field_name]): _add_ordered(report.empty_fields, field_name); _add_ordered(report.skipped_empty_fields, field_name); continue
        wrote_any = False
        for target in targets:
            value = data[field_name]
            if (
                field_name == "teaching_method"
                and _is_narrow_teaching_method_target(target, mappings)
                and _should_compact_teaching_method(value)
            ):
                value = COMPACT_TEACHING_METHOD
            if _fill_one_table_target(document, field_name, value, target, report):
                wrote_any = True
        if wrote_any: _add_ordered(report.filled_fields, field_name); _add_ordered(report.table_fields_filled, field_name)


def _remaining_placeholders(document):
    fields = []
    for p in iter_paragraphs(document):
        for m in PLACEHOLDER_PATTERN.finditer(p.text): _add_ordered(fields, m.group(1).strip())
    return fields


def fill_docx_template(template_path, data, output_path, repeat_fill_mode: str | None = None):
    template_path = Path(template_path); output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    analysis = analyze_template(template_path)
    report = FillReport(output_path=str(output_path))
    report.repeat_fill_mode = _normalize_repeat_fill_mode(repeat_fill_mode)
    repeated_sections = _detect_repeated_lesson_sections(analysis)
    report.repeated_sections_detected = len(repeated_sections)

    document = Document(str(template_path))
    for p in iter_paragraphs(document): _replace_paragraph(p, data, report)

    mappings = _filter_mappings_for_repeat_mode(
        analysis.get("table_mappings", {}),
        repeated_sections,
        report.repeat_fill_mode,
    )
    _fill_table_mappings(document, data, mappings, report)
    if repeated_sections:
        report.filled_sections = sum(1 for section in repeated_sections if report.section_write_counts.get(section, 0) > 0)

    for fn in analysis.get("mapped_fields", []):
        if fn not in data: _add_ordered(report.missing_fields, fn)
        elif _is_effectively_empty(data[fn]): _add_ordered(report.empty_fields, fn); _add_ordered(report.skipped_empty_fields, fn)
        if fn not in report.filled_fields: _add_ordered(report.unfilled_template_fields, fn)

    report.remaining_placeholders = _remaining_placeholders(document)
    report.filled_non_empty_count = len(report.filled_fields)
    tf_count = len(analysis.get("mapped_fields", []))
    if tf_count > 0 and report.filled_non_empty_count == 0:
        report.errors.append("生成失败：检测到输出可能为空白模板，未写入任何非空字段。")
    elif tf_count > 0 and report.filled_non_empty_count < max(1, tf_count * 0.5):
        report.warnings.append("警告：本次只填入少量模板字段，请检查字段映射和生成结果。")

    _build_field_reports(analysis, report)
    document.save(str(output_path))
    return report

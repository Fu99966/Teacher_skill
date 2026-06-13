from __future__ import annotations

import cgi
import json
import mimetypes
import re
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse

from . import deepseek_client
from .agent_observer import build_teacher_diagnostic_report
from .agent_core.memory import AgentMemoryStore
from .agent_core.executor import AgentExecutor
from .agent_core.evaluator import evaluate_lesson_output
from .agent_core.planner import build_plan
from .agent_core.task_router import route_task
from .agent_core.tool_registry import build_lesson_tool_registry
from .artifact_naming import unique_artifact_name, unique_upload_name
from .deepseek_client import DeepSeekError
from .history_store import HistoryStore
from .lesson_generator import (
    DEFAULT_CLASS_TYPE,
    DEFAULT_GENERATION_DEPTH,
    DEFAULT_STUDENT_LEVEL,
    DEFAULT_TEACHING_STYLE,
    LessonGenerationError,
    check_generation_health,
    extract_class_name_from_request,
    extract_lesson_title_from_request,
    normalize_lesson_field_aliases,
    refine_lesson_field,
    sanitize_lesson_title,
    sanitize_material_hint,
)
from .sample_template import create_sample_template
from .template_parser import analyze_template
from .template_profile import TemplateProfileStore
from .material_ingestion import extract_material_from_upload, merge_material_text
from .output_quality import inspect_docx_delivery_quality
from .workflow import LessonRequest, TeacherWorkflow, build_workflow_schema


def _create_placeholder_template(path: Path) -> None:
    """Generate a minimal docx with standard {{placeholders}}."""
    if path.exists():
        return
    from docx import Document
    from docx.table import _Cell as _DCell
    doc = Document()
    doc.add_heading("教案模板", 0)
    doc.add_paragraph("请替换以下占位符：")

    fields = [
        ("课题", "{{lesson_title}}"),
        ("教学目的", "{{teaching_goals}}"),
        ("重点难点", "{{teaching_key_difficult}}"),
        ("主要教学内容", "{{teaching_process}}"),
        ("教学方法的运用", "{{teaching_method}}"),
        ("作业", "{{homework}}"),
        ("课后小记", "{{reflection}}"),
    ]
    table = doc.add_table(rows=len(fields), cols=2, style="Table Grid")
    for i, (label, placeholder) in enumerate(fields):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WEB_ROOT = PROJECT_ROOT / "web"
OUTPUT_DIR = PROJECT_ROOT / "outputs"
UPLOAD_DIR = OUTPUT_DIR / "uploads"
PREVIEW_DIR = OUTPUT_DIR / "previews"
HISTORY_DB = OUTPUT_DIR / "teacher_skill_history.sqlite3"
AGENT_MEMORY_DB = OUTPUT_DIR / "teacher_skill_agent_memory.sqlite3"
TEMPLATE_DIR = PROJECT_ROOT / "templates"
SAMPLE_MATERIAL = PROJECT_ROOT / "examples" / "sample_material.md"
SAMPLE_TEMPLATE = TEMPLATE_DIR / "sample_lesson_template.docx"
MAX_JSON_REQUEST_BYTES = 5 * 1024 * 1024
MAX_MULTIPART_REQUEST_BYTES = 50 * 1024 * 1024


class RequestTooLargeError(ValueError):
    pass


def _json_bytes(payload: dict, status: int = HTTPStatus.OK) -> tuple[int, bytes, str]:
    return status, json.dumps(payload, ensure_ascii=False).encode("utf-8"), "application/json; charset=utf-8"


def _form_value(form: cgi.FieldStorage, name: str, default: str = "") -> str:
    item = form[name] if name in form else None
    if item is None or item.filename:
        return default
    value = item.value
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="replace")
    return str(value)


def _material_from_form(form: cgi.FieldStorage, typed_material: str) -> tuple[str, dict | None]:
    item = form["material_file"] if "material_file" in form else None
    if item is None or not item.filename:
        return typed_material, None
    extraction = extract_material_from_upload(Path(item.filename).name, item.file)
    return merge_material_text(typed_material, extraction), extraction.to_dict()


def _parse_probe(query: str) -> bool:
    """Parse ?probe=1/true/yes/on from query string using parse_qs."""
    from urllib.parse import parse_qs as _parse_qs
    params = _parse_qs(query or "")
    val = (params.get("probe") or [""])[0].strip().lower()
    return val in ("1", "true", "yes", "on")


def _form_bool(form: cgi.FieldStorage, name: str, default: bool = False) -> bool:
    value = _form_value(form, name, "1" if default else "")
    return value.strip().lower() in {"1", "true", "yes", "on", "strict"}


def _template_id_for(path: Path) -> str:
    if path == SAMPLE_TEMPLATE:
        return "__sample__"
    return path.name


def _template_from_id(template_id: str) -> Path:
    if template_id == "__sample__":
        if not SAMPLE_TEMPLATE.exists():
            create_sample_template(SAMPLE_TEMPLATE)
        return SAMPLE_TEMPLATE

    path = (UPLOAD_DIR / Path(template_id).name).resolve()
    if UPLOAD_DIR.resolve() not in path.parents or not path.exists():
        raise ValueError("模板已失效，请重新上传模板")
    return path


def _template_mode_label(template_path: Path) -> str:
    return "system" if template_path == SAMPLE_TEMPLATE else "upload"


def _normalize_repeat_fill_mode(value: str | None, template_mode: str) -> str:
    if template_mode != "upload":
        return "all"
    return "all" if str(value or "").strip() == "all" else "first_only"


def _beginner_summary(evaluation_report: dict | None, is_generic_material: bool, template_mode: str) -> str:
    checks_ok = bool(evaluation_report and evaluation_report.get("passed"))
    quality_text = "已完成教研审阅和自动检查" if checks_ok else "已完成教案生成，建议下载前再快速预览"
    template_text = "系统标准模板" if template_mode == "system" else "学校上传模板"
    material_text = "本次未填写教材内容，已生成通用版。" if is_generic_material else "已结合教材内容生成。"
    return f"{quality_text}，并写入{template_text}。{material_text}"


def _infer_single_prompt_defaults(text: str, defaults: dict[str, str]) -> dict[str, str]:
    """Fill the three teacher-facing essentials from a natural-language prompt."""
    merged = dict(defaults)
    normalized = re.sub(r"\s+", " ", (text or "").strip())
    explicit_class_name = extract_class_name_from_request(normalized)
    if explicit_class_name:
        merged["grade"] = explicit_class_name

    if not merged.get("title"):
        title = extract_lesson_title_from_request(normalized)
        if title:
            merged["title"] = title

    if not merged.get("grade"):
        class_name = extract_class_name_from_request(normalized)
        if class_name:
            merged["grade"] = class_name
        else:
            grade_match = re.search(r"([0-9０-９]{1,4}级[\u4e00-\u9fffA-Za-z0-9０-９]*)", normalized)
            if grade_match:
                merged["grade"] = grade_match.group(1).strip()

    if not merged.get("subject"):
        subject_keywords = (
            "物联网", "信息技术", "人工智能", "计算机", "语文", "数学", "英语",
            "物理", "化学", "生物", "科学", "历史", "地理", "政治",
        )
        for keyword in subject_keywords:
            if keyword in normalized:
                merged["subject"] = keyword
                break

    hour_match = re.search(r"([0-9０-９一二三四五六七八九十]+)\s*课时", normalized)
    if hour_match and ((not merged.get("class_hour")) or merged.get("class_hour") == "1课时"):
        merged["class_hour"] = hour_match.group(0).strip()

    if not merged.get("class_type"):
        if "实训" in normalized:
            merged["class_type"] = "实训课"
        elif "公开课" in normalized:
            merged["class_type"] = "公开课"

    if not merged.get("teaching_style") and ("项目式" in normalized or "项目驱动" in normalized or "PBL" in normalized.upper()):
        merged["teaching_style"] = "项目驱动(PBL)"

    return merged


def _sanitize_agent_task_title(task: Any, agent_request: str, fallback_title: str = "") -> Any:
    clean_title = sanitize_lesson_title(getattr(task, "title", ""), agent_request, fallback_title)
    if hasattr(task, "title"):
        task.title = clean_title
    if hasattr(task, "missing_fields") and clean_title:
        task.missing_fields = [field for field in task.missing_fields if field != "title"]
    return task


def _normalize_teacher_web_fields(fields: dict[str, Any], task: Any, agent_request: str = "") -> dict[str, Any]:
    """Keep the single-prompt UI populated even when the system template uses legacy keys."""
    normalized = dict(fields or {})
    normalized["lesson_title"] = sanitize_lesson_title(
        str(normalized.get("lesson_title") or ""),
        agent_request,
        str(getattr(task, "title", "") or normalized.get("title", "")),
    )
    normalized.setdefault("subject", getattr(task, "subject", ""))
    task_grade = str(getattr(task, "grade", "") or "").strip()
    if task_grade:
        normalized["grade"] = task_grade
        normalized["class_name"] = task_grade
    else:
        normalized.setdefault("grade", "")
        normalized.setdefault("class_name", "")
    normalized.setdefault("class_hour", getattr(task, "class_hour", ""))
    normalized.setdefault("class_type", getattr(task, "class_type", ""))

    key_points = str(normalized.get("key_points") or "").strip()
    difficult_points = str(normalized.get("difficult_points") or "").strip()
    if "teaching_key_difficult" not in normalized and (key_points or difficult_points):
        parts = []
        if key_points:
            parts.append(f"重点：{key_points}")
        if difficult_points:
            parts.append(f"难点：{difficult_points}")
        normalized["teaching_key_difficult"] = "\n".join(parts)

    method = str(normalized.get("teaching_method") or "").strip()
    if not method:
        style = str(getattr(task, "teaching_style", "") or "")
        class_type = str(getattr(task, "class_type", "") or "")
        class_hour_text = str(getattr(task, "class_hour", "") or normalized.get("class_hour", ""))
        hour_match = re.search(r"(\d+)", class_hour_text.translate(str.maketrans("０１２３４５６７８９", "0123456789")))
        hour_count = int(hour_match.group(1)) if hour_match else 1
        process_text = str(normalized.get("teaching_process") or "")
        if hour_count >= 9 or "课时分配" in process_text or "DRC" in process_text or "Gerber" in process_text:
            if "PCB" in str(getattr(task, "title", "") or normalized.get("lesson_title", "")).upper():
                method = "项目教学法、任务驱动法、演示教学法、分组协作、巡回指导、作品展示评价。围绕原理图绘制、PCB布局布线、DRC检查、Gerber输出和作品展示推进。"
            else:
                method = "项目教学法、任务驱动法、演示教学法、分组协作、巡回指导和作品展示评价。"
        elif "项目" in style or "PBL" in style.upper() or "实训" in class_type:
            method = "项目驱动、小组协作、任务实践、案例分析。"
        elif "公开课" in class_type:
            method = "情境导入、启发提问、互动探究、评价反馈。"
        else:
            method = "讲授启发、任务练习、小组讨论、课堂反馈。"
        normalized["teaching_method"] = method

    return normalize_lesson_field_aliases(normalized, agent_request)


class TeacherAgentHandler(BaseHTTPRequestHandler):
    server_version = "TeacherAgentWeb/0.1"

    def log_message(self, format: str, *args) -> None:
        print("[%s] %s" % (self.log_date_time_string(), format % args))

    def _send(
        self,
        status: int,
        body: bytes,
        content_type: str,
        filename: str | None = None,
        include_body: bool = True,
    ) -> None:
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        origin = str(self.headers.get("Origin") or "").strip()
        if origin and self._is_allowed_origin(origin):
            self.send_header("Access-Control-Allow-Origin", origin)
            self.send_header("Vary", "Origin")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, DELETE, HEAD, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.send_header("Content-Length", str(len(body)))
        if filename:
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")
        self.end_headers()
        if include_body:
            self.wfile.write(body)

    def _is_allowed_origin(self, origin: str) -> bool:
        parsed_origin = urlparse(origin)
        if parsed_origin.scheme not in {"http", "https"} or not parsed_origin.netloc:
            return False
        host_header = str(self.headers.get("Host") or "").strip()
        if not host_header:
            return False
        if parsed_origin.netloc.lower() == host_header.lower():
            return True

        parsed_host = urlparse(f"//{host_header}")
        loopback_hosts = {"localhost", "127.0.0.1", "::1"}
        return (
            (parsed_origin.hostname or "").lower() in loopback_hosts
            and (parsed_host.hostname or "").lower() in loopback_hosts
            and parsed_origin.port == parsed_host.port
        )

    def _validate_mutating_origin(self) -> None:
        origin = str(self.headers.get("Origin") or "").strip()
        if origin and not self._is_allowed_origin(origin):
            raise PermissionError("只允许当前教案助手页面访问本地写入接口。")

    def _validate_post_size(self) -> None:
        raw_length = str(self.headers.get("Content-Length") or "0").strip()
        try:
            content_length = int(raw_length)
        except ValueError as exc:
            raise ValueError("Content-Length 无效") from exc
        if content_length < 0:
            raise ValueError("Content-Length 无效")
        content_type = str(self.headers.get("Content-Type") or "").lower()
        limit = MAX_MULTIPART_REQUEST_BYTES if "multipart/form-data" in content_type else MAX_JSON_REQUEST_BYTES
        if content_length > limit:
            size_mb = max(1, round(limit / (1024 * 1024)))
            raise RequestTooLargeError(f"请求内容过大，请将文件或内容控制在 {size_mb} MB 以内。")

    def _read_json_payload(self) -> dict[str, Any]:
        content_length = int(self.headers.get("Content-Length", "0") or "0")
        if content_length <= 0:
            return {}
        value = json.loads(self.rfile.read(content_length).decode("utf-8"))
        if not isinstance(value, dict):
            raise ValueError("请求内容必须是 JSON 对象")
        return value

    def _send_file(self, path: Path, as_attachment: bool = False) -> None:
        if not path.exists() or not path.is_file():
            self._send(*_json_bytes({"error": "文件不存在"}, HTTPStatus.NOT_FOUND))
            return

        content_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
        filename = path.name if as_attachment else None
        self._send(HTTPStatus.OK, path.read_bytes(), content_type, filename)

    def _send_file_headers(self, path: Path, as_attachment: bool = False) -> None:
        if not path.exists() or not path.is_file():
            self._send(*_json_bytes({"error": "文件不存在"}, HTTPStatus.NOT_FOUND), include_body=False)
            return

        content_type = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
        filename = path.name if as_attachment else None
        self._send(HTTPStatus.OK, path.read_bytes(), content_type, filename, include_body=False)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = unquote(parsed.path)

        if path in {"/", "/index.html"}:
            self._send_file(WEB_ROOT / "index.html")
            return

        if path.startswith("/static/"):
            static_path = (WEB_ROOT / path.removeprefix("/")).resolve()
            if WEB_ROOT.resolve() not in static_path.parents:
                self._send(*_json_bytes({"error": "非法路径"}, HTTPStatus.BAD_REQUEST))
                return
            self._send_file(static_path)
            return

        if path == "/api/sample-material":
            material = SAMPLE_MATERIAL.read_text(encoding="utf-8") if SAMPLE_MATERIAL.exists() else ""
            self._send(*_json_bytes({"material": material}))
            return

        if path == "/api/workflow-schema":
            self._send(*_json_bytes(build_workflow_schema()))
            return

        if path == "/api/history":
            history = HistoryStore(HISTORY_DB).list_documents()
            self._send(*_json_bytes({"items": history}))
            return

        if path == "/api/config/model":
            self._send(*_json_bytes(deepseek_client.get_model_config_public()))
            return

        if path == "/api/llm-health":
            probe = _parse_probe(parsed.query)
            self._send(*_json_bytes({"llm": check_generation_health(probe=probe).to_dict()}))
            return

        if path == "/download/sample-template":
            if not SAMPLE_TEMPLATE.exists():
                create_sample_template(SAMPLE_TEMPLATE)
            self._send_file(SAMPLE_TEMPLATE, as_attachment=True)
            return

        if path == "/download/placeholder-template":
            pt_path = OUTPUT_DIR / "placeholder_template.docx"
            _create_placeholder_template(pt_path)
            self._send_file(pt_path, as_attachment=True)
            return

        if path.startswith("/download/"):
            name = Path(path.removeprefix("/download/")).name
            self._send_file(OUTPUT_DIR / name, as_attachment=True)
            return

        if path.startswith("/preview/"):
            name = Path(path.removeprefix("/preview/")).name
            self._send_file(PREVIEW_DIR / name, as_attachment=False)
            return

        if path == "/health":
            self._send(*_json_bytes({"ok": True}))
            return

        if path == "/api/model/health":
            from .deepseek_client import check_deepseek_health
            probe = _parse_probe(parsed.query)
            ds = check_deepseek_health(probe=probe)
            self._send(*_json_bytes({
                "configured": ds.configured,
                "provider": "deepseek",
                "model": ds.model,
                "status": ds.status,
                "message": ds.message,
                "base_url": ds.base_url,
                "error_type": ds.error_type,
            }))
            return

        # ── Agent GET routes ──
        agent_get_match = re.match(r"^/api/agent/([^/]+)$", path)
        if agent_get_match:
            try:
                self._handle_agent_get(agent_get_match.group(1))
            except KeyError:
                self._send(*_json_bytes({"error": "Session not found"}, HTTPStatus.NOT_FOUND))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        self._send(*_json_bytes({"error": "页面不存在"}, HTTPStatus.NOT_FOUND))

    def do_HEAD(self) -> None:
        parsed = urlparse(self.path)
        path = unquote(parsed.path)

        if path in {"/", "/index.html"}:
            self._send_file_headers(WEB_ROOT / "index.html")
            return

        if path == "/download/sample-template":
            if not SAMPLE_TEMPLATE.exists():
                create_sample_template(SAMPLE_TEMPLATE)
            self._send_file_headers(SAMPLE_TEMPLATE, as_attachment=True)
            return

        if path.startswith("/download/"):
            name = Path(path.removeprefix("/download/")).name
            self._send_file_headers(OUTPUT_DIR / name, as_attachment=True)
            return

        if path.startswith("/preview/"):
            name = Path(path.removeprefix("/preview/")).name
            self._send_file_headers(PREVIEW_DIR / name, as_attachment=False)
            return

        self._send(*_json_bytes({"ok": True}), include_body=False)

    def do_OPTIONS(self) -> None:
        try:
            self._validate_mutating_origin()
        except PermissionError as exc:
            self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.FORBIDDEN))
            return
        self._send(HTTPStatus.NO_CONTENT, b"", "text/plain; charset=utf-8", include_body=False)

    def do_DELETE(self) -> None:
        try:
            self._validate_mutating_origin()
        except PermissionError as exc:
            self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.FORBIDDEN))
            return
        parsed = urlparse(self.path)
        path = unquote(parsed.path)

        if path == "/api/config/model":
            try:
                self._send(*_json_bytes(deepseek_client.clear_model_config()))
            except OSError as exc:
                self._send(*_json_bytes({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        match = re.fullmatch(r"/api/history/([A-Za-z0-9_-]+)", path)
        if match:
            history = HistoryStore(HISTORY_DB)
            record = history.get_document(match.group(1))
            if not record:
                self._send(*_json_bytes({"ok": False, "error": "导出记录不存在"}, HTTPStatus.NOT_FOUND))
                return

            output_name = str(record.get("output_name") or "")
            if not output_name or Path(output_name).name != output_name:
                self._send(*_json_bytes({"ok": False, "error": "导出文件路径不安全"}, HTTPStatus.BAD_REQUEST))
                return

            output_root = OUTPUT_DIR.resolve()
            output_path = (OUTPUT_DIR / output_name).resolve()
            if output_path.parent != output_root:
                self._send(*_json_bytes({"ok": False, "error": "导出文件路径不安全"}, HTTPStatus.BAD_REQUEST))
                return

            deleted_file = False
            if output_path.exists():
                if not output_path.is_file():
                    self._send(*_json_bytes({"ok": False, "error": "导出目标不是文件"}, HTTPStatus.BAD_REQUEST))
                    return
                output_path.unlink()
                deleted_file = True
            deleted_history = history.delete_document(match.group(1))
            self._send(*_json_bytes({
                "ok": True,
                "deleted_file": deleted_file,
                "deleted_history": deleted_history,
            }))
            return

        self._send(*_json_bytes({"ok": False, "error": "接口不存在"}, HTTPStatus.NOT_FOUND))

    def do_POST(self) -> None:
        try:
            self._validate_mutating_origin()
            self._validate_post_size()
        except PermissionError as exc:
            self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.FORBIDDEN))
            return
        except RequestTooLargeError as exc:
            self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.REQUEST_ENTITY_TOO_LARGE))
            return
        except ValueError as exc:
            self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.BAD_REQUEST))
            return

        parsed = urlparse(self.path)
        if parsed.path == "/api/config/model":
            try:
                self._send(*_json_bytes(deepseek_client.save_model_config(self._read_json_payload())))
            except (DeepSeekError, ValueError) as exc:
                message = exc.user_message if isinstance(exc, DeepSeekError) else str(exc)
                self._send(*_json_bytes({"ok": False, "error": message}, HTTPStatus.BAD_REQUEST))
            except OSError as exc:
                self._send(*_json_bytes({"ok": False, "error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/config/model/test":
            try:
                result = deepseek_client.test_model_config(self._read_json_payload())
                result.pop("api_key", None)
                self._send(*_json_bytes(result))
            except (DeepSeekError, ValueError) as exc:
                message = exc.user_message if isinstance(exc, DeepSeekError) else str(exc)
                self._send(*_json_bytes({"ok": False, "error": message}, HTTPStatus.BAD_REQUEST))
            return

        if parsed.path == "/api/draft":
            try:
                self._handle_draft()
            except DeepSeekError as exc:
                self._send(*_json_bytes({"error": exc.user_message, "llm_error": exc.to_dict()}, HTTPStatus.BAD_GATEWAY))
            except LessonGenerationError as exc:
                self._send(
                    *_json_bytes(
                        {"error": str(exc), "llm_error": {"message": str(exc), "type": "generation_error"}},
                        HTTPStatus.BAD_GATEWAY,
                    )
                )
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/export":
            try:
                self._handle_export()
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/refine-field":
            try:
                self._handle_refine_field()
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/remember-edit":
            try:
                self._handle_remember_edit()
            except ValueError as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.BAD_REQUEST))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/agent-preview":
            try:
                self._handle_agent_preview()
            except ValueError as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.BAD_REQUEST))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/agent-run":
            try:
                self._handle_agent_run()
            except ValueError as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.BAD_REQUEST))
            except DeepSeekError as exc:
                self._send(*_json_bytes({"error": exc.user_message, "llm_error": exc.to_dict()}, HTTPStatus.BAD_GATEWAY))
            except LessonGenerationError as exc:
                self._send(
                    *_json_bytes(
                        {"error": str(exc), "llm_error": {"message": str(exc), "type": "generation_error"}},
                        HTTPStatus.BAD_GATEWAY,
                    )
                )
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/step1-diagnose":
            try:
                self._handle_step1_diagnose()
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/step2-generate":
            try:
                self._handle_step2_generate()
            except DeepSeekError as exc:
                self._send(*_json_bytes({"error": exc.user_message, "llm_error": exc.to_dict()}, HTTPStatus.BAD_GATEWAY))
            except LessonGenerationError as exc:
                self._send(*_json_bytes({"error": str(exc), "llm_error": {"message": str(exc), "type": "generation_error"}}, HTTPStatus.BAD_GATEWAY))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path == "/api/step4-fill":
            try:
                self._handle_step4_fill()
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        # ── Agent API routes ──
        agent_start_match = re.match(r"^/api/agent/start$", parsed.path)
        if agent_start_match:
            try:
                self._handle_agent_start()
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        agent_continue_match = re.match(r"^/api/agent/([^/]+)/continue$", parsed.path)
        if agent_continue_match:
            try:
                self._handle_agent_continue(agent_continue_match.group(1))
            except KeyError:
                self._send(*_json_bytes({"error": "Session not found"}, HTTPStatus.NOT_FOUND))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        agent_cancel_match = re.match(r"^/api/agent/([^/]+)/cancel$", parsed.path)
        if agent_cancel_match:
            try:
                self._handle_agent_cancel(agent_cancel_match.group(1))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        agent_repair_match = re.match(r"^/api/agent/([^/]+)/repair$", parsed.path)
        if agent_repair_match:
            try:
                self._handle_agent_repair(agent_repair_match.group(1))
            except Exception as exc:
                self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))
            return

        if parsed.path != "/api/generate":
            self._send(*_json_bytes({"error": "接口不存在"}, HTTPStatus.NOT_FOUND))
            return

        try:
            self._handle_generate()
        except DeepSeekError as exc:
            self._send(*_json_bytes({"error": exc.user_message, "llm_error": exc.to_dict()}, HTTPStatus.BAD_GATEWAY))
        except LessonGenerationError as exc:
            self._send(
                *_json_bytes(
                    {"error": str(exc), "llm_error": {"message": str(exc), "type": "generation_error"}},
                    HTTPStatus.BAD_GATEWAY,
                )
            )
        except Exception as exc:
            self._send(*_json_bytes({"error": str(exc)}, HTTPStatus.INTERNAL_SERVER_ERROR))

    def _handle_step1_diagnose(self) -> None:
        """Step 1: upload template, return template analysis."""
        form = self._read_multipart_form()
        template_path = self._save_template(form)
        analysis = analyze_template(template_path)
        self._send(*_json_bytes({"analysis": analysis, "template_path": str(template_path)}))

    def _handle_step2_generate(self) -> None:
        """Step 2: generate lesson fields from course info + template."""
        from .lesson_generator import draft_lesson_document_fields_with_source
        form = self._read_multipart_form()
        template_path = self._save_template(form)
        subject = _form_value(form, "subject", "")
        grade = _form_value(form, "grade", "")
        title = _form_value(form, "title", "")
        class_hour = _form_value(form, "class_hour", "1课时")
        material, material_extraction = _material_from_form(form, _form_value(form, "material", ""))
        class_type = _form_value(form, "class_type", DEFAULT_CLASS_TYPE)
        teaching_style = _form_value(form, "teaching_style", DEFAULT_TEACHING_STYLE)
        strict_ai = _form_bool(form, "strict_ai", False)
        template_analysis = analyze_template(template_path)
        template_fields = template_analysis.get("mapped_fields") or None
        fields, backend = draft_lesson_document_fields_with_source(
            subject, grade, title, material, class_hour,
            class_type, teaching_style, DEFAULT_STUDENT_LEVEL, DEFAULT_GENERATION_DEPTH,
            template_fields, strict_ai, template_analysis.get("field_context"),
        )
        self._send(*_json_bytes({
            "fields": fields,
            "generation_backend": backend,
            "template_analysis": template_analysis,
            "template_fields": template_fields,
            "template_id": _template_id_for(template_path),
        }))

    def _handle_step4_fill(self) -> None:
        """Step 4: fill template with provided fields JSON and return download URL."""
        import json as _json
        from .docx_filler import fill_docx_template
        form = self._read_multipart_form()
        template_path = self._save_template(form)
        repeat_fill_mode = _normalize_repeat_fill_mode(
            _form_value(form, "repeat_fill_mode", ""),
            "upload",
        )
        fields_json = _form_value(form, "fields_json", "{}")
        fields = _json.loads(fields_json)
        if not isinstance(fields, dict) or not fields:
            raise ValueError("fields_json 格式错误或为空")
        template_analysis = analyze_template(template_path)
        required_fields = template_analysis.get("required_fields", [])
        tm_value = str(fields.get("teaching_method") or "").strip()

        # Validate teaching_method is non-empty
        if "teaching_method" in required_fields and not tm_value:
            self._send(*_json_bytes({
                "error": "teaching_method_empty",
                "message": "「教学方法的运用」为空，不能导出。请返回第3步补充教学方法内容后再导出 Word。",
                "fill_report": None,
                "download_url": "",
                "output_name": "",
                "template_analysis": template_analysis,
            }, HTTPStatus.UNPROCESSABLE_ENTITY))
            return

        grade = str(fields.get("grade") or fields.get("class_name") or "年级")
        subject = str(fields.get("subject") or "学科")
        title = str(fields.get("lesson_title") or "教案")
        output_name = unique_artifact_name(f"{grade}-{subject}-{title}-教案", ".docx")
        output_path = OUTPUT_DIR / output_name
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        fill_report = fill_docx_template(
            template_path,
            fields,
            output_path,
            repeat_fill_mode=repeat_fill_mode,
        )
        output_quality_report = inspect_docx_delivery_quality(
            output_path,
            repeat_fill_mode=repeat_fill_mode,
        )
        fwc = fill_report.field_write_counts
        tm_count = fwc.get("teaching_method", 0)
        self._send(*_json_bytes({
            "fill_report": fill_report.to_dict(),
            "download_url": f"/download/{quote(output_name)}",
            "output_name": output_name,
            "template_analysis": template_analysis,
            "output_quality_report": output_quality_report,
            "teaching_method_writes": tm_count,
            "teaching_method_status": "已填写" if tm_count >= 2 else ("部分填写" if tm_count > 0 else "未填写"),
        }))

    def _handle_agent_start(self) -> None:
        """POST /api/agent/start – Start agent, pause at teacher_review_gate."""
        import uuid
        from .agent_core.checkpoint import AgentCheckpointStore
        from .agent_core.graph_planner import build_graph
        from .agent_core.state import AgentRunState
        from .agent_core.tool_registry import build_agent_tool_registry
        from .agent_core.executor import AgentExecutor
        from .agent_core.task_router import route_task

        form = self._read_multipart_form()
        template_path = self._save_template(form)
        template_id = _template_id_for(template_path)
        analysis = analyze_template(template_path)

        agent_request = _form_value(form, "agent_request", "")
        subject = _form_value(form, "subject", "物联网")
        grade = _form_value(form, "grade", "24物联网1班")
        title = _form_value(form, "title", "未命名课题")
        material = _form_value(form, "material", "")
        class_hour = _form_value(form, "class_hour", "1课时")
        class_type = _form_value(form, "class_type", DEFAULT_CLASS_TYPE)
        teaching_style = _form_value(form, "teaching_style", DEFAULT_TEACHING_STYLE)

        task = route_task(agent_request or f"生成{subject}{grade}{title}教案", {
            "subject": subject, "grade": grade, "title": title,
            "class_hour": class_hour, "class_type": class_type, "teaching_style": teaching_style,
        })
        session_id = uuid.uuid4().hex[:12]
        task_dict = task.to_dict()
        task_dict["material"] = material

        state = AgentRunState(
            session_id=session_id, status="initialized", task=task_dict,
            current_node="", next_action="", template_path=str(template_path),
            template_id=template_id, template_analysis=analysis,
        )

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        registry = build_agent_tool_registry(
            output_dir=OUTPUT_DIR, preview_dir=PREVIEW_DIR,
            history_db=HISTORY_DB, memory_db=AGENT_MEMORY_DB,
        )
        checkpoint = AgentCheckpointStore(OUTPUT_DIR)
        executor = AgentExecutor(registry, checkpoint)

        graph = build_graph(task)
        result = executor.run(graph, state)

        self._send(*_json_bytes(result.to_dict()))

    def _handle_agent_get(self, session_id: str) -> None:
        """GET /api/agent/{session_id} – Get current agent state."""
        from .agent_core.checkpoint import AgentCheckpointStore
        checkpoint = AgentCheckpointStore(OUTPUT_DIR)
        state = checkpoint.load(session_id)
        if state is None:
            raise KeyError(session_id)
        self._send(*_json_bytes(state.to_dict()))

    def _handle_agent_continue(self, session_id: str) -> None:
        """POST /api/agent/{session_id}/continue – Continue after teacher review."""
        import json as _json
        from .agent_core.checkpoint import AgentCheckpointStore
        from .agent_core.graph_planner import build_graph
        from .agent_core.state import AgentRunState
        from .agent_core.tool_registry import build_agent_tool_registry
        from .agent_core.executor import AgentExecutor
        from .agent_core.task_router import AgentTask

        checkpoint = AgentCheckpointStore(OUTPUT_DIR)
        state = checkpoint.load(session_id)
        if state is None:
            raise KeyError(session_id)

        content_type = self.headers.get("Content-Type", "")
        teacher_edits: dict[str, Any] = {}
        if "application/json" in content_type:
            content_length = int(self.headers.get("Content-Length", "0") or "0")
            body = self.rfile.read(content_length).decode("utf-8")
            teacher_edits = _json.loads(body).get("teacher_edits", {})
        elif "multipart/form-data" in content_type:
            form = self._read_multipart_form()
            edits_json = _form_value(form, "teacher_edits", "{}")
            teacher_edits = _json.loads(edits_json)

        if teacher_edits:
            fields = state.fields or {}
            fields.update(teacher_edits)
            state.fields = fields
            state.teacher_edits = teacher_edits
            state.status = "fields_generated"

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        registry = build_agent_tool_registry(
            output_dir=OUTPUT_DIR, preview_dir=PREVIEW_DIR,
            history_db=HISTORY_DB, memory_db=AGENT_MEMORY_DB,
        )
        executor = AgentExecutor(registry, checkpoint)
        # Rebuild graph for task
        task = AgentTask(
            raw_request=str(state.task.get("raw_request", "")),
            task_type="lesson_plan", subject=str(state.task.get("subject", "")),
            grade=str(state.task.get("grade", "")), title=str(state.task.get("title", "")),
            class_hour=str(state.task.get("class_hour", "1课时")),
            class_type=str(state.task.get("class_type", "新授课")),
            teaching_style=str(state.task.get("teaching_style", "常规启发式")),
            student_level=str(state.task.get("student_level", "常规混合水平")),
            generation_depth=str(state.task.get("generation_depth", "标准")),
            missing_fields=[], confidence=0.9, notes=[],
        )
        group = build_graph(task)
        result = executor.continue_after_review(group, state)
        self._send(*_json_bytes(result.to_dict()))

    def _handle_agent_cancel(self, session_id: str) -> None:
        """POST /api/agent/{session_id}/cancel"""
        from .agent_core.checkpoint import AgentCheckpointStore
        checkpoint = AgentCheckpointStore(OUTPUT_DIR)
        checkpoint.delete(session_id)
        self._send(*_json_bytes({"cancelled": True}))

    def _handle_agent_repair(self, session_id: str) -> None:
        """POST /api/agent/{session_id}/repair – Attempt repair."""
        from .agent_core.checkpoint import AgentCheckpointStore
        from .agent_core.repair import repair_state

        checkpoint = AgentCheckpointStore(OUTPUT_DIR)
        state = checkpoint.load(session_id)
        if state is None:
            raise KeyError(session_id)

        state = repair_state(state)
        checkpoint.save(state)
        self._send(*_json_bytes(state.to_dict()))

    def _read_multipart_form(self) -> cgi.FieldStorage:
        content_type = self.headers.get("Content-Type", "")
        if "multipart/form-data" not in content_type:
            raise ValueError("请使用表单提交")

        return cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": content_type},
            encoding="utf-8",
            errors="replace",
        )

    def _read_lesson_form(self) -> tuple[cgi.FieldStorage, str, str, str, str, str, str, str, str, str, bool, str, Path, dict]:
        form = self._read_multipart_form()

        subject = _form_value(form, "subject", "语文")
        grade = _form_value(form, "grade", "四年级")
        title = _form_value(form, "title", "未命名课题")
        class_hour = _form_value(form, "class_hour", "1课时")
        material = _form_value(form, "material", "")
        class_type = _form_value(form, "class_type", DEFAULT_CLASS_TYPE)
        teaching_style = _form_value(form, "teaching_style", DEFAULT_TEACHING_STYLE)
        student_level = _form_value(form, "student_level", DEFAULT_STUDENT_LEVEL)
        generation_depth = _form_value(form, "generation_depth", DEFAULT_GENERATION_DEPTH)
        strict_ai = _form_bool(form, "strict_ai", False)
        creative_mode = _form_value(form, "creative_mode", "常规稳妥")

        template_path = self._save_template(form)
        template_analysis = analyze_template(template_path)
        return (
            form,
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            strict_ai,
            creative_mode,
            template_path,
            template_analysis,
        )

    def _read_agent_form(self) -> tuple[
        cgi.FieldStorage,
        str,
        str,
        str,
        str,
        str,
        str,
        str,
        str,
        str,
        Path,
        dict,
        bool,
        str,
        str,
        bool,
    ]:
        form = self._read_multipart_form()

        subject = _form_value(form, "subject", "")
        grade = _form_value(form, "grade", "")
        title = _form_value(form, "title", "")
        class_hour = _form_value(form, "class_hour", "1课时")
        material, material_extraction = _material_from_form(
            form,
            _form_value(form, "material", ""),
        )
        class_type = _form_value(form, "class_type", DEFAULT_CLASS_TYPE)
        teaching_style = _form_value(form, "teaching_style", DEFAULT_TEACHING_STYLE)
        student_level = _form_value(form, "student_level", DEFAULT_STUDENT_LEVEL)
        generation_depth = _form_value(form, "generation_depth", DEFAULT_GENERATION_DEPTH)
        strict_ai = _form_bool(form, "strict_ai", False)
        creative_mode = _form_value(form, "creative_mode", "更像公开课")
        template_mode = _form_value(form, "template_mode", "system").strip() or "system"
        if template_mode not in {"system", "upload"}:
            raise ValueError("模板模式无效，请选择系统模板或上传模板")
        if template_mode == "upload":
            template_item = form["template"] if "template" in form else None
            if template_item is None or not template_item.filename:
                raise ValueError("请上传学校 Word 模板，或选择使用系统标准模板")

        template_path = self._save_template(form, allow_sample=template_mode == "system")
        template_analysis = analyze_template(template_path)
        if material_extraction:
            template_analysis = dict(template_analysis)
            template_analysis["material_extraction"] = material_extraction
        return (
            form,
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            template_path,
            template_analysis,
            strict_ai,
            creative_mode,
            _template_mode_label(template_path),
            not material.strip(),
        )

    def _handle_draft(self) -> None:
        (
            _,
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            strict_ai,
            creative_mode,
            template_path,
            template_analysis,
        ) = self._read_lesson_form()
        request = LessonRequest(
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            strict_ai,
            creative_mode,
        )
        workflow = TeacherWorkflow(history_db=HISTORY_DB)
        result = workflow.draft(request, template_path, _template_id_for(template_path), template_analysis)

        # Keep this local variable referenced so template analysis failures are
        # surfaced before generation starts.
        result["template_analysis"] = result.get("template_analysis") or template_analysis
        result["llm_status"] = check_generation_health(probe=False).to_dict()
        self._send(*_json_bytes(result))

    def _handle_refine_field(self) -> None:
        content_length = int(self.headers.get("Content-Length", "0") or "0")
        payload = json.loads(self.rfile.read(content_length).decode("utf-8"))

        field = str(payload.get("field") or "")
        value = str(payload.get("value") or "")
        action = str(payload.get("action") or "more_vivid")
        instruction = str(payload.get("instruction") or "")
        if not field:
            raise ValueError("缺少字段名")

        refined = refine_lesson_field(field, value, action, instruction)
        self._send(*_json_bytes({"field": field, "value": refined}))

    def _handle_remember_edit(self) -> None:
        content_length = int(self.headers.get("Content-Length", "0") or "0")
        payload = json.loads(self.rfile.read(content_length).decode("utf-8"))
        fields = payload.get("fields")
        if not isinstance(fields, dict) or not fields:
            raise ValueError("缺少可记忆的教案字段")
        task = payload.get("request_context") if isinstance(payload.get("request_context"), dict) else {}
        template_id = str(payload.get("template_id") or task.get("template_id") or "__sample__")
        record = AgentMemoryStore(AGENT_MEMORY_DB).remember_teacher_edit(
            template_id=template_id,
            task=task,
            fields=fields,
        )
        self._send(*_json_bytes({"ok": True, "memory": record, "message": "已记住本次修改，下次同类教案会优先参考。"}))

    def _read_agent_preview_payload(self) -> dict:
        content_type = self.headers.get("Content-Type", "")
        if "application/json" in content_type:
            content_length = int(self.headers.get("Content-Length", "0") or "0")
            body = self.rfile.read(content_length).decode("utf-8")
            return json.loads(body or "{}")

        if "multipart/form-data" in content_type:
            form = self._read_multipart_form()
            return {
                "agent_request": _form_value(form, "agent_request", ""),
                "subject": _form_value(form, "subject", ""),
                "grade": _form_value(form, "grade", ""),
                "title": _form_value(form, "title", ""),
                "class_hour": _form_value(form, "class_hour", ""),
                "class_type": _form_value(form, "class_type", ""),
                "teaching_style": _form_value(form, "teaching_style", ""),
                "student_level": _form_value(form, "student_level", ""),
                "generation_depth": _form_value(form, "generation_depth", ""),
            }

        if "application/x-www-form-urlencoded" in content_type:
            content_length = int(self.headers.get("Content-Length", "0") or "0")
            values = parse_qs(self.rfile.read(content_length).decode("utf-8"), keep_blank_values=True)
            return {key: items[0] if items else "" for key, items in values.items()}

        raise ValueError("请使用 JSON 或表单提交")

    def _handle_agent_preview(self) -> None:
        payload = self._read_agent_preview_payload()
        agent_request = str(payload.get("agent_request") or "")
        if not agent_request.strip():
            raise ValueError("请先输入备课需求")

        defaults = {
            "subject": str(payload.get("subject") or ""),
            "grade": str(payload.get("grade") or ""),
            "title": str(payload.get("title") or ""),
            "class_hour": str(payload.get("class_hour") or ""),
            "class_type": str(payload.get("class_type") or ""),
            "teaching_style": str(payload.get("teaching_style") or ""),
            "student_level": str(payload.get("student_level") or ""),
            "generation_depth": str(payload.get("generation_depth") or ""),
        }
        defaults = _infer_single_prompt_defaults(agent_request, defaults)
        defaults["title"] = sanitize_lesson_title(defaults.get("title", ""), agent_request)
        task = route_task(agent_request, defaults)
        _sanitize_agent_task_title(task, agent_request, defaults.get("title", ""))
        plan = build_plan(task)
        self._send(
            *_json_bytes(
                {
                    "agent_task": task.to_dict(),
                    "agent_plan": [step.to_dict() for step in plan],
                    "needs_input": bool(task.missing_fields),
                    "missing_fields": task.missing_fields,
                }
            )
        )

    def _handle_agent_run(self) -> None:
        (
            form,
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            template_path,
            template_analysis,
            strict_ai,
            creative_mode,
            actual_template_mode,
            is_generic_material,
        ) = self._read_agent_form()
        agent_request = _form_value(form, "agent_request", "")
        if not agent_request.strip():
            raise ValueError("请先输入 Agent 指令")
        repeat_fill_mode = _normalize_repeat_fill_mode(
            _form_value(form, "repeat_fill_mode", ""),
            actual_template_mode,
        )

        defaults = {
            "subject": subject,
            "grade": grade,
            "title": title,
            "class_hour": class_hour,
            "class_type": class_type if class_type != DEFAULT_CLASS_TYPE else "",
            "teaching_style": teaching_style if teaching_style != DEFAULT_TEACHING_STYLE else "",
            "student_level": student_level if student_level != DEFAULT_STUDENT_LEVEL else "",
            "generation_depth": generation_depth if generation_depth != DEFAULT_GENERATION_DEPTH else "",
        }
        defaults = _infer_single_prompt_defaults(agent_request, defaults)
        defaults["title"] = sanitize_lesson_title(defaults.get("title", ""), agent_request)
        task = route_task(agent_request, defaults)
        _sanitize_agent_task_title(task, agent_request, defaults.get("title", ""))
        plan = build_plan(task)

        if task.missing_fields:
            self._send(
                *_json_bytes(
                    {
                        "needs_input": True,
                        "agent_task": task.to_dict(),
                        "agent_plan": [step.to_dict() for step in plan],
                        "missing_fields": task.missing_fields,
                        "message": "Agent 需要补齐学科、年级或课题后再执行。",
                    },
                    HTTPStatus.BAD_REQUEST,
                )
            )
            return

        if task.task_type != "lesson_plan":
            self._send(
                *_json_bytes(
                    {
                        "needs_input": True,
                        "agent_task": task.to_dict(),
                        "agent_plan": [step.to_dict() for step in plan],
                        "message": "当前 Agent MVP 只支持教案生成任务。",
                    },
                    HTTPStatus.BAD_REQUEST,
                )
            )
            return

        clean_material = sanitize_material_hint(material, agent_request, task.title)
        lesson_request = LessonRequest(
            task.subject, task.grade, task.title,
            task.class_hour, clean_material,
            task.class_type, task.teaching_style,
            task.student_level, task.generation_depth,
            strict_ai, creative_mode,
        )
        # ── Delegate to new Agent state machine ──
        import uuid
        from .agent_core.checkpoint import AgentCheckpointStore
        from .agent_core.state import AgentRunState
        from .agent_core.tool_registry import build_agent_tool_registry
        from .agent_core.graph_planner import build_graph as build_agent_graph
        from .agent_core.executor import AgentExecutor

        task_dict = task.to_dict()
        task_dict["material"] = clean_material
        task_dict["repeat_fill_mode"] = repeat_fill_mode
        session_id = uuid.uuid4().hex[:12]

        state = AgentRunState(
            session_id=session_id, status="initialized", task=task_dict,
            current_node="", next_action="", template_path=str(template_path),
            template_id=_template_id_for(template_path), template_analysis=template_analysis,
        )

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        registry = build_agent_tool_registry(
            output_dir=OUTPUT_DIR, preview_dir=PREVIEW_DIR,
            history_db=HISTORY_DB, memory_db=AGENT_MEMORY_DB,
        )
        checkpoint = AgentCheckpointStore(OUTPUT_DIR)
        executor = AgentExecutor(registry, checkpoint)
        agent_graph = build_agent_graph(task)
        result = executor.run(agent_graph, state)
        result.state.fields = _normalize_teacher_web_fields(result.state.fields or {}, task, agent_request)

        export_result = result.state.export_result or {}
        llm_status = check_generation_health(probe=False).to_dict()
        generation_backend = "agent"
        for node in agent_graph:
            if node.id != "draft_fields":
                continue
            match = re.search(r"generation_backend:\s*([^,，;；\s]+)", node.detail or "")
            if match:
                generation_backend = match.group(1).strip()
            break
        if not strict_ai and not llm_status.get("configured") and generation_backend == "agent":
            generation_backend = "local_fallback"
        draft_result: dict[str, Any] = {
            "fields": result.state.fields or {},
            "generation_backend": generation_backend,
            "review_report": result.state.review_report,
        }
        teacher_diagnostic_report = build_teacher_diagnostic_report(
            template_analysis=export_result.get("template_analysis") or template_analysis,
            fill_report=export_result.get("fill_report"),
            evaluation_report=result.state.evaluation_report,
            fields=result.state.fields or {},
            template_profile=result.state.template_profile,
        ).to_dict()

        payload: dict[str, Any] = {
            **draft_result,
            "session_id": session_id,
            "status": result.status,
            "next_action": result.next_action,
            "fields": result.state.fields or {},
            "template_fields": template_analysis["mapped_fields"],
            "template_analysis": export_result.get("template_analysis") or template_analysis,
            "template_id": _template_id_for(template_path),
            "agent_task": task.to_dict(),
            "agent_trace": result.state.trace,
            "agent_failed": result.failed,
            "evaluation_report": result.state.evaluation_report,
            "output_name": export_result.get("output_name"),
            "download_url": export_result.get("download_url"),
            "preview_url": export_result.get("preview_url"),
            "fill_report": export_result.get("fill_report"),
            "output_quality_report": export_result.get("output_quality_report"),
            "teacher_diagnostic_report": teacher_diagnostic_report,
            "teacher_report": result.state.teacher_report,
            "workflow_trace": result.state.trace,
            "template_mode": actual_template_mode,
            "repeat_fill_mode": repeat_fill_mode,
            "is_generic_material": is_generic_material,
            "llm_status": llm_status,
            "mode": actual_template_mode,
            "visible_sections": ["fields", "trace", "download"],
            "technical_details_available": True,
            "beginner_summary": _beginner_summary(
                evaluation_report=result.state.evaluation_report,
                is_generic_material=is_generic_material,
                template_mode=actual_template_mode,
            ),
            "professional_diagnostics": {
                "table_mappings": template_analysis.get("table_mappings", {}),
                "field_write_counts": (export_result.get("fill_report") or {}).get("field_write_counts", {}),
                "template_errors": template_analysis.get("errors", []),
                "output_quality_report": export_result.get("output_quality_report"),
            },
        }
        if result.failed:
            self._send(*_json_bytes(payload, HTTPStatus.INTERNAL_SERVER_ERROR))
        else:
            self._send(*_json_bytes(payload, HTTPStatus.OK))

    def _handle_export(self) -> None:
        content_length = int(self.headers.get("Content-Length", "0") or "0")
        payload = json.loads(self.rfile.read(content_length).decode("utf-8"))

        fields = payload.get("fields")
        template_id = payload.get("template_id")
        if not isinstance(fields, dict):
            raise ValueError("缺少教案字段")
        if not isinstance(template_id, str):
            raise ValueError("缺少模板信息")

        request_context = payload.get("request_context") if isinstance(payload.get("request_context"), dict) else {}
        agent_request = str(
            request_context.get("raw_request")
            or request_context.get("agent_request")
            or payload.get("agent_request")
            or ""
        )
        fallback_title = str(request_context.get("title") or fields.get("title") or "")
        fields = dict(fields)
        fields["lesson_title"] = sanitize_lesson_title(
            str(fields.get("lesson_title") or ""),
            agent_request,
            fallback_title,
        )
        fields = normalize_lesson_field_aliases(fields, agent_request)

        template_path = _template_from_id(template_id)
        repeat_fill_mode = _normalize_repeat_fill_mode(
            str(payload.get("repeat_fill_mode") or request_context.get("repeat_fill_mode") or ""),
            _template_mode_label(template_path),
        )
        workflow = TeacherWorkflow()
        result = workflow.export_document(
            fields,
            template_path,
            OUTPUT_DIR,
            PREVIEW_DIR,
            repeat_fill_mode=repeat_fill_mode,
        )
        output_path = OUTPUT_DIR / str(result.get("output_name") or "")
        evaluation = evaluate_lesson_output(
            fields=fields,
            output_path=output_path,
            download_url=result.get("download_url"),
            template_analysis=result.get("template_analysis"),
            fill_report=result.get("fill_report"),
        )
        result["evaluation_report"] = evaluation.to_dict()
        teacher_diagnostic = build_teacher_diagnostic_report(
            template_analysis=result.get("template_analysis"),
            fill_report=result.get("fill_report"),
            evaluation_report=result["evaluation_report"],
            fields=fields,
        )
        result["teacher_diagnostic_report"] = teacher_diagnostic.to_dict()
        result["teacher_report"] = {"summary": teacher_diagnostic.to_markdown(), "passed": teacher_diagnostic.status == "passed"}

        prior_trace = payload.get("workflow_trace")
        if isinstance(prior_trace, list):
            result["workflow_trace"] = prior_trace + result["workflow_trace"]

        review_report = payload.get("review_report") if isinstance(payload.get("review_report"), dict) else {}
        generation_backend = str(payload.get("generation_backend") or "manual_export")
        template_mode = result["template_analysis"].get("mode", "unknown")
        result["workflow_trace"].append(
            {
                "node": "history_store",
                "label": "历史记录",
                "status": "done",
                "detail": "已写入本地 SQLite 历史库，便于教师找回最近导出。",
                "elapsed_ms": result["workflow_trace"][-1]["elapsed_ms"] if result["workflow_trace"] else 0,
            }
        )
        history_item = HistoryStore(HISTORY_DB).save_document(
            fields=fields,
            request_context=request_context,
            generation_backend=generation_backend,
            template_mode=template_mode,
            output_name=result["output_name"],
            download_url=result["download_url"],
            preview_url=result["preview_url"],
            review_report=review_report,
            workflow_trace=result["workflow_trace"],
        )
        result["history_item"] = history_item

        if evaluation.passed and (result.get("fill_report") or {}).get("filled_non_empty_count", 0) > 0:
            try:
                profile_store = TemplateProfileStore(OUTPUT_DIR)
                profile_id = profile_store.template_fingerprint(template_path)
                profile_store.save_successful_mapping(
                    profile_id,
                    (result.get("template_analysis") or {}).get("table_mappings", {}),
                    result.get("fill_report") or {},
                    mapped_fields=(result.get("template_analysis") or {}).get("mapped_fields", []),
                    repeat_fill_mode=repeat_fill_mode,
                    known_risks=teacher_diagnostic.reasons,
                )
            except Exception:
                pass

        if not evaluation.passed:
            result["error"] = evaluation.summary
            self._send(*_json_bytes(result, HTTPStatus.BAD_REQUEST))
            return

        self._send(*_json_bytes(result))

    def _handle_generate(self) -> None:
        (
            _,
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            strict_ai,
            creative_mode,
            template_path,
            template_analysis,
        ) = self._read_lesson_form()

        request = LessonRequest(
            subject,
            grade,
            title,
            class_hour,
            material,
            class_type,
            teaching_style,
            student_level,
            generation_depth,
            strict_ai,
            creative_mode,
        )
        workflow = TeacherWorkflow(history_db=HISTORY_DB)
        draft_result = workflow.draft(request, template_path, _template_id_for(template_path), template_analysis)
        export_result = workflow.export_document(draft_result["fields"], template_path, OUTPUT_DIR, PREVIEW_DIR)
        output_path = OUTPUT_DIR / str(export_result.get("output_name") or "")
        evaluation = evaluate_lesson_output(
            fields=draft_result["fields"],
            output_path=output_path,
            download_url=export_result.get("download_url"),
            template_analysis=export_result.get("template_analysis"),
            fill_report=export_result.get("fill_report"),
        )
        template_mode = export_result["template_analysis"].get("mode", "unknown")
        export_result["workflow_trace"].append(
            {
                "node": "history_store",
                "label": "历史记录",
                "status": "done",
                "detail": "已写入本地 SQLite 历史库，便于教师找回最近导出。",
                "elapsed_ms": export_result["workflow_trace"][-1]["elapsed_ms"] if export_result["workflow_trace"] else 0,
            }
        )
        history_item = HistoryStore(HISTORY_DB).save_document(
            fields=draft_result["fields"],
            request_context=request.to_dict(),
            generation_backend=str(draft_result["generation_backend"]),
            template_mode=template_mode,
            output_name=export_result["output_name"],
            download_url=export_result["download_url"],
            preview_url=export_result["preview_url"],
            review_report=draft_result["review_report"],
            workflow_trace=export_result["workflow_trace"],
        )

        payload = {
            **draft_result,
            "template_fields": template_analysis["mapped_fields"],
            "template_analysis": export_result["template_analysis"],
            "output_name": export_result["output_name"],
            "download_url": export_result["download_url"],
            "preview_url": export_result["preview_url"],
            "fill_report": export_result.get("fill_report"),
            "output_quality_report": export_result.get("output_quality_report"),
            "evaluation_report": evaluation.to_dict(),
            "workflow_trace": export_result["workflow_trace"],
            "history_item": history_item,
            "llm_status": check_generation_health(probe=False).to_dict(),
        }
        if not evaluation.passed:
            payload["error"] = evaluation.summary
            self._send(*_json_bytes(payload, HTTPStatus.BAD_REQUEST))
            return
        self._send(*_json_bytes(payload))

    def _save_template(self, form: cgi.FieldStorage, allow_sample: bool = False) -> Path:
        template_item = form["template"] if "template" in form else None
        if template_item is None or not template_item.filename:
            if allow_sample:
                if not SAMPLE_TEMPLATE.exists():
                    create_sample_template(SAMPLE_TEMPLATE)
                return SAMPLE_TEMPLATE
            raise ValueError("请上传学校 Word 模板")

        original_name = Path(template_item.filename).name
        if not original_name.lower().endswith(".docx"):
            raise ValueError("请上传 .docx 模板")

        UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        upload_name = unique_upload_name(original_name)
        upload_path = UPLOAD_DIR / upload_name
        with upload_path.open("wb") as file:
            file.write(template_item.file.read())
        return upload_path


def run(host: str = "127.0.0.1", port: int = 8765) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    if not SAMPLE_TEMPLATE.exists():
        create_sample_template(SAMPLE_TEMPLATE)

    server = ThreadingHTTPServer((host, port), TeacherAgentHandler)
    print(f"教案助手 Teacher Skill 已启动：http://{host}:{port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("教案助手 Teacher Skill 已停止。")
    finally:
        server.server_close()


if __name__ == "__main__":
    run()

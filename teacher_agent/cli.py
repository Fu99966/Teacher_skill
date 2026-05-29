from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from .deepseek_client import DeepSeekError
from .docx_filler import fill_docx_template
from .lesson_generator import (
    DEFAULT_CLASS_TYPE,
    DEFAULT_GENERATION_DEPTH,
    DEFAULT_STUDENT_LEVEL,
    DEFAULT_TEACHING_STYLE,
    LessonGenerationError,
    build_lesson_prompt,
    draft_lesson_document_fields_with_source,
    write_lesson_json,
)
from .sample_template import create_sample_template
from .template_parser import analyze_template


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


def _read_text(path: str | Path) -> str:
    try:
        return Path(path).read_text(encoding="utf-8")
    except FileNotFoundError as exc:
        raise ValueError(f"教材文件不存在：{path}") from exc


def _load_json(path: str | Path) -> dict:
    try:
        data = json.loads(Path(path).read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ValueError(f"JSON 数据文件不存在：{path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON 数据文件格式错误：{exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("JSON 数据必须是对象，不能是数组或纯文本。")
    return data


def _template_fields(template: str | None) -> tuple[list[str] | None, dict | None]:
    if not template:
        return None, None
    analysis = analyze_template(template)
    if analysis.get("needs_template_markers"):
        raise ValueError("; ".join(analysis.get("errors") or ["模板中未识别到可填字段。"]))
    return analysis["mapped_fields"], analysis


def _print_json(payload: dict) -> None:
    print(json.dumps(payload, ensure_ascii=False, indent=2))


def cmd_scan_template(args: argparse.Namespace) -> None:
    analysis = analyze_template(args.template)
    _print_json({"template": args.template, **analysis})


def cmd_fill_template(args: argparse.Namespace) -> None:
    data = _load_json(args.data)
    report = fill_docx_template(args.template, data, args.output)
    payload = {"success": not report.errors, "output": str(report.path), "fill_report": report.to_dict()}
    _print_json(payload)
    if report.errors:
        raise ValueError("; ".join(report.errors))


def cmd_draft_lesson(args: argparse.Namespace) -> None:
    material = _read_text(args.material_file)
    template_fields, template_analysis = _template_fields(args.template)
    fields, backend = draft_lesson_document_fields_with_source(
        args.subject,
        args.grade,
        args.title,
        material,
        args.class_hour,
        args.class_type,
        args.teaching_style,
        args.student_level,
        args.generation_depth,
        template_fields,
        args.strict_ai,
        (template_analysis or {}).get("field_context"),
    )
    output = write_lesson_json(fields, args.output)
    _print_json(
        {
            "output": str(output),
            "generation_backend": backend,
            "template_fields": template_fields,
            "fields": fields,
        }
    )


def cmd_generate(args: argparse.Namespace) -> None:
    material = _read_text(args.material_file)
    template_fields, template_analysis = _template_fields(args.template)
    fields, backend = draft_lesson_document_fields_with_source(
        args.subject,
        args.grade,
        args.title,
        material,
        args.class_hour,
        args.class_type,
        args.teaching_style,
        args.student_level,
        args.generation_depth,
        template_fields,
        args.strict_ai,
        (template_analysis or {}).get("field_context"),
    )
    report = fill_docx_template(args.template, fields, args.output)
    payload = {
        "success": not report.errors,
        "output": str(report.path),
        "generation_backend": backend,
        "template_fields": template_fields,
        "fields": fields,
        "fill_report": report.to_dict(),
    }
    _print_json(payload)
    if report.errors:
        raise ValueError("; ".join(report.errors))


def cmd_prompt(args: argparse.Namespace) -> None:
    material = _read_text(args.material_file)
    template_fields, template_analysis = _template_fields(args.template)
    prompt = build_lesson_prompt(
        args.subject,
        args.grade,
        args.title,
        material,
        args.class_hour,
        args.class_type,
        args.teaching_style,
        args.student_level,
        args.generation_depth,
        template_fields,
        (template_analysis or {}).get("field_context"),
    )
    print(prompt)


def cmd_create_sample_template(args: argparse.Namespace) -> None:
    output = create_sample_template(args.output)
    print(f"Created sample template: {output}")


def cmd_diagnose_template(args: argparse.Namespace) -> None:
    """Diagnose a real template: analyze structure, generate fields, fill, and report."""
    from .docx_filler import fill_docx_template
    from .template_parser import analyze_template

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Step 1: Parse template
    analysis = analyze_template(args.template)
    (output_dir / "template_analysis.json").write_text(
        json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Step 2: Write table structure
    table_structure = analysis.get("tables", [])
    (output_dir / "table_structure.json").write_text(
        json.dumps(table_structure, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Step 3: Generate fields
    material = _read_text(args.material_file)
    template_fields, template_analysis = _template_fields(args.template)
    fields, backend = draft_lesson_document_fields_with_source(
        args.subject,
        args.grade,
        args.title,
        material,
        args.class_hour,
        args.class_type,
        args.teaching_style,
        args.student_level,
        args.generation_depth,
        template_fields,
        args.strict_ai,
        (template_analysis or {}).get("field_context"),
    )
    (output_dir / "generated_fields.json").write_text(
        json.dumps(fields, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Step 4: Fill template
    output_docx = output_dir / "output.docx"
    report = fill_docx_template(args.template, fields, output_docx)
    (output_dir / "fill_report.json").write_text(
        json.dumps(report.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Step 5: Extract text from output Word
    from docx import Document
    try:
        output_doc = Document(str(output_docx))
        text_parts: list[str] = []
        for paragraph in output_doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        text_parts.append(t)
        output_text = "\n".join(text_parts)
    except Exception as exc:
        output_text = f"[无法读取输出Word文本: {exc}]"

    (output_dir / "output_text.txt").write_text(output_text, encoding="utf-8")

    # Step 6: Root cause analysis
    root_cause_parts: list[str] = [
        "# 诊断报告：root_cause.md",
        "",
        f"## 模板分析",
        f"- 模板文件: {args.template}",
        f"- 识别字段数: {analysis.get('fillable_count', 0)}",
        f"- 识别字段: {', '.join(analysis.get('mapped_fields', []))}",
        f"- 模式: {analysis.get('mode', 'unknown')}",
        f"- 表格数: {analysis.get('table_count', 0)}",
        "",
        f"## 生成结果",
        f"- 生成后端: {backend}",
        f"- 生成字段数: {len(fields)}",
        f"- 非空字段数: {sum(1 for v in fields.values() if v and v.strip())}",
        "",
        f"## 填充结果",
        f"- 已填字段: {', '.join(report.filled_fields) if report.filled_fields else '无'}",
        f"- 空字段(跳过): {', '.join(report.skipped_empty_fields) if report.skipped_empty_fields else '无'}",
        f"- 缺失字段: {', '.join(report.missing_fields) if report.missing_fields else '无'}",
        f"- 未填充字段: {', '.join(report.unfilled_template_fields) if report.unfilled_template_fields else '无'}",
        f"- filled_non_empty_count: {report.filled_non_empty_count}",
        f"- table_write_count: {report.table_write_count}",
        "",
        f"## 根因判断",
    ]

    if report.errors:
        root_cause_parts.append(f"❌ **失败**: {'; '.join(report.errors)}")
    elif analysis.get("needs_template_markers"):
        root_cause_parts.append("❌ **失败**: 字段识别失败 - 模板中未识别到任何可填字段。")
    elif not report.filled_fields:
        root_cause_parts.append("❌ **失败**: 写入定位失败 - 识别到了字段但未能写入任何内容。")
    elif report.filled_non_empty_count == 0:
        root_cause_parts.append("❌ **失败**: 生成字段为空 - 所有字段内容均为空。")
    else:
        # Check output text for key content
        key_checks = []
        title_val = fields.get("lesson_title", "")
        if title_val and title_val.strip():
            found_title = title_val.strip() in output_text
            key_checks.append(f"lesson_title ('{title_val.strip()[:30]}') -> {'✅存在于输出Word' if found_title else '❌未在输出Word中找到'}")
        
        goals_val = fields.get("teaching_goals", "")
        if goals_val and goals_val.strip():
            snippet = goals_val.strip()[:30]
            found = snippet in output_text
            key_checks.append(f"teaching_goals ({snippet}...) -> {'✅存在于输出Word' if found else '❌未在输出Word中找到'}")
        
        if key_checks:
            root_cause_parts.append("\n### 关键词验证\n" + "\n".join(f"- {c}" for c in key_checks))
        
        all_found = all("✅" in c for c in key_checks) if key_checks else False
        if all_found:
            root_cause_parts.append("\n✅ **成功**: 内容已写入Word文档。")
        else:
            root_cause_parts.append("\n⚠️ **部分成功**: 填充报告显示有内容填入，但部分关键词未在输出Word中验证到。")

    (output_dir / "root_cause.md").write_text("\n".join(root_cause_parts), encoding="utf-8")

    # Print summary
    payload = {
        "template": args.template,
        "output_dir": str(output_dir),
        "analysis": {
            "mapped_fields": analysis.get("mapped_fields", []),
            "fillable_count": analysis.get("fillable_count", 0),
            "mode": analysis.get("mode"),
        },
        "generation_backend": backend,
        "fill_report": report.to_dict(),
        "output_text_preview": output_text[:500] if output_text else "(空)",
    }
    _print_json(payload)
    if report.errors:
        raise ValueError("; ".join(report.errors))


def cmd_web(args: argparse.Namespace) -> None:
    from .web_app import run

    run(args.host, args.port)


def _add_lesson_args(parser: argparse.ArgumentParser) -> None:
    parser.add_argument("--subject", required=True)
    parser.add_argument("--grade", required=True)
    parser.add_argument("--title", required=True)
    parser.add_argument("--material-file", required=True)
    parser.add_argument("--class-hour", default="1课时")
    parser.add_argument("--class-type", default=DEFAULT_CLASS_TYPE)
    parser.add_argument("--teaching-style", default=DEFAULT_TEACHING_STYLE)
    parser.add_argument("--student-level", default=DEFAULT_STUDENT_LEVEL)
    parser.add_argument("--generation-depth", default=DEFAULT_GENERATION_DEPTH)
    strict = parser.add_mutually_exclusive_group()
    strict.add_argument("--strict-ai", dest="strict_ai", action="store_true", help="DeepSeek 失败时直接报错")
    strict.add_argument("--no-strict-ai", dest="strict_ai", action="store_false", help="DeepSeek 失败时使用本地 fallback")
    parser.set_defaults(strict_ai=True)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Teacher Skill document tools")
    subparsers = parser.add_subparsers(required=True)

    scan = subparsers.add_parser("scan-template", help="scan placeholders and table fields in a docx template")
    scan.add_argument("template")
    scan.set_defaults(func=cmd_scan_template)

    fill = subparsers.add_parser("fill-template", help="fill a docx template with JSON data")
    fill.add_argument("--template", required=True)
    fill.add_argument("--data", required=True)
    fill.add_argument("--output", required=True)
    fill.set_defaults(func=cmd_fill_template)

    draft = subparsers.add_parser("draft-lesson", help="draft lesson fields JSON")
    _add_lesson_args(draft)
    draft.add_argument("--template", help="optional docx template; dynamic fields come from this template")
    draft.add_argument("--output", required=True)
    draft.set_defaults(func=cmd_draft_lesson)

    generate = subparsers.add_parser("generate", help="scan template, draft fields, fill Word, and report")
    _add_lesson_args(generate)
    generate.add_argument("--template", required=True)
    generate.add_argument("--output", required=True)
    generate.set_defaults(func=cmd_generate)

    prompt = subparsers.add_parser("lesson-prompt", help="print an LLM prompt for lesson JSON")
    _add_lesson_args(prompt)
    prompt.add_argument("--template", help="optional docx template; dynamic fields come from this template")
    prompt.set_defaults(func=cmd_prompt)

    sample = subparsers.add_parser("create-sample-template", help="create a sample placeholder docx")
    sample.add_argument("--output", required=True)
    sample.set_defaults(func=cmd_create_sample_template)

    web = subparsers.add_parser("web", help="start the teacher web app")
    web.add_argument("--host", default="127.0.0.1")
    web.add_argument("--port", type=int, default=8765)
    web.set_defaults(func=cmd_web)

    diagnose = subparsers.add_parser("diagnose-template", help="diagnose a real template: analyze, generate, fill, and report")
    diagnose.add_argument("--template", required=True)
    _add_lesson_args(diagnose)
    diagnose.add_argument("--output-dir", required=True)
    diagnose.set_defaults(func=cmd_diagnose_template)

    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    try:
        args.func(args)
    except (ValueError, DeepSeekError, LessonGenerationError) as exc:
        print(f"错误：{getattr(exc, 'user_message', str(exc))}", file=sys.stderr)
        raise SystemExit(1) from exc


if __name__ == "__main__":
    main()

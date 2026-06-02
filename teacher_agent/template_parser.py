from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn

from .docx_grid import GridCell, parse_table_grid, find_cell_by_grid

PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}\r\n\t<>]+?)\s*\}\}")

# ── Field label mapping ──
FIELD_LABELS: dict[str, tuple[str, ...]] = {
    "lesson_title": ("课题", "课题名称", "题目", "课题（含章节号）", "课题(含章节号)"),
    "subject": ("学科", "课程", "课程名称", "科目"),
    "grade": ("年级", "授课年级", "适用年级"),
    "class_name": ("班级", "授课班级"),
    "class_type": ("授课类型", "课型"),
    "class_hour": ("课时", "课时安排", "授课课时", "学时", "课时数"),
    "teaching_date": ("授课日期",),
    "teaching_goals": (
        "教学目的", "教学目的与要求", "教学目的及要求",
        "教学目标", "教学目标与要求", "教学目标及要求",
        "学习目标", "核心素养目标", "知识目标", "能力目标", "情感目标",
    ),
    "key_points": ("教学重点", "重点", "重点内容"),
    "difficult_points": ("教学难点", "难点", "难点内容"),
    "teaching_key_difficult": ("教学重难点", "重难点", "教学重点与难点", "重点难点", "重点和难点"),
    "teaching_preparation": ("教学准备", "课前准备", "教具准备", "教学资源"),
    "teaching_environment": ("对教学环境的要求", "教学环境", "教学环境要求"),
    "teaching_aids": ("教具挂图", "教具", "挂图", "教学用具", "教学挂图"),
    "student_analysis": ("学情分析", "学生分析", "学习者分析"),
    "teaching_process": (
        "主要教学内容", "教学内容", "主要教学内容安排",
        "教学过程", "教学流程", "教学环节", "课堂过程", "过程设计", "教学活动",
    ),
    "teaching_method": (
        "教学方法的运用", "教学方法运用", "教学方法",
        "教法", "学法", "教学方式",
    ),
    "teacher_activity": ("教师活动", "教师行为", "教师指导"),
    "student_activity": ("学生活动", "学生行为", "学习活动"),
    "design_intent": ("设计意图", "设计说明", "活动意图"),
    "blackboard_design": ("板书设计", "板书", "板书规划"),
    "homework": ("作业设计", "作业", "课后作业", "课后任务"),
    "reflection": ("课后小记", "课后小结", "教学反思", "课后反思", "教学后记"),
}

NEXT_ROW_PREFERRED_FIELDS = {"teaching_process", "teaching_method"}

REQUIRED_FIELDS = {"teaching_method", "teaching_process"}

FIELD_LABELS_BY_ALIAS: list[tuple[str, str]] = sorted(
    [(f, a) for f, aliases in FIELD_LABELS.items() for a in aliases],
    key=lambda x: -len(x[1]),
)

DIRECT_LABEL_FIELDS = {
    "教材分析", "课程标准", "核心素养", "学习任务", "学习活动",
    "评价任务", "评价标准", "教学评价", "教学资源", "信息技术应用",
    "课堂小结", "课后拓展", "安全教育", "德育渗透", "跨学科融合",
    "二次备课", "个性化修改", "教研组意见", "审批意见",
    "授课日期", "授课班级", "授课类型", "准备",
}


# ── Helpers ──

def _normalize_label(text: str) -> str:
    text = PLACEHOLDER_PATTERN.sub("", text or "")
    text = re.sub(r"[\s:：；;、，,。.·\-—–（）()\[\]【】<>《》]+", "", text)
    return text.strip()


def _is_blankish(text: str) -> bool:
    normalized = _normalize_label(text)
    return not normalized or set(normalized) <= {"_", "-", "—", "一", ".", "。"}


def _match_field(label: str) -> str | None:
    normalized = _normalize_label(label)
    if not normalized:
        return None
    for field, alias in FIELD_LABELS_BY_ALIAS:
        if normalized == _normalize_label(alias) or _normalize_label(alias) in normalized:
            return field
    return None


def _direct_field_from_label(label: str) -> str | None:
    normalized = _normalize_label(label)
    if not normalized:
        return None
    if normalized in DIRECT_LABEL_FIELDS:
        return normalized
    if re.search(r"[\u4e00-\u9fff]", normalized) and 2 <= len(normalized) <= 18:
        return normalized
    return None


def _add_ordered(target: list[str], field: str) -> None:
    if field and field not in target:
        target.append(field)


def find_placeholders_in_text(text: str) -> list[str]:
    fields: list[str] = []
    for match in PLACEHOLDER_PATTERN.finditer(text or ""):
        field = re.sub(r"[\r\n\t<>]", "", match.group(1).strip())
        if field and field not in fields:
            fields.append(field)
    return fields


# ── Fill target selection ──

def _choose_table_target(
    field: str,
    grid: list[list[GridCell]],
    label_row: int,
    label_grid_col: int,
    label_grid_span: int,
) -> tuple[int, int, int | None, str, str]:
    """Returns (target_grid_col, target_grid_span, target_row_override, mapping_type, target_type)."""
    num_cols = len(grid[0]) if grid else 0

    def _blank_cell_on_row(row_idx: int, col_start: int, col_end: int) -> tuple[int, int] | None:
        if row_idx < 0 or row_idx >= len(grid):
            return None
        unique_cells: list[GridCell] = []
        seen: set[tuple[int, int]] = set()
        for gcell in grid[row_idx]:
            if gcell is None:
                continue
            key = (gcell.row, gcell.grid_col)
            if key in seen:
                continue
            seen.add(key)
            if _is_blankish(gcell.text):
                unique_cells.append(gcell)
        for gcell in unique_cells:
            if gcell.grid_col == col_start:
                return (gcell.grid_col, gcell.grid_span)
        for gcell in unique_cells:
            if gcell.grid_col >= col_start and gcell.grid_col < num_cols:
                return (gcell.grid_col, gcell.grid_span)
        for gcell in unique_cells:
            if gcell.grid_col <= col_start < gcell.grid_col + gcell.grid_span:
                return (gcell.grid_col, gcell.grid_span)
        return None

    # A) NEXT_ROW_PREFERRED: check next-row first
    if field in NEXT_ROW_PREFERRED_FIELDS:
        nr = label_row + 1
        if nr < len(grid):
            next_target = _blank_cell_on_row(nr, label_grid_col, label_grid_col + label_grid_span)
            if next_target is not None:
                return (next_target[0], next_target[1], nr, "table_cell", "next_row_cell")

    # B) Right-side blank
    rc = label_grid_col + label_grid_span
    if rc < num_cols:
        gcell = grid[label_row][rc]
        if gcell is not None and (_is_blankish(gcell.text) or "{{" in gcell.text):
            return (gcell.grid_col, gcell.grid_span, None, "table_cell", "right_cell")

    # C) Next-row (fallback for non-preferred)
    nr = label_row + 1
    if nr < len(grid):
        next_target = _blank_cell_on_row(nr, label_grid_col, label_grid_col + label_grid_span)
        if next_target is not None:
            return (next_target[0], next_target[1], nr, "table_cell", "next_row_cell")

    # D) Append in-place
    return (label_grid_col, label_grid_span, None, "table_cell_append", "append_label_cell")


# ── Table scanning ──

def _scan_table_labels(
    table,
    *,
    table_index: int,
    location: str,
    section: int | None,
    mapped_fields: list[str],
    field_context: dict[str, list[dict[str, Any]]],
    table_mappings: dict[str, list[dict[str, Any]]],
) -> tuple[list[list[GridCell]], list[list[str]]]:
    grid = parse_table_grid(table)
    num_rows = len(grid)
    num_cols = len(grid[0]) if grid else 0

    row_texts: list[list[str]] = []
    for ri in range(num_rows):
        row_texts.append([grid[ri][c].text if (grid[ri][c] and grid[ri][c].grid_col == c) else "" for c in range(num_cols)])

    scanned: set[int] = set()
    for ri in range(num_rows):
        for gc in range(num_cols):
            gcell = grid[ri][gc]
            if gcell is None or gcell.grid_col != gc:
                continue
            cid = id(gcell.cell._tc) if gcell.cell else (ri * 10000 + gc)
            if cid in scanned:
                continue
            scanned.add(cid)

            text = gcell.text
            for pf in find_placeholders_in_text(text):
                _add_ordered(mapped_fields, pf)
                field_context.setdefault(pf, []).append({"source": "placeholder", "location": location, "section": section, "table": table_index, "row": ri, "grid_col": gc, "text": text})

            field = _match_field(text)
            if not field:
                field = _direct_field_from_label(text)
            if not field:
                continue

            _add_ordered(mapped_fields, field)
            tgt_col, tgt_span, tgt_row_override, map_type, target_type = _choose_table_target(field, grid, ri, gcell.grid_col, gcell.grid_span)

            effective_row = tgt_row_override if tgt_row_override is not None else ri
            effective_target_text = ""
            tgt_physical = gcell.physical_col
            if effective_row < num_rows and tgt_col < num_cols:
                tgt_gcell = grid[effective_row][tgt_col]
                if tgt_gcell is not None:
                    effective_target_text = tgt_gcell.text
                    tgt_physical = tgt_gcell.physical_col

            table_mappings.setdefault(field, []).append({
                "field": field, "type": map_type, "target_type": target_type,
                "location": location, "section": section, "table": table_index,
                "label_row": ri, "label_grid_col": gcell.grid_col, "label_grid_span": gcell.grid_span,
                "row": effective_row, "grid_col": tgt_col, "grid_span": tgt_span,
                "physical_col": tgt_physical, "col": tgt_physical,
                "label": text, "target_text": effective_target_text,
                "required": field in REQUIRED_FIELDS,
            })
            field_context.setdefault(field, []).append({
                "source": "table_label", "location": location, "section": section,
                "table": table_index, "field": field, "label": text,
                "label_row": ri, "label_grid_col": gcell.grid_col,
                "target_row": effective_row, "target_grid_col": tgt_col,
                "target_type": target_type, "mapping_type": map_type,
            })

    return grid, row_texts


# ── Paragraph iteration ──

def _paragraph_location(paragraph) -> str:
    parent = paragraph._p.getparent()
    while parent is not None:
        tag = str(parent.tag)
        if tag.endswith("}hdr"): return "header"
        if tag.endswith("}ftr"): return "footer"
        if tag.endswith("}tc"): return "table"
        parent = parent.getparent()
    return "body"


def iter_paragraphs(document: Document) -> Iterable:
    for p in document.paragraphs: yield p
    for t in document.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: yield p
    for s in document.sections:
        for p in s.header.paragraphs: yield p
        for p in s.footer.paragraphs: yield p
        for t in s.header.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs: yield p
        for t in s.footer.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs: yield p


# ── Public API ──

def scan_template(path: str | Path) -> list[str]:
    document = Document(str(path))
    fields: list[str] = []
    for p in iter_paragraphs(document):
        for f in find_placeholders_in_text(p.text):
            _add_ordered(fields, f)
    return fields


def _scan_placeholders(document: Document) -> tuple[list[str], dict[str, list[dict[str, Any]]]]:
    placeholders: list[str] = []
    occurrences: dict[str, list[dict[str, Any]]] = {}
    for pi, p in enumerate(iter_paragraphs(document)):
        for f in find_placeholders_in_text(p.text):
            _add_ordered(placeholders, f)
            occurrences.setdefault(f, []).append({"source": "placeholder", "location": _paragraph_location(p), "paragraph": pi, "text": p.text})
    return placeholders, occurrences


def analyze_template(path: str | Path) -> dict[str, Any]:
    document = Document(str(path))
    placeholders, field_context = _scan_placeholders(document)
    mapped_fields: list[str] = []
    table_mappings: dict[str, list[dict[str, Any]]] = {}
    tables: list[dict[str, Any]] = []

    for f in placeholders:
        _add_ordered(mapped_fields, f)

    for ti, table in enumerate(document.tables):
        grid, row_texts = _scan_table_labels(table, table_index=ti, location="body_table", section=None, mapped_fields=mapped_fields, field_context=field_context, table_mappings=table_mappings)
        cells = []
        seen_c = set()
        for grid_row in grid:
            for gc, gcell in enumerate(grid_row):
                if gcell is None: continue
                cid = (gcell.row, gcell.grid_col)
                if cid in seen_c: continue
                seen_c.add(cid)
                cells.append({"row": gcell.row, "physical_col": gcell.physical_col, "grid_col": gcell.grid_col, "grid_span": gcell.grid_span, "text": gcell.text, "normalized_text": gcell.normalized_text})
        tables.append({"index": ti, "location": "body_table", "rows": row_texts, "grid_cells": cells, "num_rows": len(grid), "num_cols": len(grid[0]) if grid else 0})

    for si, section in enumerate(document.sections):
        for ti, table in enumerate(section.header.tables):
            grid, row_texts = _scan_table_labels(table, table_index=ti, location="header_table", section=si, mapped_fields=mapped_fields, field_context=field_context, table_mappings=table_mappings)
            tables.append({"index": ti, "section": si, "location": "header_table", "rows": row_texts})
        for ti, table in enumerate(section.footer.tables):
            grid, row_texts = _scan_table_labels(table, table_index=ti, location="footer_table", section=si, mapped_fields=mapped_fields, field_context=field_context, table_mappings=table_mappings)
            tables.append({"index": ti, "section": si, "location": "footer_table", "rows": row_texts})

    errors: list[str] = []
    if not mapped_fields:
        errors.append('未识别到可填写字段。请在模板中加入 {{field_name}} 占位符，或使用可识别的表格标签，例如"教学目标""教学过程""作业设计"。')

    # Mark required fields
    required_field_list: list[str] = [f for f in mapped_fields if f in REQUIRED_FIELDS]

    return {
        "placeholders": placeholders, "placeholder_occurrences": field_context,
        "table_mappings": table_mappings, "mapped_fields": mapped_fields,
        "field_context": field_context, "tables": tables,
        "table_count": len(document.tables), "paragraph_count": len(document.paragraphs),
        "mode": "mixed" if placeholders and table_mappings else ("placeholder" if placeholders else "table_mapping"),
        "fillable_count": len(mapped_fields), "needs_template_markers": not mapped_fields,
        "required_fields": required_field_list,
        "warnings": [], "errors": errors,
    }

from __future__ import annotations

import time
from dataclasses import asdict, dataclass
from pathlib import Path
import re
from urllib.parse import quote

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from .agent_observer import build_teacher_diagnostic_report
from .docx_filler import fill_docx_template
from .few_shot_examples import select_few_shot_examples
from .history_store import HistoryStore
from .lesson_generator import (
    clean_cn_punctuation,
    draft_lesson_document_fields_with_source,
    is_pcb_project_lesson,
    is_stm32_smart_car_project_lesson,
    normalize_lesson_field_aliases,
    sanitize_lesson_title,
)
from .preview_renderer import render_docx_pdf_preview
from .rag_context import build_knowledge_context
from .teacher_agents import review_lesson_quality, revise_lesson_after_review
from .template_parser import analyze_template


@dataclass
class LessonRequest:
    subject: str
    grade: str
    title: str
    class_hour: str
    material: str
    class_type: str
    teaching_style: str
    student_level: str
    generation_depth: str
    strict_ai: bool = False
    creative_mode: str = "常规稳妥"

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class WorkflowTraceEvent:
    node: str
    label: str
    status: str
    detail: str
    elapsed_ms: int

    def to_dict(self) -> dict:
        return asdict(self)


PCB_HOUR_ALLOCATION_ROWS = [
    ("第一阶段", "项目导入与 PCB 基础认知", "4课时", "项目任务书、PCB设计流程图"),
    ("第二阶段", "原理图设计与元件封装检查", "6课时", "原理图文件、封装检查记录"),
    ("第三阶段", "PCB布局与布线规范训练", "8课时", "PCB布局布线文件"),
    ("第四阶段", "DRC检查与问题修改", "6课时", "DRC检查记录、修改说明"),
    ("第五阶段", "Gerber文件输出与项目文档整理", "4课时", "Gerber文件、BOM表、项目说明书"),
    ("第六阶段", "作品展示、互评与总结提升", "4课时", "展示汇报、评价表、反思记录"),
]

STM32_SMART_CAR_HOUR_ALLOCATION_ROWS = [
    ("第一阶段", "项目导入与STM32智能小车结构认知", "4课时", "项目任务书、系统结构图"),
    ("第二阶段", "STM32开发环境搭建与基础外设训练", "6课时", "工程文件、GPIO/定时器实验记录"),
    ("第三阶段", "电机驱动与PWM调速控制", "6课时", "电机控制程序、调速测试记录"),
    ("第四阶段", "循迹与避障传感器调试", "6课时", "传感器测试记录、循迹/避障程序"),
    ("第五阶段", "智能小车综合联调与故障排查", "6课时", "综合运行视频、问题修改记录"),
    ("第六阶段", "作品展示、评价反馈与总结提升", "4课时", "展示汇报、评价表、个人反思"),
]


def _is_system_template_path(template_path: Path) -> bool:
    return Path(template_path).name == "sample_lesson_template.docx"


def _remove_text_stage_rows(lines: list[str]) -> list[str]:
    stage_pattern = re.compile(r"^第[一二三四五六]阶段[：:].*")
    return [line for line in lines if not stage_pattern.match(line.strip())]


def _split_top_level_process_lines(text: str) -> list[str]:
    marked = re.sub(r"(?<!^)(?=[一二三四五六]、|教材依据[：:])", "\n", text)
    return [line.strip() for line in marked.splitlines() if line.strip()]


def _split_hour_allocation_sections(text: str) -> tuple[list[str], list[str]]:
    match = re.search(r"二、课时分配表[：:]?", text)
    if not match:
        return _split_top_level_process_lines(text), []

    before = text[: match.start()].strip()
    after = text[match.end() :].strip()
    next_section = re.search(r"(?=三、)", after)
    post = after[next_section.start() :].strip() if next_section else ""

    pre_lines = _split_top_level_process_lines(before)
    pre_lines.append("二、课时分配表：见下表。")
    post_lines = _split_top_level_process_lines(post)
    return pre_lines, post_lines


def _insert_paragraph_after_table(table, text: str, template_paragraph) -> Paragraph:
    paragraph_element = OxmlElement("w:p")
    table._tbl.addnext(paragraph_element)
    paragraph = Paragraph(paragraph_element, template_paragraph._parent)
    paragraph.style = template_paragraph.style
    paragraph.paragraph_format.line_spacing = template_paragraph.paragraph_format.line_spacing
    paragraph.paragraph_format.space_after = template_paragraph.paragraph_format.space_after
    paragraph.paragraph_format.left_indent = template_paragraph.paragraph_format.left_indent
    paragraph.add_run(text)
    return paragraph


def _insert_table_after_paragraph(document: Document, paragraph, rows: list[tuple[str, str, str, str]]) -> object:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header = ["阶段", "主要内容", "课时", "阶段产出"]
    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            row.cells[index].text = value

    for row in table.rows:
        for cell in row.cells:
            for paragraph_in_cell in cell.paragraphs:
                paragraph_in_cell.paragraph_format.space_after = Pt(0)

    table_element = table._tbl
    table_element.getparent().remove(table_element)
    paragraph._p.addnext(table_element)
    return table


def _enhance_system_template_docx(output_path: Path, fields: dict, template_path: Path) -> bool:
    if not _is_system_template_path(template_path):
        return False
    lesson_title = str(fields.get("lesson_title") or "")
    class_hour = str(fields.get("class_hour") or "")
    class_type = str(fields.get("class_type") or "")
    if is_pcb_project_lesson(lesson_title, class_hour, class_type):
        allocation_rows = PCB_HOUR_ALLOCATION_ROWS
    elif is_stm32_smart_car_project_lesson(lesson_title, class_hour, class_type):
        allocation_rows = STM32_SMART_CAR_HOUR_ALLOCATION_ROWS
    else:
        return False

    document = Document(str(output_path))
    inserted = False
    for paragraph in document.paragraphs:
        text = clean_cn_punctuation(paragraph.text)
        if "二、课时分配表" not in text:
            continue

        source_text = clean_cn_punctuation(str(fields.get("teaching_process") or text))
        if "二、课时分配表" not in source_text:
            source_text = text
        pre_lines, post_lines = _split_hour_allocation_sections(source_text)
        if not any("二、课时分配表" in line for line in pre_lines):
            continue
        post_lines = _remove_text_stage_rows(post_lines)
        paragraph.text = "\n".join(pre_lines)
        table = _insert_table_after_paragraph(document, paragraph, allocation_rows)
        post_text = "\n".join(line for line in post_lines if line.strip())
        if post_text:
            _insert_paragraph_after_table(table, clean_cn_punctuation(post_text), paragraph)
        inserted = True
        break

    if inserted:
        document.save(str(output_path))
    return inserted


def build_workflow_schema() -> dict:
    return {
        "version": "Teacher Skill V9",
        "name": "Lesson Assistant Workflow",
        "nodes": [
            {"id": "app_input", "label": "应用输入", "layer": "应用层"},
            {"id": "template_analyzer", "label": "模板解析", "layer": "编排层"},
            {"id": "knowledge_context", "label": "RAG 上下文", "layer": "知识层"},
            {"id": "anti_repetition", "label": "历史反重复", "layer": "质量层"},
            {"id": "lesson_writer", "label": "执教老师 Agent", "layer": "Agent 层"},
            {"id": "teaching_reviewer", "label": "教研组长 Agent", "layer": "Agent 层"},
            {"id": "lesson_reviser", "label": "二次修订 Agent", "layer": "Agent 层"},
            {"id": "doc_renderer", "label": "Word 渲染器", "layer": "工具层"},
            {"id": "history_store", "label": "历史记录", "layer": "数据层"},
        ],
        "edges": [
            ["app_input", "template_analyzer"],
            ["template_analyzer", "knowledge_context"],
            ["knowledge_context", "anti_repetition"],
            ["anti_repetition", "lesson_writer"],
            ["lesson_writer", "teaching_reviewer"],
            ["teaching_reviewer", "lesson_reviser"],
            ["lesson_reviser", "doc_renderer"],
            ["doc_renderer", "history_store"],
        ],
    }


def _safe_filename(value: str, fallback: str = "lesson") -> str:
    import re

    value = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", value).strip(" ._")
    return value[:80] or fallback


class TeacherWorkflow:
    def __init__(self, history_db: Path | None = None) -> None:
        self._started_at = time.perf_counter()
        self.trace: list[WorkflowTraceEvent] = []
        self.history_db = history_db

    def _mark(self, node: str, label: str, status: str, detail: str) -> None:
        elapsed_ms = int((time.perf_counter() - self._started_at) * 1000)
        self.trace.append(WorkflowTraceEvent(node, label, status, detail, elapsed_ms))

    def draft(
        self,
        request: LessonRequest,
        template_path: Path,
        template_id: str,
        template_analysis: dict | None = None,
    ) -> dict:
        self._mark("app_input", "应用输入", "done", "已接收课程信息、Word 模板和可选补充资料。")

        template_analysis = template_analysis or analyze_template(template_path)
        if template_analysis.get("needs_template_markers"):
            message = "; ".join(template_analysis.get("errors") or ["模板中未识别到可填字段。"])
            raise ValueError(message)
        mode = "占位符" if template_analysis["placeholders"] else "表格标签"
        self._mark(
            "template_analyzer",
            "模板解析",
            "done",
            f"已识别 {len(template_analysis['mapped_fields'])} 个字段，采用{mode}映射。",
        )

        knowledge_context = build_knowledge_context(
            request.material,
            subject=request.subject,
            title=request.title,
            class_type=request.class_type,
            teaching_style=request.teaching_style,
        )
        self._mark("knowledge_context", "RAG 上下文", "done", knowledge_context.source_summary)

        enhanced_material = knowledge_context.enhanced_material(request.material)
        anti_repetition_context = self._build_anti_repetition_context(request)
        if anti_repetition_context:
            self._mark(
                "anti_repetition",
                "历史反重复",
                "done",
                "已读取近期相似教案，生成时会避开重复导入、活动顺序、作业和板书表达。",
            )
        else:
            self._mark("anti_repetition", "历史反重复", "done", "未发现近期相似教案，本次按新任务生成。")
        few_shot_examples = select_few_shot_examples(
            request.subject,
            request.grade,
            request.class_type,
            request.creative_mode,
        )
        template_fields = template_analysis["mapped_fields"] or None
        field_map, generation_backend = draft_lesson_document_fields_with_source(
            request.subject,
            request.grade,
            request.title,
            enhanced_material,
            request.class_hour,
            request.class_type,
            request.teaching_style,
            request.student_level,
            request.generation_depth,
            template_fields,
            request.strict_ai,
            template_analysis.get("field_context"),
            request.creative_mode,
            anti_repetition_context,
            few_shot_examples,
        )
        self._mark("lesson_writer", "执教老师 Agent", "done", f"已根据模板生成 {len(field_map)} 个字段。")

        context = {
            **request.to_dict(),
            "template_fields": template_analysis["mapped_fields"],
            "knowledge_summary": knowledge_context.source_summary,
            "anti_repetition_context": anti_repetition_context,
            "creative_mode": request.creative_mode,
            "strict_ai": request.strict_ai,
        }
        fields = dict(field_map)
        review_report = review_lesson_quality(fields, context)
        self._mark(
            "teaching_reviewer",
            "教研组长 Agent",
            "done",
            f"预审完成，评分 {review_report.score}。",
        )

        fields, revision_backend = revise_lesson_after_review(fields, review_report, context, template_fields)
        field_map.update(fields)
        self._mark("lesson_reviser", "二次修订 Agent", "done", "已根据审阅意见完成修订。")

        return {
            "fields": field_map,
            "template_fields": template_analysis["mapped_fields"],
            "template_analysis": template_analysis,
            "template_id": template_id,
            "generation_backend": generation_backend,
            "revision_backend": revision_backend,
            "review_report": review_report.to_dict(),
            "knowledge_report": knowledge_context.to_dict(),
            "quality_controls": {
                "strict_ai": request.strict_ai,
                "creative_mode": request.creative_mode,
                "anti_repetition_used": bool(anti_repetition_context),
                "few_shot_used": bool(few_shot_examples),
            },
            "workflow_trace": [event.to_dict() for event in self.trace],
            "workflow_schema": build_workflow_schema(),
        }

    def _build_anti_repetition_context(self, request: LessonRequest) -> str:
        if not self.history_db:
            return ""
        try:
            documents = HistoryStore(self.history_db).find_similar_documents(
                subject=request.subject,
                grade=request.grade,
                title=request.title,
                limit=3,
            )
        except Exception:
            return ""
        if not documents:
            return ""
        chunks = []
        for index, item in enumerate(documents, start=1):
            chunks.append(
                "\n".join(
                    [
                        f"相似教案{index}：{item.get('grade','')}{item.get('subject','')}《{item.get('lesson_title','')}》",
                        f"课型/风格：{item.get('class_type','')} / {item.get('teaching_style','')}",
                        f"教学过程摘要：{item.get('teaching_process','')}",
                        f"作业摘要：{item.get('homework','')}",
                        f"板书摘要：{item.get('blackboard_design','')}",
                    ]
                )
            )
        return "\n\n".join(chunks)

    def export_document(
        self,
        fields: dict,
        template_path: Path,
        output_dir: Path,
        preview_dir: Path,
        repeat_fill_mode: str | None = None,
    ) -> dict:
        title = sanitize_lesson_title(str(fields.get("lesson_title") or ""), "", str(fields.get("title") or "教案"))
        fields = dict(fields)
        fields["lesson_title"] = title
        fields = normalize_lesson_field_aliases(fields)
        grade = str(fields.get("grade") or "年级")
        subject = str(fields.get("subject") or "学科")
        safe_title = _safe_filename(f"{grade}-{subject}-{title}-教案")
        output_name = f"{safe_title}-{time.strftime('%Y%m%d-%H%M%S')}.docx"
        output_path = output_dir / output_name

        actual_repeat_mode = repeat_fill_mode or ("all" if _is_system_template_path(template_path) else "first_only")
        fill_report = fill_docx_template(template_path, fields, output_path, repeat_fill_mode=actual_repeat_mode)
        _enhance_system_template_docx(output_path, fields, template_path)
        template_analysis = analyze_template(template_path)
        teacher_diagnostic_report = build_teacher_diagnostic_report(
            template_analysis=template_analysis,
            fill_report=fill_report.to_dict(),
            evaluation_report={"passed": not fill_report.errors},
            fields=fields,
        ).to_dict()
        preview_pdf = render_docx_pdf_preview(output_path, preview_dir)
        preview_url = f"/preview/{quote(preview_pdf.name)}" if preview_pdf else None
        self._mark("doc_renderer", "Word 渲染器", "done", "已按原 Word 模板写入字段并生成下载文件。")

        return {
            "output_name": output_name,
            "download_url": f"/download/{quote(output_name)}",
            "preview_url": preview_url,
            "template_analysis": template_analysis,
            "fill_report": fill_report.to_dict(),
            "teacher_diagnostic_report": teacher_diagnostic_report,
            "workflow_trace": [event.to_dict() for event in self.trace],
        }

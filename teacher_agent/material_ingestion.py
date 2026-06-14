from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader

from .docx_security import DocxSecurityError, validate_docx_bytes


@dataclass
class MaterialExtraction:
    text: str
    source_name: str
    warnings: list[str]

    def to_dict(self) -> dict:
        return {
            "text": self.text,
            "source_name": self.source_name,
            "warnings": self.warnings,
            "char_count": len(self.text),
        }


def extract_material_from_upload(filename: str, file_obj: BinaryIO) -> MaterialExtraction:
    raw = file_obj.read()
    return extract_material_bytes(filename, raw)


def extract_material_bytes(filename: str, raw: bytes) -> MaterialExtraction:
    suffix = Path(filename or "").suffix.lower()
    warnings: list[str] = []
    text = ""

    if suffix in {".txt", ".md"}:
        text = _decode_text(raw)
    elif suffix == ".docx":
        text = _extract_docx_text(raw, warnings)
    elif suffix == ".pdf":
        text = _extract_pdf_text_best_effort(raw, warnings)
    else:
        warnings.append("暂不支持该教材文件类型，已忽略文件内容。")

    return MaterialExtraction(text=_normalize_material_text(text), source_name=filename, warnings=warnings)


def merge_material_text(typed_material: str, extracted: MaterialExtraction | None) -> str:
    typed = (typed_material or "").strip()
    if not extracted or not extracted.text.strip():
        return typed
    if typed:
        return f"{typed}\n\n# 上传教材资料：{extracted.source_name}\n{extracted.text}"
    return extracted.text


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_docx_text(raw: bytes, warnings: list[str]) -> str:
    try:
        validate_docx_bytes(raw, source_name="上传的 Word 教材")
        document = Document(BytesIO(raw))
    except DocxSecurityError as exc:
        warnings.append(str(exc))
        return ""
    except Exception as exc:
        warnings.append(f"Word 教材读取失败：{exc}")
        return ""

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf_text_best_effort(raw: bytes, warnings: list[str]) -> str:
    try:
        reader = PdfReader(BytesIO(raw), strict=False)
    except Exception as exc:
        warnings.append(f"PDF 教材读取失败：{exc}")
        return ""

    if reader.is_encrypted:
        warnings.append("PDF 教材已加密，无法读取；请上传未加密版本。")
        return ""

    parts: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:
            warnings.append(f"PDF 第 {page_number} 页文本提取失败：{exc}")
            continue
        if text:
            parts.append(text)

    if not parts:
        warnings.append("PDF 中未检测到可提取文本；如果是扫描版，请先进行 OCR 后再上传。")
    return "\n\n".join(parts)


def _normalize_material_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    compact: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank:
                compact.append("")
            blank = True
            continue
        compact.append(line)
        blank = False
    return "\n".join(compact).strip()
